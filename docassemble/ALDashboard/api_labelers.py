"""
Flask endpoints for the DOCX and PDF labeler tools.

These provide interactive browser-based interfaces for:
- al/docx-labeler: Add Jinja2 labels to DOCX templates
- al/pdf-labeler: Add/edit PDF form fields

Both tools use AI to suggest labels and follow AssemblyLine conventions.
"""

import base64
import copy
import io
import json
import os
import re
import tempfile
import time
import uuid
from contextlib import contextmanager
from urllib.parse import quote, urlsplit
from typing import Any, Dict, List, Optional, Set

from flask import Response, jsonify, request

try:
    from flask_cors import cross_origin
except ImportError:  # pragma: no cover - exercised by subprocess import tests

    def cross_origin(*args: Any, **kwargs: Any):
        if args and callable(args[0]) and len(args) == 1 and not kwargs:
            return args[0]

        def decorator(func: Any) -> Any:
            return func

        return decorator


try:
    from flask_login import current_user
except ImportError:  # pragma: no cover - exercised by subprocess import tests
    current_user = type(
        "AnonymousCurrentUser",
        (),
        {"is_authenticated": False, "id": None, "email": None},
    )()

from docassemble.base.config import daconfig
import docassemble.base.functions
from docassemble.base.util import log
from docassemble.webapp.app_object import app, csrf
from docassemble.webapp.server import api_verify, jsonify_with_status, r
from docassemble.webapp.worker_common import workerapp

from .api_dashboard_utils import (
    DashboardAPIValidationError,
    _extract_repair_options,
    _format_pdf_fields_for_ui_payload,
    _validate_upload_size,
    coerce_async_flag,
    decode_base64_content,
    docx_labeler_suggest_payload_from_options,
    merge_raw_options,
    pdf_fields_detect_payload_from_options,
    pdf_fields_relabel_payload_from_options,
    parse_bool,
    validate_docx_payload_from_options,
)
from .pdf_export_utils import build_pdf_export_fields_per_page

__all__ = []

LABELER_BASE_PATH = "/al"

OPENAI_LABELER_MODELS = [
    "gpt-5-mini",
    "gpt-5",
    "gpt-4.1-mini",
    "gpt-4.1",
    "gpt-5-nano",
]
GEMINI_LABELER_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5",
    "gemini-3-pro",
]
CLAUDE_LABELER_MODELS = [
    "claude-4.5-sonnet",
    "claude-4.6-sonnet",
    "claude-4.6-opus",
]
LABELER_PROVIDER_MODEL_SETS = {
    "openai": OPENAI_LABELER_MODELS,
    "gemini": GEMINI_LABELER_MODELS,
    "claude": CLAUDE_LABELER_MODELS,
}
LABELER_MODEL_SET_PRIORITY = [
    OPENAI_LABELER_MODELS,
    GEMINI_LABELER_MODELS,
    CLAUDE_LABELER_MODELS,
]
LABELER_DEFAULT_MODEL = "gpt-5-mini"
LABELER_JOB_KEY_PREFIX = "da:aldashboard:labeler-job:"
LABELER_JOB_EXPIRE_SECONDS = 24 * 60 * 60
ASYNC_CELERY_MODULE = "docassemble.ALDashboard.api_dashboard_worker"


def _sanitize_checkbox_export_value(raw_value: Any) -> str:
    """Normalize a checkbox export value into a safe PDF name token.

    Args:
        raw_value: Raw export value supplied by the browser UI.

    Returns:
        str: A sanitized export token suitable for pikepdf ``Name`` objects.
    """
    token = str(raw_value or "").strip() or "Yes"
    token = re.sub(r"[^A-Za-z0-9_.-]+", "_", token).strip("_")
    return token or "Yes"


def _parse_optional_json_field(raw_value: Any, *, field_name: str) -> Optional[Any]:
    """Parse an optional JSON field that may be serialized as a string."""
    if raw_value is None:
        return None
    if isinstance(raw_value, str):
        text = raw_value.strip()
        if not text:
            return None
        try:
            return json.loads(text)
        except json.JSONDecodeError as exc:
            raise DashboardAPIValidationError(
                f"{field_name} must be valid JSON."
            ) from exc
    return raw_value


def _looks_like_name_email_address_phone_field(field_name: Any) -> bool:
    """Return True when a field name should keep auto-size enabled."""
    return bool(
        re.search(
            r"(name|address|street|city|state|zip|postal|phone|phone_number|email|cell)",
            str(field_name or ""),
            flags=re.IGNORECASE,
        )
    )


def _get_named_pdf_parent(field_obj: Any) -> Optional[Any]:
    """Walk up a PDF widget tree until a named field container is found.

    Args:
        field_obj: The current pikepdf annotation or parent object.

    Returns:
        Optional[Any]: The nearest object exposing ``T``, or ``None`` if missing.
    """
    if hasattr(field_obj, "T"):
        return field_obj
    if hasattr(field_obj, "Parent"):
        return _get_named_pdf_parent(field_obj.Parent)
    return None


def _rename_checkbox_export_state(
    field_obj: Any, export_value: str, pikepdf_module: Any
) -> None:
    """Rename checkbox appearance states to match a custom export value.

    Args:
        field_obj: The widget or parent field object to update.
        export_value: The desired export token for checked state.
        pikepdf_module: Imported ``pikepdf`` module used to create ``Name`` values.
    """
    desired_state = pikepdf_module.Name(
        "/" + _sanitize_checkbox_export_value(export_value)
    )
    current_state = pikepdf_module.Name("/Yes")
    if hasattr(field_obj, "AP"):
        for appearance_key in ("/N", "/D", "/R"):
            if appearance_key not in field_obj.AP:
                continue
            appearance_dict = field_obj.AP[appearance_key]
            if (
                current_state in appearance_dict
                and desired_state not in appearance_dict
            ):
                appearance_dict[desired_state] = appearance_dict[current_state]
                del appearance_dict[current_state]
    for attr_name in ("V", "AS", "DV"):
        if (
            hasattr(field_obj, attr_name)
            and getattr(field_obj, attr_name) == current_state
        ):
            setattr(field_obj, attr_name, desired_state)


def _apply_checkbox_export_values(
    pdf_path: str, value_by_field_name: Dict[str, str]
) -> None:
    """Apply custom checkbox export values to a written PDF in place.

    Args:
        pdf_path: Path to the PDF file to modify.
        value_by_field_name: Mapping of field names to export values.
    """
    desired = {
        str(name): _sanitize_checkbox_export_value(value)
        for name, value in value_by_field_name.items()
        if str(name).strip() and _sanitize_checkbox_export_value(value) != "Yes"
    }
    if not desired:
        return

    import pikepdf

    with pikepdf.open(pdf_path, allow_overwriting_input=True) as pdf:
        for page in pdf.pages:
            if "/Annots" not in page:
                continue
            for annot in page.Annots:  # type: ignore[attr-defined]
                try:
                    if annot.Type != "/Annot" or annot.Subtype != "/Widget":
                        continue
                    named_parent = _get_named_pdf_parent(annot)
                    if not named_parent or not hasattr(named_parent, "T"):
                        continue
                    field_name = str(named_parent.T)
                    if field_name not in desired:
                        continue
                    export_value = desired[field_name]
                    _rename_checkbox_export_state(named_parent, export_value, pikepdf)
                    if named_parent is not annot:
                        _rename_checkbox_export_state(annot, export_value, pikepdf)
                except Exception:  # nosec B112
                    continue
        pdf.save(pdf_path)


def _collect_fields_with_explicit_background(
    fields_data: List[Dict[str, Any]],
) -> Set[str]:
    """Collect field names where the client explicitly set a background color.

    Args:
        fields_data: Field definitions submitted by the labeler UI.

    Returns:
        Set[str]: Field names whose background color should be preserved.
    """
    explicit: Set[str] = set()
    for field in fields_data:
        try:
            name = str(field.get("name", "")).strip()
            background = field.get("backgroundColor")
            if name and isinstance(background, str) and background.strip():
                explicit.add(name)
        except Exception:  # nosec B112
            continue
    return explicit


def _apply_pdf_field_visual_defaults(
    pdf_path: str,
    *,
    explicit_background_fields: Optional[Set[str]] = None,
    preserve_button_appearances: bool = False,
) -> None:
    """Enforce default transparent field backgrounds and no borders.

    Fields with explicit background colors are preserved. For all other fields,
    this clears widget background appearance metadata so viewers regenerate a
    transparent background, and it removes border drawing.

    Args:
        pdf_path: Path to the written PDF file.
        explicit_background_fields: Field names whose background should be kept.
        preserve_button_appearances: When True, the appearance stream (``/AP``)
            is not deleted for button fields (checkboxes, radio buttons, push
            buttons).  Use this when the fields were copied from an existing PDF
            whose button appearances are already correct, rather than freshly
            generated by a library that may embed non-transparent backgrounds.
    """
    import pikepdf

    explicit = {str(name).strip() for name in (explicit_background_fields or set())}

    def _set_no_border(obj: Any) -> None:
        obj["/Border"] = pikepdf.Array([0, 0, 0])
        bs = obj.get("/BS")
        if not isinstance(bs, pikepdf.Dictionary):
            bs = pikepdf.Dictionary()
        bs["/W"] = 0
        obj["/BS"] = bs
        mk = obj.get("/MK")
        if isinstance(mk, pikepdf.Dictionary) and "/BC" in mk:
            del mk["/BC"]
            if len(mk) == 0:
                del obj["/MK"]

    def _clear_background(obj: Any) -> None:
        mk = obj.get("/MK")
        if isinstance(mk, pikepdf.Dictionary) and "/BG" in mk:
            del mk["/BG"]
            if len(mk) == 0:
                del obj["/MK"]

    def _is_button_widget(named_parent: Optional[Any]) -> bool:
        if named_parent is None or not hasattr(named_parent, "get"):
            return False
        return str(named_parent.get("/FT", "")) == "/Btn"

    with pikepdf.open(pdf_path, allow_overwriting_input=True) as pdf:
        acroform = pdf.Root.get("/AcroForm")
        if isinstance(acroform, pikepdf.Dictionary):
            acroform["/NeedAppearances"] = True

        for page in pdf.pages:
            if "/Annots" not in page:
                continue
            for annot in page.Annots:  # type: ignore[attr-defined]
                try:
                    if annot.Type != "/Annot" or annot.Subtype != "/Widget":
                        continue

                    named_parent = _get_named_pdf_parent(annot)
                    field_name = ""
                    if named_parent and hasattr(named_parent, "T"):
                        field_name = str(named_parent.T).strip()

                    _set_no_border(annot)
                    if named_parent and named_parent is not annot:
                        _set_no_border(named_parent)

                    if field_name in explicit:
                        continue

                    _clear_background(annot)
                    if named_parent and named_parent is not annot:
                        _clear_background(named_parent)

                    # Remove widget appearance so viewers can regenerate using
                    # transparent/default settings instead of reportlab defaults.
                    # Button fields (checkboxes/radios) copied from existing PDFs
                    # must keep their /AP stream because it encodes checked vs.
                    # unchecked visual states that PDF viewers cannot reliably
                    # reconstruct on their own.
                    if "/AP" in annot:
                        if not preserve_button_appearances or not _is_button_widget(
                            named_parent
                        ):
                            del annot["/AP"]
                except Exception:  # nosec B112
                    continue

        pdf.save(pdf_path)


def _normalize_provider_family(family_name: Optional[str]) -> str:
    """Normalize model provider aliases into the labeler provider keys.

    Args:
        family_name: Raw provider family name.

    Returns:
        str: One of ``openai``, ``gemini``, or ``claude``.
    """
    family = str(family_name or "").strip().lower()
    if family in {"google", "gemini"}:
        return "gemini"
    if family in {"anthropic", "claude"}:
        return "claude"
    return "openai"


def _build_labeler_model_catalog() -> Dict[str, Any]:
    """Build model metadata for labeler UIs using ALToolbox llms helpers."""
    default_model = LABELER_DEFAULT_MODEL
    provider_family = "openai"
    recommended_models = list(OPENAI_LABELER_MODELS)
    available_models: List[str] = []

    try:
        from docassemble.ALToolbox.llms import (
            detect_model_family,
            get_default_model,
            get_first_available_model_set,
            list_available_models,
        )

        available_models = list_available_models()
        selected_set = get_first_available_model_set(
            LABELER_MODEL_SET_PRIORITY,
            require_full_set=False,
            return_partial_if_needed=True,
            fallback_to_first_small_model=True,
        )
        if selected_set:
            recommended_models = selected_set
            provider_family = _normalize_provider_family(
                detect_model_family(selected_set[0])
            )
        else:
            medium_default = get_default_model("medium")
            provider_family = _normalize_provider_family(
                detect_model_family(medium_default)
            )
            recommended_models = list(
                LABELER_PROVIDER_MODEL_SETS.get(provider_family, OPENAI_LABELER_MODELS)
            )

        if provider_family == "openai":
            default_model = "gpt-5-mini"
        elif recommended_models:
            default_model = recommended_models[0]
        else:
            default_model = get_default_model("medium")

        if default_model not in recommended_models and default_model:
            recommended_models = [default_model] + recommended_models
    except Exception as exc:
        log(
            f"ALDashboard: failed to build model catalog from ALToolbox.llms; using fallback list ({exc!r})",
            "warning",
        )

    return {
        "default_model": default_model or LABELER_DEFAULT_MODEL,
        "recommended_models": recommended_models,
        "available_models": available_models,
        "provider_family": provider_family,
    }


def _labeler_session_identity() -> Dict[str, Any]:
    """Return session identity details for browser users."""
    try:
        if current_user.is_authenticated:
            email_value = getattr(current_user, "email", None)
            if email_value is not None:
                email_value = str(email_value)
            return {
                "is_authenticated": True,
                "email": email_value,
                "user_id": getattr(current_user, "id", None),
            }
    except Exception:  # nosec B110
        pass
    return {"is_authenticated": False, "email": None, "user_id": None}


def _labeler_ai_auth_check() -> bool:
    """AI features require API key auth or a logged-in browser user."""
    if api_verify():
        return True
    identity = _labeler_session_identity()
    return bool(identity.get("is_authenticated"))


def _safe_labeler_return_target(raw_target: Optional[str]) -> Optional[str]:
    """Allow only same-origin or relative return targets for auth redirects."""
    target = str(raw_target or "").strip()
    if not target:
        return None

    parsed = urlsplit(target)
    if parsed.scheme or parsed.netloc:
        if parsed.netloc != request.host:
            return None
        path = parsed.path or "/"
    else:
        path = parsed.path or target

    if not path.startswith("/") or path.startswith("//"):
        return None

    if parsed.query:
        return f"{path}?{parsed.query}"
    return path


def _labeler_auth_return_target() -> str:
    """Resolve the page the labeler should return to after auth."""
    explicit_target = _safe_labeler_return_target(request.args.get("next"))
    if explicit_target:
        return explicit_target

    referer_target = _safe_labeler_return_target(request.headers.get("Referer"))
    if referer_target:
        return referer_target

    return LABELER_BASE_PATH


def _labeler_playground_auth_check() -> bool:
    """Check whether the current browser session can access playground data.

    Returns:
        bool: ``True`` when the current user is authenticated.
    """
    try:
        return bool(current_user.is_authenticated)
    except Exception:
        return False


def _playground_auth_fail(request_id: str):
    """Build the standard JSON auth failure for playground endpoints.

    Args:
        request_id: Correlation ID for the current request.

    Returns:
        Response: A 401 JSON response describing the auth error.
    """
    return jsonify_with_status(
        {
            "success": False,
            "request_id": request_id,
            "error": {
                "type": "auth_error",
                "message": "Login required for Playground interview access.",
            },
        },
        401,
    )


def _normalize_playground_project(project: Optional[str]) -> str:
    """Validate and normalize a Playground project name.

    Args:
        project: Raw project name from the request.

    Returns:
        str: A normalized Playground project name.
    """
    value = str(project or "default").strip() or "default"
    if "/" in value or "\\" in value or value.startswith("."):
        raise DashboardAPIValidationError(
            "Invalid Playground project.", status_code=400
        )
    return value


def _normalize_playground_filename(filename: Optional[str]) -> str:
    """Validate and normalize a Playground YAML filename.

    Args:
        filename: Raw filename from the request.

    Returns:
        str: A sanitized YAML filename.
    """
    value = os.path.basename(str(filename or "").strip())
    if not value or value in {".", ".."}:
        raise DashboardAPIValidationError(
            "Playground YAML filename is required.", status_code=400
        )
    if not value.lower().endswith((".yml", ".yaml")):
        raise DashboardAPIValidationError(
            "Playground file must be a YAML interview.", status_code=400
        )
    return value


@contextmanager
def _playground_user_context(user_id: int):
    """Temporarily impersonate a Playground user in docassemble thread state.

    Args:
        user_id: The authenticated user ID to expose to Playground helpers.

    Yields:
        None: Control returns to the caller while thread context is overridden.
    """
    original_info = copy.deepcopy(
        getattr(docassemble.base.functions.this_thread, "current_info", {}) or {}
    )
    current_info = copy.deepcopy(original_info)
    current_info.setdefault("user", {})
    current_info["user"].update({"is_anonymous": False, "theid": user_id})
    docassemble.base.functions.this_thread.current_info = current_info
    try:
        yield
    finally:
        docassemble.base.functions.this_thread.current_info = original_info


def _list_playground_projects() -> List[str]:
    """List available Playground projects for the current user.

    Returns:
        List[str]: Sorted Playground project names, always including ``default``.
    """
    from docassemble.webapp.files import SavedFile

    uid = getattr(current_user, "id", None)
    if uid is None:
        return ["default"]
    playground = SavedFile(uid, fix=False, section="playground")
    projects = playground.list_of_dirs() or []
    projects = [proj for proj in projects if isinstance(proj, str) and proj]
    if "default" not in projects:
        projects.append("default")
    return sorted(set(projects))


def _list_playground_yaml_files(project: str) -> List[Dict[str, str]]:
    """List YAML interviews available in a Playground project.

    Args:
        project: Playground project name.

    Returns:
        List[Dict[str, str]]: Label/value entries for YAML files in the project.
    """
    from docassemble.webapp.playground import Playground

    uid = getattr(current_user, "id", None)
    if uid is None:
        return []
    with _playground_user_context(uid):
        playground = Playground(project=project)
        return [
            {"filename": filename, "label": filename}
            for filename in playground.file_list
            if isinstance(filename, str)
            and filename.lower().endswith((".yml", ".yaml"))
        ]


def _get_playground_variable_info(project: str, filename: str) -> Dict[str, Any]:
    """Extract variable names from a Playground YAML interview.

    Args:
        project: Playground project containing the YAML file.
        filename: YAML filename to inspect.

    Returns:
        Dict[str, Any]: Variable metadata including all and top-level names.
    """
    from docassemble.webapp.playground import Playground

    uid = getattr(current_user, "id", None)
    if uid is None:
        raise DashboardAPIValidationError(
            "Login required for Playground interview access.", status_code=401
        )
    with _playground_user_context(uid):
        pg = Playground(project=project)
        if filename not in pg.file_list:
            raise DashboardAPIValidationError(
                "Selected Playground YAML file was not found.", status_code=404
            )
        variable_info = pg.variables_from_file(filename)
    if not isinstance(variable_info, dict):
        variable_info = {}
    all_names = sorted(
        str(name).strip()
        for name in (variable_info.get("all_names_reduced") or [])
        if str(name).strip()
    )
    top_level_names = sorted(
        {name.split(".", 1)[0].split("[", 1)[0] for name in all_names if name}
    )
    return {
        "project": project,
        "filename": filename,
        "all_names": all_names,
        "top_level_names": top_level_names,
    }


def _normalize_installed_package_name(package_name: Optional[str]) -> str:
    """Validate a docassemble package name used for installed interviews.

    Args:
        package_name: Raw package name from the request.

    Returns:
        str: A validated ``docassemble.*`` package name.
    """
    value = str(package_name or "").strip()
    if not value or "/" in value or "\\" in value or ":" in value:
        raise DashboardAPIValidationError(
            "Installed interview package name is required.", status_code=400
        )
    if not value.startswith("docassemble."):
        raise DashboardAPIValidationError(
            "Installed interview package must start with docassemble.",
            status_code=400,
        )
    return value


def _normalize_installed_interview_path(interview_path: Optional[str]) -> str:
    """Validate a ``package:file`` interview reference for installed YAML.

    Args:
        interview_path: Raw installed interview path from the request.

    Returns:
        str: A normalized ``package:file`` reference.
    """
    value = str(interview_path or "").strip()
    if ":" not in value:
        raise DashboardAPIValidationError(
            "Installed interview path must be in package:file format.",
            status_code=400,
        )
    package_name, filename = value.split(":", 1)
    package_name = _normalize_installed_package_name(package_name)
    filename = _normalize_playground_filename(filename)
    return f"{package_name}:{filename}"


def _extract_variable_names_from_var_json(variable_json: Any) -> Dict[str, List[str]]:
    """Normalize docassemble variable-inspection JSON into flat name lists.

    Args:
        variable_json: Raw variable metadata returned by docassemble helpers.

    Returns:
        Dict[str, List[str]]: All variable names and their top-level roots.
    """
    if not isinstance(variable_json, dict):
        variable_json = {}

    names: List[str] = []
    seen = set()

    def add_name(entry: Any) -> None:
        """Extract one candidate variable name from a heterogeneous JSON entry.

        Args:
            entry: A raw item from the variable-inspection payload.
        """
        candidate = None
        if isinstance(entry, str):
            candidate = entry
        elif isinstance(entry, dict):
            for key in ("name", "variable", "var_name", "varName"):
                if isinstance(entry.get(key), str):
                    candidate = entry.get(key)
                    break
            if candidate is None and isinstance(entry.get(0), str):
                candidate = entry.get(0)
        elif isinstance(entry, (list, tuple)) and entry:
            if isinstance(entry[0], str):
                candidate = entry[0]
        if candidate is None:
            return
        candidate = str(candidate).strip()
        if not candidate or candidate in seen:
            return
        seen.add(candidate)
        names.append(candidate)

    for key in ("var_list", "undefined_names"):
        values = variable_json.get(key)
        if isinstance(values, list):
            for entry in values:
                add_name(entry)

    names.sort()
    top_level_names = sorted(
        {name.split(".", 1)[0].split("[", 1)[0] for name in names if name}
    )
    return {"all_names": names, "top_level_names": top_level_names}


def _list_installed_interview_packages() -> List[str]:
    """List installed docassemble packages that contain YAML interviews.

    Returns:
        List[str]: Sorted package names that expose question files.
    """
    from .aldashboard import list_question_files_in_docassemble_packages

    package_map = list_question_files_in_docassemble_packages()
    return sorted(
        package_name
        for package_name, filenames in package_map.items()
        if isinstance(package_name, str) and filenames
    )


def _list_installed_interview_files(package_name: str) -> List[Dict[str, str]]:
    """List YAML interview files for an installed docassemble package.

    Args:
        package_name: Installed package name.

    Returns:
        List[Dict[str, str]]: Label/value entries for YAML interviews.
    """
    from .aldashboard import list_question_files_in_docassemble_packages

    package_map = list_question_files_in_docassemble_packages()
    filenames = package_map.get(package_name, [])
    return [
        {
            "filename": filename,
            "label": filename,
            "interview_path": f"{package_name}:{filename}",
        }
        for filename in sorted(filenames)
        if isinstance(filename, str) and filename.lower().endswith((".yml", ".yaml"))
    ]


def _get_installed_interview_variable_info(interview_path: str) -> Dict[str, Any]:
    """Inspect an installed interview and return its discovered variable names.

    Args:
        interview_path: Normalized ``package:file`` interview reference.

    Returns:
        Dict[str, Any]: Variable metadata including all and top-level names.
    """
    from docassemble.base.parse import InterviewStatus, interview_source_from_string
    from docassemble.webapp.server import current_info, get_vars_in_use

    normalized_path = _normalize_installed_interview_path(interview_path)
    try:
        interview_source = interview_source_from_string(normalized_path, testing=True)
        interview = interview_source.get_interview()
    except Exception as exc:
        raise DashboardAPIValidationError(
            f"Unable to load installed interview: {exc}", status_code=400
        ) from exc

    device_id = request.cookies.get("ds", None)
    interview_status = InterviewStatus(
        current_info=current_info(
            yaml=normalized_path, req=request, action=None, device_id=device_id
        )
    )
    try:
        variable_json, vocab_list, vocab_dict, ac_list = (
            get_vars_in_use(  # pylint: disable=unused-variable
                interview,
                interview_status,
                debug_mode=False,
                return_json=True,
                use_playground=False,
                current_project="default",
            )
        )
    except Exception as exc:
        raise DashboardAPIValidationError(
            f"Unable to analyze installed interview variables: {exc}",
            status_code=400,
        ) from exc

    parsed_names = _extract_variable_names_from_var_json(variable_json)
    return {
        "interview_path": normalized_path,
        "all_names": parsed_names["all_names"],
        "top_level_names": parsed_names["top_level_names"],
    }


def _get_static_content(filename: str) -> str:
    """Read a static file from the data/static directory."""
    import importlib.resources

    try:
        ref = (
            importlib.resources.files("docassemble.ALDashboard")
            / "data"
            / "static"
            / filename
        )
        with importlib.resources.as_file(ref) as path:
            if path.exists():
                return path.read_text(encoding="utf-8")
    except Exception:  # nosec B110
        pass
    return ""


def _get_template_content(filename: str) -> str:
    """Read a template file from the data/templates directory."""
    import importlib.resources

    try:
        ref = (
            importlib.resources.files("docassemble.ALDashboard")
            / "data"
            / "templates"
            / filename
        )
        with importlib.resources.as_file(ref) as path:
            if path.exists():
                return path.read_text(encoding="utf-8")
    except Exception:  # nosec B110
        pass
    return ""


def _render_template_content(
    filename: str, *, bootstrap_data: Optional[Dict[str, Any]] = None
) -> str:
    """Read a template and inject bootstrap JSON when requested."""
    html_content = _get_template_content(filename)
    if not html_content:
        return ""
    if bootstrap_data is None:
        return html_content
    bootstrap_json = json.dumps(bootstrap_data, sort_keys=True)
    bootstrap_json = (
        bootstrap_json.replace("<", "\\u003c")
        .replace(">", "\\u003e")
        .replace("&", "\\u0026")
        .replace("\u2028", "\\u2028")
        .replace("\u2029", "\\u2029")
    )
    return html_content.replace(
        "__LABELER_BOOTSTRAP_JSON__",
        bootstrap_json,
    )


def _parse_initial_playground_source(
    project: Optional[str],
    filename: Optional[str],
    *,
    allowed_extensions: tuple[str, ...],
) -> Dict[str, str]:
    """Normalize optional labeler page query params for initial template load.

    Args:
        project: Optional ``project`` query parameter.
        filename: Optional ``filename`` query parameter.
        allowed_extensions: Allowed filename extensions for this labeler page.

    Returns:
        Dict[str, str]: Sanitized initial source keys for UI bootstrap.
    """
    raw_project = str(project or "").strip()
    raw_filename = str(filename or "").strip()
    if not raw_project and not raw_filename:
        return {}

    normalized_project = (
        _normalize_playground_project(raw_project) if raw_project else "default"
    )
    if not raw_filename:
        return {"project": normalized_project}

    normalized_filename = _normalize_template_filename(
        raw_filename, allowed_extensions=allowed_extensions
    )
    return {"project": normalized_project, "filename": normalized_filename}


def _labeler_initial_playground_source_from_request(
    *, allowed_extensions: tuple[str, ...]
) -> Dict[str, str]:
    """Read and validate optional ``project``/``filename`` query params.

    Invalid query input is ignored so the labeler UI still loads.
    """
    try:
        return _parse_initial_playground_source(
            request.args.get("project"),
            request.args.get("filename"),
            allowed_extensions=allowed_extensions,
        )
    except DashboardAPIValidationError as exc:
        log(
            "ALDashboard: Ignoring invalid labeler query params "
            f"project={request.args.get('project')!r} filename={request.args.get('filename')!r} "
            f"({exc.message})",
            "warning",
        )
        return {}


def _build_docx_labeler_bootstrap() -> Dict[str, Any]:
    """Build bootstrap data for the DOCX labeler page."""
    return {
        "apiBasePath": LABELER_BASE_PATH,
        "initialPlaygroundSource": _labeler_initial_playground_source_from_request(
            allowed_extensions=(".docx",)
        ),
    }


def _build_pdf_labeler_bootstrap() -> Dict[str, Any]:
    """Build bootstrap data for the PDF labeler page."""
    from .labeler_config import get_pdf_labeler_ui_config

    pdf_ui_config = get_pdf_labeler_ui_config()
    return {
        "apiBasePath": LABELER_BASE_PATH,
        "initialPlaygroundSource": _labeler_initial_playground_source_from_request(
            allowed_extensions=(".pdf",)
        ),
        "branding": pdf_ui_config.get("branding", {}),
        "pdf": {
            "fieldNameLibrary": pdf_ui_config.get("field_name_library", {}),
        },
    }


def _auth_fail(request_id: str):
    """Build the standard JSON auth failure for labeler API endpoints.

    Args:
        request_id: Correlation ID for the current request.

    Returns:
        Response: A 403 JSON response describing the auth failure.
    """
    return jsonify_with_status(
        {
            "success": False,
            "request_id": request_id,
            "error": {"type": "auth_error", "message": "Access denied."},
        },
        403,
    )


def _ai_auth_fail(request_id: str):
    """Build the standard JSON auth failure for AI-enabled labeler endpoints.

    Args:
        request_id: Correlation ID for the current request.

    Returns:
        Response: A 401 JSON response describing the AI auth requirement.
    """
    return jsonify_with_status(
        {
            "success": False,
            "request_id": request_id,
            "error": {
                "type": "auth_error",
                "message": "Login required for AI features in the labeler.",
            },
        },
        401,
    )


def _labeler_async_is_configured() -> bool:
    """Check whether the labeler Celery worker module is configured.

    Returns:
        bool: ``True`` when async labeler jobs are enabled in docassemble config.
    """
    celery_modules = daconfig.get("celery modules", []) or []
    return ASYNC_CELERY_MODULE in celery_modules


def _labeler_job_key(job_id: str) -> str:
    """Build the Redis key used to store labeler job metadata.

    Args:
        job_id: Public labeler job identifier.

    Returns:
        str: Redis key for the job mapping.
    """
    return LABELER_JOB_KEY_PREFIX + job_id


def _store_labeler_job_mapping(
    job_id: str, task_id: str, extra: Optional[Dict[str, Any]] = None
) -> None:
    """Persist a labeler job-to-task mapping in Redis.

    Args:
        job_id: Public labeler job identifier.
        task_id: Celery task ID backing the job.
        extra: Optional extra metadata to store alongside the mapping.
    """
    payload = {"id": task_id, "created_at": time.time()}
    if extra:
        payload.update(extra)
    pipe = r.pipeline()
    pipe.set(_labeler_job_key(job_id), json.dumps(payload))
    pipe.expire(_labeler_job_key(job_id), LABELER_JOB_EXPIRE_SECONDS)
    pipe.execute()


def _fetch_labeler_job_mapping(job_id: str) -> Optional[Dict[str, Any]]:
    """Load stored labeler job metadata from Redis.

    Args:
        job_id: Public labeler job identifier.

    Returns:
        Optional[Dict[str, Any]]: Stored job metadata, or ``None`` if unavailable.
    """
    raw = r.get(_labeler_job_key(job_id))
    if raw is None:
        return None
    try:
        return json.loads(raw.decode())
    except Exception:
        return None


def _normalize_labeler_result_for_json(value: Any) -> Any:
    """Convert async task result objects into JSON-serializable structures.

    Args:
        value: Raw task result value.

    Returns:
        Any: A JSON-safe version of the result.
    """
    if isinstance(value, dict):
        return {
            str(key): _normalize_labeler_result_for_json(item)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [_normalize_labeler_result_for_json(item) for item in value]
    if isinstance(value, tuple):
        return [_normalize_labeler_result_for_json(item) for item in value]
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if hasattr(value, "value"):
        inner_value = getattr(value, "value")
        if inner_value is not None:
            return _normalize_labeler_result_for_json(inner_value)
    if hasattr(value, "__dict__"):
        return _normalize_labeler_result_for_json(vars(value))
    return str(value)


def _format_pdf_auto_detect_labeler_data(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize shared PDF detect payloads into the labeler UI response shape."""
    fields = payload.get("positioned_fields")
    if not isinstance(fields, list):
        fields = []
    return {
        "filename": str(payload.get("input_filename") or "upload.pdf"),
        "page_count": payload.get("page_count"),
        "fields": fields,
        "pdf_base64": payload.get("pdf_base64"),
    }


def _format_pdf_relabel_labeler_data(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize shared PDF relabel payloads into the labeler UI response shape."""
    return {
        "filename": str(
            payload.get("output_filename")
            or payload.get("input_filename")
            or "upload.pdf"
        ),
        "output_filename": payload.get("output_filename"),
        "fields": payload.get("fields", []),
        "fields_old": payload.get("fields_old", []),
        "pdf_base64": payload.get("pdf_base64"),
    }


def _queue_labeler_async_job(
    task: Any, *, kind: str, request_id: str, job_path: str = "/pdf-labeler/api/jobs"
):
    """Store async labeler metadata and return the queued-job response.

    Args:
        task: Celery task object returned by ``delay``.
        kind: Short label describing the queued job kind.
        request_id: Correlation ID for the current request.

    Returns:
        Response: A 202 JSON response describing the queued job.
    """
    job_id = str(uuid.uuid4())
    _store_labeler_job_mapping(job_id, task.id, extra={"kind": kind})
    return jsonify_with_status(
        {
            "success": True,
            "request_id": request_id,
            "status": "queued",
            "job_id": job_id,
            "job_url": f"{LABELER_BASE_PATH}{job_path}/{job_id}",
        },
        202,
    )


# =============================================================================
# Playground Template Management (shared by DOCX and PDF labelers)
# =============================================================================

_TEMPLATE_EXTENSIONS_DOCX: tuple[str, ...] = (".docx",)
_TEMPLATE_EXTENSIONS_PDF: tuple[str, ...] = (".pdf",)
_TEMPLATE_EXTENSIONS_ALL: tuple[str, ...] = (
    _TEMPLATE_EXTENSIONS_DOCX + _TEMPLATE_EXTENSIONS_PDF
)


def _normalize_template_filename(
    filename: Optional[str],
    *,
    allowed_extensions: tuple[str, ...] = _TEMPLATE_EXTENSIONS_ALL,
) -> str:
    """Validate a Playground template filename against allowed extensions.

    Args:
        filename: Raw filename from the request.
        allowed_extensions: Allowed file extensions for the current operation.

    Returns:
        str: A normalized template filename.
    """
    value = os.path.basename(str(filename or "").strip())
    if not value or value in {".", ".."}:
        raise DashboardAPIValidationError(
            "Template filename is required.", status_code=400
        )
    if not value.lower().endswith(allowed_extensions):
        raise DashboardAPIValidationError(
            f"Template file must be one of: {', '.join(allowed_extensions)}",
            status_code=400,
        )
    return value


def _list_playground_template_files(
    project: str, *, extensions: tuple[str, ...] = _TEMPLATE_EXTENSIONS_ALL
) -> List[Dict[str, str]]:
    """List uploaded DOCX or PDF templates in a Playground project.

    Args:
        project: Playground project name.
        extensions: File extensions to include.

    Returns:
        List[Dict[str, str]]: Template metadata for files in the template folder.
    """
    from docassemble.webapp.backend import directory_for
    from docassemble.webapp.files import SavedFile

    uid = getattr(current_user, "id", None)
    if uid is None:
        return []
    area = SavedFile(uid, fix=True, section="playgroundtemplate")
    the_directory = directory_for(area, project)
    if not os.path.isdir(the_directory):
        return []
    results = []
    for fname in sorted(os.listdir(the_directory)):
        if not fname.lower().endswith(extensions):
            continue
        fpath = os.path.join(the_directory, fname)
        if os.path.isfile(fpath):
            results.append(
                {
                    "filename": fname,
                    "label": fname,
                    "size": os.path.getsize(fpath),
                }
            )
    return results


def _load_playground_template_file(project: str, filename: str) -> bytes:
    """Load raw bytes for a Playground template file.

    Args:
        project: Playground project name.
        filename: Template filename to load.

    Returns:
        bytes: File contents for the requested template.
    """
    from docassemble.webapp.backend import directory_for
    from docassemble.webapp.files import SavedFile

    uid = getattr(current_user, "id", None)
    if uid is None:
        raise DashboardAPIValidationError("Login required.", status_code=401)
    area = SavedFile(uid, fix=True, section="playgroundtemplate")
    the_directory = directory_for(area, project)
    filepath = os.path.join(the_directory, filename)
    real_dir = os.path.realpath(the_directory)
    real_file = os.path.realpath(filepath)
    if not real_file.startswith(real_dir + os.sep):
        raise DashboardAPIValidationError("Invalid template path.", status_code=400)
    if not os.path.isfile(filepath):
        raise DashboardAPIValidationError(
            f"Template file '{filename}' not found in project '{project}'.",
            status_code=404,
        )
    with open(filepath, "rb") as fh:
        return fh.read()


def _save_playground_template_file(
    project: str, filename: str, content: bytes
) -> Dict[str, Any]:
    """Save a DOCX or PDF template into a Playground project.

    Args:
        project: Playground project name.
        filename: Template filename to write.
        content: Raw file bytes to persist.

    Returns:
        Dict[str, Any]: Metadata about the saved template file.
    """
    from docassemble.webapp.backend import directory_for
    from docassemble.webapp.files import SavedFile

    uid = getattr(current_user, "id", None)
    if uid is None:
        raise DashboardAPIValidationError("Login required.", status_code=401)
    area = SavedFile(uid, fix=True, section="playgroundtemplate")
    the_directory = directory_for(area, project)
    real_dir = os.path.realpath(the_directory)
    filepath = os.path.join(the_directory, filename)
    real_file = os.path.realpath(filepath)
    if not real_file.startswith(real_dir + os.sep):
        raise DashboardAPIValidationError("Invalid template path.", status_code=400)
    os.makedirs(the_directory, exist_ok=True)
    is_new = not os.path.isfile(filepath)
    with open(filepath, "wb") as fh:
        fh.write(content)
    area.finalize()
    return {
        "project": project,
        "filename": filename,
        "size": len(content),
        "created": is_new,
    }


@app.route(f"{LABELER_BASE_PATH}/labeler/api/playground-projects", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_playground_projects() -> Response:
    """Shared endpoint: list playground projects for both labelers.

    Returns:
        Response: A JSON response containing available Playground projects.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {"projects": _list_playground_projects()},
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: labeler playground-projects {request_id} error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/labeler/api/playground-files", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_playground_files() -> Response:
    """Shared endpoint: list YAML files in a playground project.

    Returns:
        Response: A JSON response containing YAML files for the selected project.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        project = _normalize_playground_project(request.args.get("project"))
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {
                    "project": project,
                    "files": _list_playground_yaml_files(project),
                },
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: labeler playground-files {request_id} error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/labeler/api/playground-variables", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_playground_variables() -> Response:
    """Shared endpoint: get variables from a playground YAML file.

    Returns:
        Response: A JSON response containing extracted variable metadata.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        project = _normalize_playground_project(request.args.get("project"))
        filename = _normalize_playground_filename(request.args.get("filename"))
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": _get_playground_variable_info(project, filename),
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: labeler playground-variables {request_id} error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/labeler/api/playground-templates", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_playground_templates() -> Response:
    """List template files (DOCX/PDF) in a playground project's template folder.

    Returns:
        Response: A JSON response containing template metadata for the project.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        project = _normalize_playground_project(request.args.get("project"))
        file_type = str(request.args.get("type") or "").strip().lower()
        extensions: tuple[str, ...]
        if file_type == "docx":
            extensions = _TEMPLATE_EXTENSIONS_DOCX
        elif file_type == "pdf":
            extensions = _TEMPLATE_EXTENSIONS_PDF
        else:
            extensions = _TEMPLATE_EXTENSIONS_ALL
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {
                    "project": project,
                    "templates": _list_playground_template_files(
                        project, extensions=extensions
                    ),
                },
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(f"ALDashboard: playground-templates {request_id} error: {exc!r}", "error")
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(
    f"{LABELER_BASE_PATH}/labeler/api/playground-templates/load", methods=["GET"]
)
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_playground_template_load() -> Response:
    """Load a template file from a playground project's template folder.

    Returns:
        Response: A JSON response containing the requested template as base64.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        project = _normalize_playground_project(request.args.get("project"))
        filename = _normalize_template_filename(request.args.get("filename"))
        content = _load_playground_template_file(project, filename)
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {
                    "project": project,
                    "filename": filename,
                    "size": len(content),
                    "file_content_base64": base64.b64encode(content).decode("ascii"),
                },
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: playground-template-load {request_id} error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(
    f"{LABELER_BASE_PATH}/labeler/api/playground-templates/save", methods=["POST"]
)
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def labeler_playground_template_save() -> Response:
    """Save a template file back to a playground project's template folder.

    Returns:
        Response: A JSON response describing the saved template file.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        post_data = request.get_json(silent=True) or {}
        if "file" in request.files:
            upload = request.files["file"]
            filename = upload.filename or ""
            content = upload.read()
            project = _normalize_playground_project(
                request.form.get("project") or post_data.get("project")
            )
        else:
            filename = str(post_data.get("filename") or "")
            project = _normalize_playground_project(post_data.get("project"))
            content = decode_base64_content(post_data.get("file_content_base64"))
        filename = _normalize_template_filename(filename)
        _validate_upload_size(content)
        result = _save_playground_template_file(project, filename, content)
        action = "created" if result["created"] else "updated"
        log(
            f"ALDashboard: playground-template-save {request_id} {action} '{filename}' in project '{project}'",
            "info",
        )
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": result,
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: playground-template-save {request_id} error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


# =============================================================================
# DOCX Labeler Routes
# =============================================================================


@app.route(f"{LABELER_BASE_PATH}/docx-labeler", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def docx_labeler_page() -> Response:
    """Serve the DOCX labeler interactive UI.

    Returns:
        Response: The rendered DOCX labeler HTML page.
    """
    log("ALDashboard: Serving DOCX labeler page", "info")
    html_content = _render_template_content(
        "docx_labeler.html", bootstrap_data=_build_docx_labeler_bootstrap()
    )
    if not html_content:
        log("ALDashboard: DOCX labeler template not found", "error")
        return Response(
            "DOCX labeler template not found.", status=500, mimetype="text/plain"
        )
    return Response(html_content, mimetype="text/html")


def _read_file_upload(
    default_filename: str, required_ext: str
) -> tuple[str, bytes, Dict[str, Any]]:
    """Read a file from a multipart upload or JSON body with base64 content.

    Args:
        default_filename: Fallback filename when none is provided.
        required_ext: Required file extension (e.g. ".docx", ".pdf").

    Returns:
        A (filename, content_bytes, post_data) tuple.
    """
    if "file" in request.files:
        upload = request.files["file"]
        filename = upload.filename or default_filename
        content = upload.read()
        post_data = dict(request.form)
    else:
        post_data = request.get_json(silent=True) or {}
        filename = str(post_data.get("filename") or default_filename)
        content = decode_base64_content(post_data.get("file_content_base64"))

    _validate_upload_size(content)
    if not filename.lower().endswith(required_ext):
        ext_label = required_ext.lstrip(".").upper()
        raise DashboardAPIValidationError(
            f"Only {ext_label} files are supported.", status_code=415
        )
    return filename, content, post_data


def _read_docx_labeler_file_request() -> tuple[str, bytes, Dict[str, Any]]:
    return _read_file_upload("upload.docx", ".docx")


def _read_pdf_labeler_file_request() -> tuple[str, bytes, Dict[str, Any]]:
    return _read_file_upload("upload.pdf", ".pdf")


def _resolve_interview_variables(
    post_data: Dict[str, Any],
) -> tuple[Optional[List[str]], str, Optional[str], Optional[str], Optional[str]]:
    """Parse preferred variable names and resolve from playground/installed interview.

    Returns:
        A tuple of (preferred_variable_names, interview_source_mode,
        selected_playground_project, selected_playground_filename,
        selected_installed_interview_path).
    """
    preferred_variable_names: Optional[List[str]] = None
    interview_source_mode = (
        str(post_data.get("interview_source_mode") or "playground").strip().lower()
        or "playground"
    )
    selected_playground_project: Optional[str] = None
    selected_playground_filename: Optional[str] = None
    selected_installed_interview_path: Optional[str] = None

    preferred_variable_names_raw = post_data.get("preferred_variable_names")
    if preferred_variable_names_raw:
        if isinstance(preferred_variable_names_raw, str):
            try:
                parsed_preferred_names = json.loads(preferred_variable_names_raw)
            except json.JSONDecodeError:
                parsed_preferred_names = [
                    item.strip()
                    for item in preferred_variable_names_raw.split(",")
                    if item.strip()
                ]
        elif isinstance(preferred_variable_names_raw, list):
            parsed_preferred_names = preferred_variable_names_raw
        else:
            parsed_preferred_names = []
        preferred_variable_names = [
            str(item).strip() for item in parsed_preferred_names if str(item).strip()
        ] or None

    use_playground_variables = parse_bool(
        post_data.get("use_playground_variables"), default=False
    )
    if use_playground_variables:
        if interview_source_mode == "installed":
            selected_installed_interview_path = _normalize_installed_interview_path(
                post_data.get("installed_interview_path")
                or (
                    (
                        str(post_data.get("installed_package") or "").strip()
                        + ":"
                        + str(post_data.get("installed_yaml_file") or "").strip()
                    )
                    if (
                        str(post_data.get("installed_package") or "").strip()
                        and str(post_data.get("installed_yaml_file") or "").strip()
                    )
                    else None
                )
            )
            if preferred_variable_names is None:
                preferred_variable_names = _get_installed_interview_variable_info(
                    selected_installed_interview_path
                )["all_names"]
        else:
            selected_playground_project = _normalize_playground_project(
                post_data.get("playground_project")
            )
            selected_playground_filename = _normalize_playground_filename(
                post_data.get("playground_yaml_file")
            )
            if preferred_variable_names is None:
                preferred_variable_names = _get_playground_variable_info(
                    selected_playground_project, selected_playground_filename
                )["all_names"]

    return (
        preferred_variable_names,
        interview_source_mode,
        selected_playground_project,
        selected_playground_filename,
        selected_installed_interview_path,
    )


def _docx_output_payload(output_path: str, output_filename: str) -> Dict[str, Any]:
    with open(output_path, "rb") as handle:
        output_bytes = handle.read()
    return {
        "filename": output_filename,
        "docx_base64": base64.b64encode(output_bytes).decode("ascii"),
    }


def _write_temp_docx(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(content)
        return tmp.name


@app.route(f"{LABELER_BASE_PATH}/labeler/api/models", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_models() -> Response:
    """Return AI model metadata for DOCX/PDF labeler settings UIs.

    Returns:
        Response: A JSON response containing labeler model metadata.
    """
    request_id = str(uuid.uuid4())
    return jsonify(
        {
            "success": True,
            "request_id": request_id,
            "data": _build_labeler_model_catalog(),
        }
    )


@app.route(f"{LABELER_BASE_PATH}/labeler/api/auth-status", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_auth_status() -> Response:
    """Return browser-session auth status for labeler UI controls.

    Returns:
        Response: A JSON response describing auth and AI availability.
    """
    request_id = str(uuid.uuid4())
    identity = _labeler_session_identity()
    next_target = _labeler_auth_return_target()
    login_url = f"/user/sign-in?next={quote(next_target, safe='')}"
    logout_url = f"/user/sign-out?next={quote(next_target, safe='')}"
    return jsonify(
        {
            "success": True,
            "request_id": request_id,
            "data": {
                "is_authenticated": bool(identity.get("is_authenticated")),
                "email": identity.get("email"),
                "user_id": identity.get("user_id"),
                "login_url": login_url,
                "logout_url": logout_url,
                "ai_enabled": _labeler_ai_auth_check(),
            },
        }
    )


@app.route(f"{LABELER_BASE_PATH}/labeler/api/installed-packages", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_installed_packages() -> Response:
    """Shared endpoint: list installed interview packages for both labelers.

    Returns:
        Response: A JSON response containing installed interview packages.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {"packages": _list_installed_interview_packages()},
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: installed-packages {request_id} server error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/labeler/api/installed-files", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_installed_files() -> Response:
    """Shared endpoint: list YAML files in an installed interview package.

    Returns:
        Response: A JSON response containing YAML files for the package.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        package_name = _normalize_installed_package_name(request.args.get("package"))
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {
                    "package": package_name,
                    "files": _list_installed_interview_files(package_name),
                },
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(f"ALDashboard: installed-files {request_id} server error: {exc!r}", "error")
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/labeler/api/installed-variables", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def labeler_installed_variables() -> Response:
    """Shared endpoint: get variables from an installed interview.

    Returns:
        Response: A JSON response containing extracted variable metadata.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_playground_auth_check():
        return _playground_auth_fail(request_id)
    try:
        interview_path = _normalize_installed_interview_path(
            request.args.get("interview_path")
        )
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": _get_installed_interview_variable_info(interview_path),
            }
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: installed-variables {request_id} server error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/docx-labeler/api/extract-runs", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def docx_labeler_extract_runs() -> Response:
    """Extract paragraph runs from a DOCX file for labeling.

    Returns:
        Response: A JSON response containing extracted DOCX run coordinates.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: extract-runs request {request_id}", "info")
    try:
        import docx

        from .docx_wrangling import defragment_docx_runs, get_docx_run_items

        filename, content, post_data = _read_docx_labeler_file_request()
        defragment_runs = parse_bool(post_data.get("defragment_runs"), default=True)

        # Write to temp file for processing
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            temp_path = tmp.name

        try:
            doc = docx.Document(temp_path)
            defragmentation = {"paragraphs_defragmented": 0, "runs_removed": 0}
            if defragment_runs:
                doc, defragmentation = defragment_docx_runs(doc)
            runs = get_docx_run_items(doc)
            paragraph_count = 0
            if runs:
                paragraph_count = max(int(item[0]) for item in runs) + 1

            log(
                f"ALDashboard: extract-runs {request_id} extracted {len(runs)} runs from {paragraph_count} paragraphs in '{filename}'",
                "info",
            )
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": filename,
                        "paragraph_count": paragraph_count,
                        "run_count": len(runs),
                        "runs": runs,
                        "defragment_runs": defragment_runs,
                        "defragmentation": defragmentation,
                    },
                }
            )
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except DashboardAPIValidationError as exc:
        log(
            f"ALDashboard: extract-runs {request_id} validation error: {exc.message}",
            "warning",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(f"ALDashboard: extract-runs {request_id} server error: {exc!r}", "error")
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/docx-labeler/api/suggest-labels", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def docx_labeler_suggest_labels() -> Response:
    """Use AI to suggest Jinja2 labels for a DOCX file.

    Returns:
        Response: A JSON response containing generated DOCX label suggestions.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: suggest-labels request {request_id}", "info")
    if not _labeler_ai_auth_check():
        log(
            f"ALDashboard: suggest-labels auth failed for request {request_id}",
            "warning",
        )
        return _ai_auth_fail(request_id)

    try:
        filename, content, post_data = _read_docx_labeler_file_request()
        post_data = merge_raw_options(post_data)

        # Extract options
        prompt_profile = (
            str(post_data.get("prompt_profile") or "standard").strip() or "standard"
        )
        optional_context = post_data.get("context_text")
        custom_prompt = post_data.get("custom_prompt")
        additional_instructions = post_data.get("additional_instructions")
        defragment_runs = parse_bool(post_data.get("defragment_runs"), default=True)
        model = post_data.get("model")
        if model is None or str(model).strip() == "":
            model = _build_labeler_model_catalog()["default_model"]
        model = str(model)
        judge_model = post_data.get("judge_model")
        if judge_model is not None:
            judge_model = str(judge_model).strip() or None
        openai_api = post_data.get("openai_api")
        openai_base_url = post_data.get("openai_base_url")
        generator_models = None
        generator_models_raw = post_data.get("generator_models")
        if generator_models_raw:
            if isinstance(generator_models_raw, str):
                try:
                    parsed_models = json.loads(generator_models_raw)
                except json.JSONDecodeError:
                    parsed_models = [
                        item.strip()
                        for item in generator_models_raw.split(",")
                        if item.strip()
                    ]
            elif isinstance(generator_models_raw, list):
                parsed_models = generator_models_raw
            else:
                parsed_models = []
            generator_models = [
                str(item).strip() for item in parsed_models if str(item).strip()
            ] or None

        # Parse custom people names if provided
        custom_people_names = None
        custom_people_raw = post_data.get("custom_people_names")
        if custom_people_raw:
            if isinstance(custom_people_raw, str):
                try:
                    custom_people_names = json.loads(custom_people_raw)
                except json.JSONDecodeError:
                    pass
            elif isinstance(custom_people_raw, list):
                custom_people_names = custom_people_raw

        (
            preferred_variable_names,
            interview_source_mode,
            selected_playground_project,
            selected_playground_filename,
            selected_installed_interview_path,
        ) = _resolve_interview_variables(post_data)

        task_payload = {
            "request_id": request_id,
            "filename": filename,
            "file_content_base64": base64.b64encode(content).decode("ascii"),
            "prompt_profile": prompt_profile,
            "context_text": optional_context,
            "custom_prompt": custom_prompt,
            "additional_instructions": additional_instructions,
            "defragment_runs": defragment_runs,
            "model": model,
            "judge_model": judge_model,
            "openai_api": openai_api,
            "openai_base_url": openai_base_url,
            "generator_models": generator_models,
            "custom_people_names": custom_people_names,
            "preferred_variable_names": preferred_variable_names,
            "interview_source_mode": interview_source_mode,
            "playground_project": selected_playground_project,
            "playground_yaml_file": selected_playground_filename,
            "installed_interview_path": selected_installed_interview_path,
        }

        if coerce_async_flag(post_data):
            if not _labeler_async_is_configured():
                return jsonify_with_status(
                    {
                        "success": False,
                        "request_id": request_id,
                        "error": {
                            "type": "async_not_configured",
                            "message": (
                                "Async mode is not configured. Add "
                                f"{ASYNC_CELERY_MODULE!r} to the docassemble "
                                "'celery modules' configuration list."
                            ),
                        },
                    },
                    503,
                )

            from .api_dashboard_worker import dashboard_docx_labeler_suggest_task

            task = dashboard_docx_labeler_suggest_task.delay(payload=task_payload)
            return _queue_labeler_async_job(
                task,
                kind="docx_suggest",
                request_id=request_id,
                job_path="/docx-labeler/api/jobs",
            )

        data = docx_labeler_suggest_payload_from_options(task_payload)
        log(
            f"ALDashboard: suggest-labels {request_id} generated {len(data.get('suggestions', []))} suggestions for '{filename}'",
            "info",
        )
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": data,
            }
        )

    except DashboardAPIValidationError as exc:
        log(
            f"ALDashboard: suggest-labels {request_id} validation error: {exc.message}",
            "warning",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(f"ALDashboard: suggest-labels {request_id} server error: {exc!r}", "error")
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/docx-labeler/api/validate-syntax", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def docx_labeler_validate_syntax() -> Response:
    """Validate DOCX Jinja syntax after simulated label and rename edits."""
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: validate-docx-syntax request {request_id}", "info")
    try:
        from .docx_wrangling import validate_docx_template_syntax

        filename, content, post_data = _read_docx_labeler_file_request()
        labels_raw = post_data.get("labels")
        renames_raw = post_data.get("renames")

        defragment_runs = parse_bool(post_data.get("defragment_runs"), default=True)

        labels = []
        if labels_raw:
            if isinstance(labels_raw, str):
                labels = json.loads(labels_raw)
            elif isinstance(labels_raw, list):
                labels = labels_raw

        renames = []
        if renames_raw:
            if isinstance(renames_raw, str):
                renames = json.loads(renames_raw)
            elif isinstance(renames_raw, list):
                renames = renames_raw

        modified_runs = []
        for label in labels:
            para = int(label.get("paragraph", 0))
            run = int(label.get("run", 0))
            text = str(label.get("text", ""))
            new_para = int(label.get("new_paragraph", 0))
            modified_runs.append((para, run, text, new_para))

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            temp_path = tmp.name

        try:
            validation = validate_docx_template_syntax(
                temp_path,
                suggestions=modified_runs,
                renames=renames,
                defragment_runs=defragment_runs,
            )
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        log(
            f"ALDashboard: validate-docx-syntax {request_id} completed with {validation.get('error_count', 0)} errors and {validation.get('warning_count', 0)} warnings for '{filename}'",
            "info",
        )
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {
                    "filename": filename,
                    "validation": validation,
                },
            }
        )

    except DashboardAPIValidationError as exc:
        log(
            f"ALDashboard: validate-docx-syntax {request_id} validation error: {exc.message}",
            "warning",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(
            f"ALDashboard: validate-docx-syntax {request_id} server error: {exc!r}",
            "error",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/docx-labeler/api/utilities", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def docx_labeler_run_utility() -> Response:
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: docx utility request {request_id}", "info")
    try:
        from .docx_wrangling import apply_jinja2_highlights, defragment_docx_runs
        from .validate_docx import (
            strip_docx_problem_controls,
            validate_docx_ooxml_schema,
        )
        import docx

        filename, content, post_data = _read_docx_labeler_file_request()
        action = str(post_data.get("action") or "").strip().lower()
        if action not in {
            "validate",
            "schema-validate",
            "defragment-runs",
            "cleanup-docx",
            "highlight-jinja2",
        }:
            raise DashboardAPIValidationError("Unknown DOCX utility action.")

        temp_path = _write_temp_docx(content)
        try:
            if action == "validate":
                data = validate_docx_payload_from_options(
                    {
                        "files": [
                            {
                                "filename": filename,
                                "file_content_base64": base64.b64encode(content).decode(
                                    "ascii"
                                ),
                            }
                        ]
                    }
                )
                return jsonify(
                    {
                        "success": True,
                        "request_id": request_id,
                        "data": {"action": action, "report": data["files"][0]},
                    }
                )

            if action == "schema-validate":
                report = validate_docx_ooxml_schema(temp_path)
                return jsonify(
                    {
                        "success": True,
                        "request_id": request_id,
                        "data": {"action": action, "report": report},
                    }
                )

            with tempfile.NamedTemporaryFile(
                suffix=".docx", delete=False
            ) as output_file:
                output_path = output_file.name

            try:
                if action == "cleanup-docx":
                    report = strip_docx_problem_controls(temp_path, output_path)
                    payload = _docx_output_payload(
                        output_path,
                        filename.replace(".docx", "-cleaned.docx"),
                    )
                else:
                    utility_doc = docx.Document(temp_path)
                    report = {}
                    if action == "defragment-runs":
                        utility_doc, report = defragment_docx_runs(utility_doc)
                    elif action == "highlight-jinja2":
                        utility_doc = apply_jinja2_highlights(utility_doc)
                        report = {"highlighted": True}
                    utility_doc.save(output_path)
                    suffix = (
                        "-defragmented.docx"
                        if action == "defragment-runs"
                        else "-highlighted.docx"
                    )
                    payload = _docx_output_payload(
                        output_path, filename.replace(".docx", suffix)
                    )

                return jsonify(
                    {
                        "success": True,
                        "request_id": request_id,
                        "data": {
                            "action": action,
                            "report": report,
                            **payload,
                        },
                    }
                )
            finally:
                if os.path.exists(output_path):
                    os.remove(output_path)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(f"ALDashboard: docx utility {request_id} server error: {exc!r}", "error")
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/docx-labeler/api/repair", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def docx_labeler_repair_docx() -> Response:
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: docx repair request {request_id}", "info")
    try:
        from .docx_repair import (
            repair_docx_xml_conservatively,
            rescue_docx_to_shell,
            roundtrip_docx_via_soffice,
        )

        filename, content, post_data = _read_docx_labeler_file_request()
        action = str(post_data.get("action") or "").strip().lower()
        if action not in {"soffice-roundtrip", "repair-xml", "rescue-shell"}:
            raise DashboardAPIValidationError("Unknown DOCX repair action.")

        temp_path = _write_temp_docx(content)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as output_file:
            output_path = output_file.name
        try:
            if action == "soffice-roundtrip":
                report = roundtrip_docx_via_soffice(temp_path, output_path)
                output_name = filename.replace(".docx", "-roundtripped.docx")
            elif action == "repair-xml":
                report = repair_docx_xml_conservatively(temp_path, output_path)
                output_name = filename.replace(".docx", "-repaired.docx")
            else:
                report = rescue_docx_to_shell(temp_path, output_path)
                output_name = filename.replace(".docx", "-rescued.docx")

            data: Dict[str, Any] = {"action": action, "report": report}
            if (
                os.path.exists(output_path)
                and os.path.getsize(output_path) > 0
                and report.get("available", True)
            ):
                data.update(_docx_output_payload(output_path, output_name))
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": data,
                }
            )
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            if os.path.exists(output_path):
                os.remove(output_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(f"ALDashboard: docx repair {request_id} server error: {exc!r}", "error")
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route(f"{LABELER_BASE_PATH}/docx-labeler/api/apply-labels", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def docx_labeler_apply_labels() -> Response:
    """Apply accepted labels and/or renames to a DOCX file and return the modified file.

    Supports two types of modifications:
    - labels: New AI-suggested labels to insert (paragraph, run, text, new_paragraph)
    - renames: Find/replace operations on existing labels (original, replacement)

    Returns:
        Response: A JSON response containing the modified DOCX as base64.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: apply-labels request {request_id}", "info")
    try:
        from .docx_wrangling import (
            apply_docx_label_renames,
            apply_jinja2_highlights,
            defragment_docx_runs,
            update_docx,
            validate_docx_template_syntax,
        )
        import docx

        filename, content, post_data = _read_docx_labeler_file_request()
        labels_raw = post_data.get("labels")
        renames_raw = post_data.get("renames")
        defragment_runs = parse_bool(post_data.get("defragment_runs"), default=True)
        apply_highlights = parse_bool(post_data.get("apply_highlights"), default=False)
        allow_invalid_syntax = parse_bool(
            post_data.get("allow_invalid_syntax"), default=False
        )

        # Parse labels (new insertions)
        labels = []
        if labels_raw:
            if isinstance(labels_raw, str):
                labels = json.loads(labels_raw)
            elif isinstance(labels_raw, list):
                labels = labels_raw

        # Parse renames (find/replace on existing)
        renames = []
        if renames_raw:
            if isinstance(renames_raw, str):
                renames = json.loads(renames_raw)
            elif isinstance(renames_raw, list):
                renames = renames_raw

        if not labels and not renames and not apply_highlights:
            raise DashboardAPIValidationError(
                "Labels, renames, or apply_highlights must be provided."
            )

        log(
            f"ALDashboard: apply-labels {request_id} processing {len(labels)} label insertions, {len(renames)} renames, apply_highlights={apply_highlights} for '{filename}'",
            "info",
        )

        # Convert labels to the format expected by update_docx
        modified_runs = []
        for label in labels:
            para = int(label.get("paragraph", 0))
            run = int(label.get("run", 0))
            text = str(label.get("text", ""))
            new_para = int(label.get("new_paragraph", 0))
            modified_runs.append((para, run, text, new_para))

        # Write to temp file for processing
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            temp_path = tmp.name

        try:
            doc = docx.Document(temp_path)
            syntax_validation = validate_docx_template_syntax(
                doc,
                suggestions=modified_runs,
                renames=renames,
                defragment_runs=defragment_runs,
            )
            if not syntax_validation.get("valid") and not allow_invalid_syntax:
                message = "DOCX contains invalid Jinja syntax."
                if syntax_validation.get("errors"):
                    message = str(
                        syntax_validation["errors"][0].get("message") or message
                    )
                return jsonify_with_status(
                    {
                        "success": False,
                        "request_id": request_id,
                        "error": {
                            "type": "validation_error",
                            "message": message,
                        },
                        "data": {
                            "validation": syntax_validation,
                        },
                    },
                    409,
                )

            defragmentation = {"paragraphs_defragmented": 0, "runs_removed": 0}

            if defragment_runs and modified_runs:
                target_paragraph_numbers = sorted(
                    {
                        para
                        for para, _run, _text, new_para in modified_runs
                        if new_para == 0 and para >= 0
                    }
                )
                if target_paragraph_numbers:
                    doc, defragmentation = defragment_docx_runs(
                        doc, paragraph_numbers=target_paragraph_numbers
                    )

            # Apply renames first (find/replace existing Jinja2 labels)
            if renames:
                apply_docx_label_renames(doc, renames)

            # Apply new label insertions
            if modified_runs:
                doc = update_docx(
                    doc,
                    modified_runs,
                    apply_jinja_highlights=apply_highlights,
                )
            elif apply_highlights:
                doc = apply_jinja2_highlights(doc)

            # Save to bytes
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            output_bytes = output_buffer.read()

            if labels or renames:
                output_suffix = "-labeled"
                if apply_highlights:
                    output_suffix += "-highlighted"
            elif apply_highlights:
                output_suffix = "-highlighted"
            else:
                output_suffix = "-updated"
            output_filename = filename.replace(".docx", f"{output_suffix}.docx")

            log(
                f"ALDashboard: apply-labels {request_id} successfully produced '{output_filename}'",
                "info",
            )
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": output_filename,
                        "docx_base64": base64.b64encode(output_bytes).decode("ascii"),
                        "defragment_runs": defragment_runs,
                        "defragmented_before_apply": bool(
                            defragmentation["paragraphs_defragmented"]
                            or defragmentation["runs_removed"]
                        ),
                        "defragmentation": defragmentation,
                        "apply_highlights": apply_highlights,
                        "validation": syntax_validation,
                    },
                }
            )
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except DashboardAPIValidationError as exc:
        log(
            f"ALDashboard: apply-labels {request_id} validation error: {exc.message}",
            "warning",
        )
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        log(f"ALDashboard: apply-labels {request_id} server error: {exc!r}", "error")
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


# =============================================================================
# PDF Labeler Routes
# =============================================================================


@app.route("/pdf-labeler", methods=["GET"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def pdf_labeler_page() -> Response:
    """Serve the PDF labeler interactive UI.

    Returns:
        Response: The rendered PDF labeler HTML page.
    """
    log("ALDashboard: Serving PDF labeler page", "info")
    html_content = _render_template_content(
        "pdf_labeler.html", bootstrap_data=_build_pdf_labeler_bootstrap()
    )
    if not html_content:
        log("ALDashboard: PDF labeler template not found", "error")
        return Response(
            "PDF labeler template not found.", status=500, mimetype="text/plain"
        )
    return Response(html_content, mimetype="text/html")


@app.route("/pdf-labeler/api/jobs/<job_id>", methods=["GET"])
@app.route("/docx-labeler/api/jobs/<job_id>", methods=["GET"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/jobs/<job_id>", methods=["GET"])
@app.route(f"{LABELER_BASE_PATH}/docx-labeler/api/jobs/<job_id>", methods=["GET"])
@csrf.exempt
@cross_origin(origins="*", methods=["GET", "HEAD"], automatic_options=True)
def pdf_labeler_job(job_id: str) -> Response:
    """Return status or result data for an async PDF labeler job.

    Args:
        job_id: Public labeler job identifier.

    Returns:
        Response: A JSON response describing job state and any available result.
    """
    request_id = str(uuid.uuid4())
    if not _labeler_ai_auth_check():
        return _ai_auth_fail(request_id)

    task_info = _fetch_labeler_job_mapping(job_id)
    if not task_info:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "not_found", "message": "Job not found."},
            },
            404,
        )

    result = workerapp.AsyncResult(id=task_info["id"])
    state = (result.state or "").upper()
    if state == "SUCCESS":
        status = "succeeded"
    elif state in {"RECEIVED", "STARTED", "RETRY"}:
        status = "running"
    elif state == "FAILURE":
        status = "failed"
    else:
        status = "queued"

    body: Dict[str, Any] = {
        "success": True,
        "request_id": request_id,
        "job_id": job_id,
        "task_id": task_info.get("id"),
        "status": status,
        "celery_state": state,
        "created_at": task_info.get("created_at"),
    }
    if state == "SUCCESS":
        data = _normalize_labeler_result_for_json(result.get())
        if task_info.get("kind") == "pdf_auto_detect" and isinstance(data, dict):
            data = _format_pdf_auto_detect_labeler_data(data)
        elif task_info.get("kind") == "pdf_relabel" and isinstance(data, dict):
            data = _format_pdf_relabel_labeler_data(data)
        body["data"] = data
    elif state == "FAILURE":
        error_obj = result.result
        body["error"] = {
            "type": getattr(error_obj, "__class__", type(error_obj)).__name__,
            "message": str(error_obj),
        }
    return jsonify(body)


@app.route("/pdf-labeler/api/detect-fields", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/detect-fields", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_detect_fields() -> Response:
    """Detect existing form fields in a PDF.

    Returns:
        Response: A JSON response containing detected PDF field metadata.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: detect-fields request {request_id}", "info")

    try:
        import formfyxer  # type: ignore[import-not-found]

        filename, content, _post_data = _read_pdf_labeler_file_request()

        # Write to temp file for processing
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            temp_path = tmp.name

        try:
            # Get existing fields with positions
            fields_per_page = formfyxer.get_existing_pdf_fields(temp_path)

            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": filename,
                        "page_count": len(fields_per_page),
                        "fields": _format_pdf_fields_for_ui_payload(fields_per_page),
                    },
                }
            )
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route("/pdf-labeler/api/accessibility-inspect", methods=["POST"])
@app.route(
    f"{LABELER_BASE_PATH}/pdf-labeler/api/accessibility-inspect", methods=["POST"]
)
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_accessibility_inspect() -> Response:
    """Inspect PDF accessibility metadata for the browser editor mode.

    Returns:
        Response: A JSON response containing field tooltip metadata, field order,
            image assets, document metadata, and tag structure summary.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: accessibility-inspect request {request_id}", "info")

    try:
        from .pdf_accessibility import PDFAccessibilityError, inspect_pdf_accessibility

        _filename, content, _post_data = _read_pdf_labeler_file_request()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
            tmp_in.write(content)
            input_path = tmp_in.name

        try:
            payload = inspect_pdf_accessibility(input_path)
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": payload,
                }
            )
        finally:
            if os.path.exists(input_path):
                os.remove(input_path)
    except PDFAccessibilityError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {
                    "type": "validation_error",
                    "message": str(exc),
                },
            },
            400,
        )
    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {
                    "type": "validation_error",
                    "message": exc.message,
                },
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {
                    "type": "server_error",
                    "message": str(exc),
                },
            },
            500,
        )


@app.route("/pdf-labeler/api/auto-detect", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/auto-detect", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_auto_detect() -> Response:
    """Use AI to automatically detect and add fields to a PDF.

    Returns:
        Response: A JSON response containing detected field positions and the
        updated PDF.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: auto-detect request {request_id}", "info")
    if not _labeler_ai_auth_check():
        log(f"ALDashboard: auto-detect auth failed for request {request_id}", "warning")
        return _ai_auth_fail(request_id)

    try:
        filename, content, post_data = _read_pdf_labeler_file_request()
        post_data = merge_raw_options(post_data)
        normalize_fields = parse_bool(post_data.get("normalize_fields"), default=True)

        # Resolve preferred variable names from Playground/installed interview.
        (
            preferred_variable_names,
            _interview_source_mode,
            _selected_playground_project,
            _selected_playground_filename,
            _selected_installed_interview_path,
        ) = _resolve_interview_variables(post_data)

        detect_payload: Dict[str, Any] = {
            "filename": filename,
            "file_content_base64": base64.b64encode(content).decode("ascii"),
            "include_pdf_base64": True,
            "include_field_positions": True,
            **post_data,
        }
        detect_payload.setdefault("relabel_with_ai", normalize_fields)
        if "preferred_variable_names" in post_data or preferred_variable_names:
            detect_payload["preferred_variable_names"] = preferred_variable_names or []

        if coerce_async_flag(post_data):
            if not _labeler_async_is_configured():
                return jsonify_with_status(
                    {
                        "success": False,
                        "request_id": request_id,
                        "error": {
                            "type": "async_not_configured",
                            "message": (
                                "Async mode is not configured. Add "
                                f"{ASYNC_CELERY_MODULE!r} to the docassemble "
                                "'celery modules' configuration list."
                            ),
                        },
                    },
                    503,
                )
            from .api_dashboard_worker import dashboard_pdf_fields_detect_task

            task = dashboard_pdf_fields_detect_task.delay(payload=detect_payload)
            return _queue_labeler_async_job(
                task, kind="pdf_auto_detect", request_id=request_id
            )

        payload = pdf_fields_detect_payload_from_options(detect_payload)
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": _format_pdf_auto_detect_labeler_data(payload),
            }
        )

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route("/pdf-labeler/api/relabel", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/relabel", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_relabel() -> Response:
    """Relabel PDF fields using AI suggestions.

    Returns:
        Response: A JSON response containing renamed field names and the updated PDF.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: pdf-relabel request {request_id}", "info")
    if not _labeler_ai_auth_check():
        log(f"ALDashboard: pdf-relabel auth failed for request {request_id}", "warning")
        return _ai_auth_fail(request_id)

    try:
        filename, content, post_data = _read_pdf_labeler_file_request()
        post_data = merge_raw_options(post_data)

        relabel_payload: Dict[str, Any] = {
            "filename": filename,
            "file_content_base64": base64.b64encode(content).decode("ascii"),
            "relabel_with_ai": True,
            "include_pdf_base64": True,
            **post_data,
        }

        if coerce_async_flag(post_data):
            if not _labeler_async_is_configured():
                return jsonify_with_status(
                    {
                        "success": False,
                        "request_id": request_id,
                        "error": {
                            "type": "async_not_configured",
                            "message": (
                                "Async mode is not configured. Add "
                                f"{ASYNC_CELERY_MODULE!r} to the docassemble "
                                "'celery modules' configuration list."
                            ),
                        },
                    },
                    503,
                )
            from .api_dashboard_worker import dashboard_pdf_fields_relabel_task

            task = dashboard_pdf_fields_relabel_task.delay(payload=relabel_payload)
            return _queue_labeler_async_job(
                task, kind="pdf_relabel", request_id=request_id
            )

        payload = pdf_fields_relabel_payload_from_options(relabel_payload)
        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": _format_pdf_relabel_labeler_data(payload),
            }
        )

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route("/pdf-labeler/api/apply-fields", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/apply-fields", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_apply_fields() -> Response:
    """Apply field definitions to a PDF and return the modified file.

    Returns:
        Response: A JSON response containing the updated PDF as base64.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: apply-fields request {request_id}", "info")

    try:
        import formfyxer  # type: ignore[import-not-found]
        from formfyxer.pdf_wrangling import FormField, FieldType, set_fields
        from reportlab.lib.colors import HexColor

        filename, content, post_data = _read_pdf_labeler_file_request()
        fields_raw = post_data.get("fields")
        accessibility_raw = post_data.get("accessibility")

        # Parse fields
        if isinstance(fields_raw, str):
            fields_data = json.loads(fields_raw)
        elif isinstance(fields_raw, list):
            fields_data = fields_raw
        else:
            raise DashboardAPIValidationError("fields is required and must be a list.")

        # Get page count from original PDF
        import pikepdf

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
            tmp_in.write(content)
            input_path = tmp_in.name

        output_path = None
        try:
            with pikepdf.open(input_path) as pdf:
                page_count = len(pdf.pages)

            checkbox_export_values = {
                str(field.get("name", "")): str(
                    field.get("checkboxExportValue", "")
                ).strip()
                for field in fields_data
                if str(field.get("type", "")).lower() == "checkbox"
                and str(field.get("checkboxExportValue", "")).strip()
            }
            explicit_background_fields = _collect_fields_with_explicit_background(
                fields_data
            )

            fields_per_page = build_pdf_export_fields_per_page(
                fields_data,
                page_count=page_count,
                form_field_cls=FormField,
                field_type_enum=FieldType,
                color_parser=HexColor,
            )

            # Apply fields using FormFyxer
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_out:
                output_path = tmp_out.name

            set_fields(input_path, output_path, fields_per_page, overwrite=True)
            _apply_checkbox_export_values(output_path, checkbox_export_values)
            _apply_pdf_field_visual_defaults(
                output_path,
                explicit_background_fields=explicit_background_fields,
            )

            accessibility_payload = _parse_optional_json_field(
                accessibility_raw, field_name="accessibility"
            )
            if accessibility_payload is None:
                accessibility_payload = {}
            if not isinstance(accessibility_payload, dict):
                raise DashboardAPIValidationError(
                    "accessibility must be a JSON object."
                )

            accessibility_enabled = parse_bool(
                accessibility_payload.get("enabled"), default=True
            )
            if accessibility_enabled:
                field_tooltips_from_fields: Dict[str, str] = {}
                field_order_from_fields: List[str] = []
                for field in fields_data:
                    field_name = str(field.get("name") or "").strip()
                    if not field_name:
                        continue
                    field_order_from_fields.append(field_name)
                    tooltip = str(field.get("tooltip") or "").strip()
                    if tooltip:
                        field_tooltips_from_fields[field_name] = tooltip

                field_tooltips_override = accessibility_payload.get("field_tooltips")
                if not isinstance(field_tooltips_override, dict):
                    field_tooltips_override = {}
                merged_field_tooltips = {
                    **field_tooltips_from_fields,
                    **{
                        str(key): str(value)
                        for key, value in field_tooltips_override.items()
                        if str(key).strip()
                    },
                }

                field_order = accessibility_payload.get("field_order")
                if isinstance(field_order, list):
                    ordered_names = [
                        str(name) for name in field_order if str(name).strip()
                    ]
                else:
                    ordered_names = field_order_from_fields

                metadata_payload = accessibility_payload.get("metadata")
                if not isinstance(metadata_payload, dict):
                    metadata_payload = {}
                image_alt_text_payload = accessibility_payload.get("image_alt_text")
                if not isinstance(image_alt_text_payload, dict):
                    image_alt_text_payload = {}

                from .pdf_accessibility import apply_pdf_accessibility_settings

                apply_pdf_accessibility_settings(
                    input_pdf_path=output_path,
                    output_pdf_path=output_path,
                    field_tooltips=merged_field_tooltips,
                    field_order=ordered_names,
                    image_alt_text={
                        str(key): str(value)
                        for key, value in image_alt_text_payload.items()
                        if str(key).strip()
                    },
                    metadata={
                        "language": str(metadata_payload.get("language") or "").strip(),
                        "title": str(metadata_payload.get("title") or "").strip(),
                        "author": str(metadata_payload.get("author") or "").strip(),
                        "subject": str(metadata_payload.get("subject") or "").strip(),
                    },
                    auto_fill_missing_tooltips=parse_bool(
                        accessibility_payload.get("auto_fill_missing_tooltips"),
                        default=True,
                    ),
                )

            # Read the output file
            with open(output_path, "rb") as f:
                output_bytes = f.read()

            output_filename = filename.replace(".pdf", "-with-fields.pdf")

            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": output_filename,
                        "pdf_base64": base64.b64encode(output_bytes).decode("ascii"),
                    },
                }
            )
        finally:
            if os.path.exists(input_path):
                os.remove(input_path)
            if output_path is not None and os.path.exists(output_path):
                os.remove(output_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route("/pdf-labeler/api/rename-fields", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/rename-fields", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_rename_fields() -> Response:
    """Rename fields in an existing PDF.

    Returns:
        Response: A JSON response containing the renamed PDF as base64.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: rename-fields request {request_id}", "info")

    try:
        import formfyxer  # type: ignore[import-not-found]

        filename, content, post_data = _read_pdf_labeler_file_request()
        mapping_raw = post_data.get("mapping")

        # Parse mapping
        if isinstance(mapping_raw, str):
            mapping = json.loads(mapping_raw)
        elif isinstance(mapping_raw, dict):
            mapping = mapping_raw
        else:
            raise DashboardAPIValidationError(
                "mapping is required and must be an object."
            )

        # Write to temp files for processing
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
            tmp_in.write(content)
            input_path = tmp_in.name

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_out:
            output_path = tmp_out.name

        try:
            formfyxer.rename_pdf_fields(input_path, output_path, mapping)

            # Read the output file
            with open(output_path, "rb") as f:
                output_bytes = f.read()

            output_filename = filename.replace(".pdf", "-renamed.pdf")

            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": output_filename,
                        "pdf_base64": base64.b64encode(output_bytes).decode("ascii"),
                    },
                }
            )
        finally:
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


# ---------------------------------------------------------------------------
# PDF repair
# ---------------------------------------------------------------------------


@app.route("/pdf-labeler/api/repair", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/repair", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_repair() -> Response:
    """Run a single PDF repair action and return the repaired file.

    Returns:
        Response: A JSON response containing repair metadata and repaired PDF bytes.
    """
    from .pdf_repair import PDFRepairError, list_repair_actions, run_repair

    request_id = str(uuid.uuid4())

    try:
        post_data = request.get_json(silent=True) or {}
        form_data = dict(request.form)
        merged = {**form_data, **post_data}

        action = str(merged.get("action") or "").strip()
        if not action:
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {"available_actions": list_repair_actions()},
                }
            )

        # Read PDF content
        if "file" in request.files:
            upload = request.files["file"]
            filename = upload.filename or "upload.pdf"
            content = upload.read()
        else:
            filename = str(merged.get("filename") or "upload.pdf")
            b64 = merged.get("file_content_base64")
            if not b64:
                raise DashboardAPIValidationError(
                    "A PDF file is required (upload or file_content_base64)."
                )
            content = decode_base64_content(b64)

        _validate_upload_size(content)
        if not filename.lower().endswith(".pdf"):
            raise DashboardAPIValidationError(
                "Only PDF files are supported.", status_code=415
            )

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
            tmp_in.write(content)
            input_path = tmp_in.name
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_out:
            output_path = tmp_out.name
        if os.path.exists(output_path):
            os.remove(output_path)

        repair_options = _extract_repair_options(merged, action)

        try:
            result = run_repair(action, input_path, output_path, options=repair_options)
            with open(output_path, "rb") as fh:
                output_bytes = fh.read()
            output_filename = f"repaired_{filename}"
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": output_filename,
                        "pdf_base64": base64.b64encode(output_bytes).decode("ascii"),
                        "repair_result": result,
                    },
                }
            )
        finally:
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except PDFRepairError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "repair_error", "message": str(exc)},
            },
            400,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route("/pdf-labeler/api/copy-fields", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/copy-fields", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_copy_fields() -> Response:
    """Copy form field positions from a source PDF onto a destination PDF.

    Returns:
        Response: A JSON response containing the destination PDF with copied fields.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: copy-fields request {request_id}", "info")

    try:
        import formfyxer  # type: ignore[import-not-found]

        source_file = request.files.get("source")
        dest_file = request.files.get("destination")
        if not source_file or not dest_file:
            raise DashboardAPIValidationError(
                "Both 'source' and 'destination' PDF files are required."
            )

        source_bytes = source_file.read()
        dest_bytes = dest_file.read()
        _validate_upload_size(source_bytes)
        _validate_upload_size(dest_bytes)

        source_name = source_file.filename or "source.pdf"
        dest_name = dest_file.filename or "destination.pdf"
        if not source_name.lower().endswith(".pdf") or not dest_name.lower().endswith(
            ".pdf"
        ):
            raise DashboardAPIValidationError(
                "Only PDF files are supported.", status_code=415
            )

        with (
            tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_src,
            tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_dst,
        ):
            tmp_src.write(source_bytes)
            source_path = tmp_src.name
            tmp_dst.write(dest_bytes)
            dest_path = tmp_dst.name

        output_path: Optional[str] = None
        try:
            result_pdf = formfyxer.swap_pdf_page(
                source_pdf=source_path, destination_pdf=dest_path
            )
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_out:
                output_path = tmp_out.name
            result_pdf.save(output_path)

            # Remove any border/color metadata that swap_pdf_page may have
            # introduced (e.g. red /MK[BC] on checkboxes), while preserving the
            # existing /AP appearance streams on button fields so that checkbox
            # and radio rendering is not broken.
            _apply_pdf_field_visual_defaults(
                output_path, preserve_button_appearances=True
            )

            # Preserve the source PDF's per-page tab order setting (/Tabs).
            # Without this, viewers fall back to visual row order even though
            # copy_pdf_fields already copied the Annots array in source order.
            import pikepdf as _pikepdf

            with (
                _pikepdf.open(source_path) as _src_pdf,
                _pikepdf.open(output_path, allow_overwriting_input=True) as _out_pdf,
            ):
                _src_page_count = len(_src_pdf.pages)
                for _page_index, _out_page in enumerate(_out_pdf.pages):
                    if (
                        _page_index < _src_page_count
                        and "/Tabs" in _src_pdf.pages[_page_index]
                    ):
                        _out_page["/Tabs"] = _src_pdf.pages[_page_index]["/Tabs"]
                    elif "/Tabs" in _out_page:
                        del _out_page["/Tabs"]  # type: ignore[operator]
                _out_pdf.save(output_path)

            from .pdf_accessibility import (
                apply_pdf_accessibility_settings,
                extract_pdf_field_tooltips,
            )

            source_field_tooltips = extract_pdf_field_tooltips(source_path)
            if source_field_tooltips:
                apply_pdf_accessibility_settings(
                    input_pdf_path=output_path,
                    output_pdf_path=output_path,
                    field_tooltips=source_field_tooltips,
                    auto_fill_missing_tooltips=False,
                )

            with open(output_path, "rb") as fh:
                output_bytes = fh.read()
            output_filename = dest_name.replace(".pdf", "-with-fields.pdf")
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": output_filename,
                        "pdf_base64": base64.b64encode(output_bytes).decode("ascii"),
                    },
                }
            )
        finally:
            for p in (source_path, dest_path):
                if os.path.exists(p):
                    os.remove(p)
            if output_path is not None and os.path.exists(output_path):
                os.remove(output_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route("/pdf-labeler/api/strip-fonts", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/strip-fonts", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_strip_fonts() -> Response:
    """Remove embedded font programs from a PDF.

    Returns:
        Response: A JSON response containing the stripped PDF and removal stats.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: strip-fonts request {request_id}", "info")

    try:
        from .pdf_repair import strip_embedded_fonts

        filename, content, _post_data = _read_pdf_labeler_file_request()

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
            tmp_in.write(content)
            input_path = tmp_in.name
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_out:
            output_path = tmp_out.name

        try:
            result = strip_embedded_fonts(input_path, output_path)
            with open(output_path, "rb") as fh:
                output_bytes = fh.read()
            return jsonify(
                {
                    "success": True,
                    "request_id": request_id,
                    "data": {
                        "filename": filename,
                        "pdf_base64": base64.b64encode(output_bytes).decode("ascii"),
                        "fonts_removed": result.get("fonts_removed", 0),
                    },
                }
            )
        finally:
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )


@app.route("/pdf-labeler/api/bulk-normalize", methods=["POST"])
@app.route(f"{LABELER_BASE_PATH}/pdf-labeler/api/bulk-normalize", methods=["POST"])
@csrf.exempt
@cross_origin(origins="*", methods=["POST", "HEAD"], automatic_options=True)
def pdf_labeler_bulk_normalize():
    """Normalize one or more PDFs and return a zip archive.

    Accepts multipart form data with one or more ``files`` and a JSON
    ``options`` field describing which normalization steps to apply.
    """
    request_id = str(uuid.uuid4())
    log(f"ALDashboard: bulk-normalize request {request_id}", "info")

    try:
        import pikepdf
        import formfyxer  # type: ignore[import-not-found]
        from formfyxer.pdf_wrangling import FormField, FieldType, set_fields
        from reportlab.lib.colors import HexColor
        from .pdf_repair import strip_embedded_fonts

        # Read options
        options_raw = request.form.get("options", "{}")
        try:
            options = json.loads(options_raw) if isinstance(options_raw, str) else {}
        except json.JSONDecodeError:
            options = {}

        # Read uploaded files
        uploaded_files = request.files.getlist("files")
        if not uploaded_files:
            raise DashboardAPIValidationError("At least one PDF file is required.")

        file_entries: list[dict[str, Any]] = []
        for upload in uploaded_files:
            filename = upload.filename or "upload.pdf"
            content = upload.read()
            _validate_upload_size(content)
            if not filename.lower().endswith(".pdf"):
                continue
            file_entries.append({"filename": filename, "content": content})

        if not file_entries:
            raise DashboardAPIValidationError("No valid PDF files were uploaded.")

        # Parse normalization options with defaults matching the UI
        norm_font = parse_bool(options.get("normalizeFont"), default=True)
        norm_font_name = (
            str(options.get("fontName") or "Helvetica").strip() or "Helvetica"
        )
        norm_font_size = parse_bool(options.get("normalizeFontSize"), default=True)
        font_size_pt = int(options.get("fontSizePt") or 10)
        norm_checkbox_style = parse_bool(
            options.get("normalizeCheckboxStyle"), default=True
        )
        checkbox_style = str(options.get("checkboxStyle") or "cross").strip() or "cross"
        checkbox_export_value = (
            str(options.get("checkboxExportValue") or "Yes").strip() or "Yes"
        )
        uniform_checkbox_size = parse_bool(
            options.get("uniformCheckboxSize"), default=True
        )
        checkbox_size_pt = int(options.get("checkboxSizePt") or 12)
        auto_size_name_address = parse_bool(
            options.get("autoSizeNameAddress"), default=True
        )
        fixed_text_height_pt = int(options.get("fixedTextHeightPt") or 14)
        remove_embedded_fonts_flag = parse_bool(
            options.get("removeEmbeddedFonts"), default=False
        )

        import zipfile as _zipfile

        zip_buffer = io.BytesIO()
        processed_count = 0
        errors: list[str] = []

        with _zipfile.ZipFile(zip_buffer, "w", _zipfile.ZIP_DEFLATED) as zf:
            for entry in file_entries:
                fname = entry["filename"]
                content = entry["content"]
                try:
                    with tempfile.NamedTemporaryFile(
                        suffix=".pdf", delete=False
                    ) as tmp_in:
                        tmp_in.write(content)
                        input_path = tmp_in.name

                    try:
                        # Detect existing fields
                        with pikepdf.open(input_path) as pdf_doc:
                            page_count = len(pdf_doc.pages)
                            page_sizes = []
                            for page in pdf_doc.pages:
                                mbox = page.get("/MediaBox")
                                if mbox:
                                    coords = [float(v) for v in mbox]
                                    page_sizes.append(
                                        {
                                            "width": coords[2] - coords[0],
                                            "height": coords[3] - coords[1],
                                        }
                                    )
                                else:
                                    page_sizes.append({"width": 612, "height": 792})

                        # Use formfyxer to detect fields
                        detected = formfyxer.get_existing_pdf_fields(input_path)
                        if not detected:
                            detected = []

                        # Build normalized field definitions
                        normalized_fields: list[dict[str, Any]] = []
                        for field in detected:
                            f: dict[str, Any] = (
                                dict(field) if isinstance(field, dict) else {}
                            )
                            field_name = str(f.get("name", f.get("var_name", "field")))
                            field_type_str = str(f.get("type", "text")).lower()
                            page_idx = int(f.get("page", f.get("pageIndex", 0)))
                            if page_idx >= page_count:
                                page_idx = 0

                            nf: dict[str, Any] = {
                                "name": field_name,
                                "type": field_type_str,
                                "pageIndex": page_idx,
                                "x": float(f.get("x", 0)),
                                "y": float(f.get("y", 0)),
                                "width": float(f.get("width", 100)),
                                "height": float(f.get("height", 20)),
                                "font": (
                                    norm_font_name
                                    if norm_font
                                    else str(f.get("font", "Helvetica"))
                                ),
                                "fontSize": (
                                    font_size_pt
                                    if norm_font_size
                                    else int(f.get("fontSize", 12) or 12)
                                ),
                                "autoSize": False,
                            }

                            if field_type_str == "checkbox" and norm_checkbox_style:
                                nf["checkboxStyle"] = checkbox_style
                                nf["checkboxExportValue"] = checkbox_export_value
                            if field_type_str == "checkbox" and uniform_checkbox_size:
                                nf["width"] = checkbox_size_pt
                                nf["height"] = checkbox_size_pt
                            if (
                                auto_size_name_address
                                and field_type_str == "text"
                                and _looks_like_name_email_address_phone_field(
                                    field_name
                                )
                            ):
                                nf["autoSize"] = True
                                nf["height"] = fixed_text_height_pt

                            normalized_fields.append(nf)

                        if not normalized_fields:
                            # No fields to normalize; include as-is
                            zf.writestr(fname, content)
                            processed_count += 1
                            continue

                        fields_per_page = build_pdf_export_fields_per_page(
                            normalized_fields,
                            page_count=page_count,
                            form_field_cls=FormField,
                            field_type_enum=FieldType,
                            color_parser=HexColor,
                        )
                        explicit_background_fields = (
                            _collect_fields_with_explicit_background(normalized_fields)
                        )

                        with tempfile.NamedTemporaryFile(
                            suffix=".pdf", delete=False
                        ) as tmp_out:
                            output_path = tmp_out.name

                        set_fields(
                            input_path, output_path, fields_per_page, overwrite=True
                        )

                        # Apply checkbox export values
                        checkbox_values = {
                            nf["name"]: nf.get("checkboxExportValue", "")
                            for nf in normalized_fields
                            if nf.get("type") == "checkbox"
                            and nf.get("checkboxExportValue")
                        }
                        if checkbox_values:
                            _apply_checkbox_export_values(output_path, checkbox_values)
                        _apply_pdf_field_visual_defaults(
                            output_path,
                            explicit_background_fields=explicit_background_fields,
                        )

                        # Strip embedded fonts if requested
                        if remove_embedded_fonts_flag:
                            with tempfile.NamedTemporaryFile(
                                suffix=".pdf", delete=False
                            ) as tmp_stripped:
                                stripped_path = tmp_stripped.name
                            strip_embedded_fonts(output_path, stripped_path)
                            os.remove(output_path)
                            output_path = stripped_path

                        with open(output_path, "rb") as fh:
                            zf.writestr(fname, fh.read())
                        os.remove(output_path)
                        processed_count += 1
                    finally:
                        if os.path.exists(input_path):
                            os.remove(input_path)
                except Exception as exc:
                    errors.append(f"{fname}: {exc}")

        zip_buffer.seek(0)
        zip_b64 = base64.b64encode(zip_buffer.read()).decode("ascii")

        return jsonify(
            {
                "success": True,
                "request_id": request_id,
                "data": {
                    "zip_base64": zip_b64,
                    "filename": "normalized-pdfs.zip",
                    "processed": processed_count,
                    "total": len(file_entries),
                    "errors": errors,
                },
            }
        )

    except DashboardAPIValidationError as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "validation_error", "message": exc.message},
            },
            exc.status_code,
        )
    except Exception as exc:
        return jsonify_with_status(
            {
                "success": False,
                "request_id": request_id,
                "error": {"type": "server_error", "message": str(exc)},
            },
            500,
        )
