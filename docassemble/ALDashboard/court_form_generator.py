"""Draft a court document from the variables an interview already gathers.

The variable report answers "what does this interview collect?". A court form
answers "what does the court want to see?", and the two are close enough that
the report's own extraction can drive both: the interview's screens become the
numbered middle of the document, and the jurisdiction profile supplies the
caption, running footer, signature block and certificate of service that the
court fixes (SuffolkLITLab/docassemble-ALDashboard#272).

Four shapes are drafted here. They share the caption and signature machinery and
differ in what surrounds the body:

``court_form``
    Caption, title, the interview's screens as numbered sections, signature.
``motion``
    The same, wrapped in a "NOW COMES" introduction and a prayer for relief,
    and followed by a certificate of service.
``affidavit``
    The same, introduced as sworn statements and closed with a jurat.
``letter``
    No caption at all: sender block, date, recipient, RE: line, body, closing.

The output is a docassemble DOCX **template**, not a filled document. Two
conventions apply to the Jinja it contains, and they differ on purpose:

* Boilerplate named by the profile -- ``{{ trial_court }}``,
  ``{{ users[0].signature }}`` -- is written through as-is. These names are the
  profile's contract with the interview, they read the way a human would write
  them, and they match what hand-authored AssemblyLine court templates use.
* Fields discovered in the interview are wrapped in ``showifdef()``, because
  whether any given one is defined at assembly time is exactly what we do not
  know. This matches the intake report's existing behavior.
"""

import re
from typing import Any, Dict, List, Mapping, Optional, Sequence, Tuple

from .court_form_profiles import (
    CourtFormProfile,
    apply_profile_styles,
    find_docx_fragment,
    load_court_form_profile,
    splice_docx_fragment,
)

try:
    import docx
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.document import Document as _DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError:  # pragma: no cover - python-docx is a hard dependency at runtime
    docx = None  # type: ignore[assignment]

__all__ = [
    "COURT_FORM_SHAPES",
    "SHAPE_LABELS",
    "build_court_form_context",
    "generate_court_form_docx",
    "generate_court_form_markdown",
    "render_blocks",
]

# Shapes this module can draft. "intake" stays with the variable report.
COURT_FORM_SHAPES: Sequence[str] = ("court_form", "motion", "affidavit", "letter")

SHAPE_LABELS: Dict[str, str] = {
    "intake": "Intake summary (field and value tables)",
    "court_form": "Court form (caption, numbered body, signature)",
    "motion": "Motion (caption, grounds, prayer for relief, certificate of service)",
    "affidavit": "Affidavit (caption, sworn statements, jurat)",
    "letter": "Letter (no caption; sender, recipient, body, closing)",
}

# Placeholders resolved while drafting rather than left for docassemble. These
# name things the drafter knows now -- the title of this document, what this
# court calls its docket number -- so baking them in produces a template an
# author can read instead of a chain of indirections.
_DRAFT_TIME_KEYS: Sequence[str] = (
    "document_title",
    "docket_label",
    "court_name",
    "jurisdiction",
    "form_code",
    "form_revision",
    "court_rule_citation",
)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_\.]*)\s*\}\}")

# Attributes that are noise in a court document body.
_SKIP_BODY_ATTRS = {"instanceName", "elements", "parent", "attr_name"}


def _require_docx() -> None:
    if docx is None:
        raise RuntimeError("python-docx is not installed in the python environment.")


def build_court_form_context(
    profile: CourtFormProfile,
    document_title: str,
    extra: Optional[Mapping[str, Any]] = None,
) -> Dict[str, str]:
    """The draft-time substitutions for one document.

    ``labels.plaintiff`` and friends resolve from the profile so a caption can
    say "Defendant(s)/Respondent(s)" in Michigan and "Defendant" in Vermont
    without either profile restating the other's wording.
    """
    context: Dict[str, str] = {
        "document_title": str(document_title or "").strip(),
        "docket_label": profile.label("docket", "Case No."),
        "court_name": profile.name,
        "jurisdiction": profile.jurisdiction,
        # Form metadata nobody has supplied yet. A visible bracket tells the
        # author what to fill in; an empty string would just leave a gap in the
        # footer that reads like a bug.
        "form_code": "[form number]",
        "form_revision": "[revision date]",
        "court_rule_citation": "[court rule citation]",
    }
    for key, value in profile.labels.items():
        context[f"labels.{key}"] = str(value)
    for key in _DRAFT_TIME_KEYS:
        if key in profile.labels:
            context[key] = str(profile.labels[key])
    if extra:
        for key, value in extra.items():
            if value is not None and str(value) != "":
                context[str(key)] = str(value)
    return context


def _substitute(text: str, context: Mapping[str, str]) -> str:
    """Fill draft-time placeholders and leave interview variables alone."""

    def replace(match: "re.Match[str]") -> str:
        name = match.group(1)
        if name in context:
            return str(context[name])
        return match.group(0)

    return _PLACEHOLDER_RE.sub(replace, str(text or ""))


def _set_table_borders(table, kind: str) -> None:
    """Apply one of the border treatments a caption needs.

    python-docx has no border API, so this writes ``w:tblBorders`` directly.
    """
    kind = str(kind or "none").strip().lower()
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")

    if kind in ("box", "grid"):
        edges = ["top", "left", "bottom", "right"]
        if kind == "grid":
            edges += ["insideH", "insideV"]
        for edge in edges:
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:color"), "000000")
            borders.append(element)
        if kind == "box":
            for edge in ("insideH", "insideV"):
                element = OxmlElement(f"w:{edge}")
                element.set(qn("w:val"), "none")
                borders.append(element)
    elif kind == "bottom":
        for edge in ("top", "left", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "none")
            borders.append(element)
        element = OxmlElement("w:bottom")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    else:
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "none")
            borders.append(element)

    tbl_pr.append(borders)


def _add_page_number_field(paragraph, align: Optional[str] = None) -> None:
    """Append a live "Page X of Y" field so pagination follows the document."""
    if align:
        alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(str(align).lower())
        if alignment is not None:
            paragraph.alignment = alignment

    if paragraph.text:
        paragraph.add_run("    ")
    paragraph.add_run("Page ")
    for instruction in ("PAGE", "NUMPAGES"):
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {instruction} "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)
        if instruction == "PAGE":
            paragraph.add_run(" of ")


def _add_horizontal_rule(paragraph) -> None:
    """A bottom border on an empty paragraph: the line under a caption."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)


def _style_or_none(document, name: Optional[str]) -> Optional[str]:
    """A style name the document actually has, so a typo cannot raise."""
    if not name:
        return None
    try:
        document.styles[str(name)]
        return str(name)
    except KeyError:
        return None


def _add_paragraph(container, document, text: str, style: Optional[str]):
    resolved = _style_or_none(document, style)
    paragraph = container.add_paragraph(text)
    if resolved:
        paragraph.style = document.styles[resolved]
    return paragraph


def render_blocks(
    document,
    container,
    blocks: Sequence[Mapping[str, Any]],
    profile: CourtFormProfile,
    context: Mapping[str, str],
) -> int:
    """Draw a profile's declarative blocks into ``container``.

    ``container`` is the document body, a header or a footer -- all of them take
    ``add_paragraph`` and ``add_table``. Returns the number of blocks drawn.
    """
    _require_docx()
    drawn = 0

    for block in blocks:
        if not isinstance(block, Mapping):
            continue
        block_type = str(block.get("type") or "paragraph").strip().lower()

        if block_type == "spacer":
            _add_paragraph(container, document, "", block.get("style"))
            drawn += 1

        elif block_type == "rule":
            paragraph = _add_paragraph(container, document, "", block.get("style"))
            _add_horizontal_rule(paragraph)
            drawn += 1

        elif block_type == "page_break":
            paragraph = container.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            drawn += 1

        elif block_type == "paragraph":
            text = _substitute(block.get("text", ""), context)
            paragraph = _add_paragraph(container, document, text, block.get("style"))
            align = block.get("align")
            if align:
                alignment = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }.get(str(align).lower())
                if alignment is not None:
                    paragraph.alignment = alignment
            if block.get("page_numbers"):
                _add_page_number_field(
                    paragraph, block.get("page_numbers_align") or "right"
                )
            drawn += 1

        elif block_type == "table":
            rows = block.get("rows")
            if not isinstance(rows, list) or not rows:
                continue
            widths = block.get("widths")
            column_count = max(
                len(row.get("cells", []) if isinstance(row, Mapping) else row)
                for row in rows
            )
            if isinstance(widths, list) and widths:
                column_count = max(column_count, len(widths))

            # Document.add_table sizes itself from the section; a header or
            # footer is a bare block container and insists on an explicit width.
            if isinstance(container, _DocxDocument):
                table = container.add_table(rows=0, cols=column_count)
            else:
                total_width = 6.5
                if isinstance(widths, list) and widths:
                    try:
                        total_width = sum(float(w) for w in widths)
                    except (TypeError, ValueError):
                        pass
                table = container.add_table(
                    rows=0, cols=column_count, width=Inches(total_width)
                )
            try:
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.autofit = False
            except Exception:
                pass
            _set_table_borders(table, block.get("borders", "none"))

            default_style = block.get("style") or "CourtCaptionText"
            for row_spec in rows:
                if isinstance(row_spec, Mapping):
                    cells = row_spec.get("cells") or []
                    row_style = row_spec.get("style") or default_style
                else:
                    cells = row_spec
                    row_style = default_style
                row = table.add_row()
                for index in range(column_count):
                    raw = cells[index] if index < len(cells) else ""
                    text = _substitute(raw, context)
                    cell = row.cells[index]
                    # A cell starts with one empty paragraph; write the first
                    # line into it so no blank line precedes the content.
                    lines = str(text).split("\n")
                    resolved = _style_or_none(document, row_style)
                    first = cell.paragraphs[0]
                    first.text = lines[0]
                    if resolved:
                        first.style = document.styles[resolved]
                    for line in lines[1:]:
                        extra = cell.add_paragraph(line)
                        if resolved:
                            extra.style = document.styles[resolved]

            if isinstance(widths, list):
                for index, width in enumerate(widths):
                    if index >= column_count:
                        break
                    try:
                        inches = Inches(float(width))
                    except (TypeError, ValueError):
                        continue
                    for row in table.rows:
                        row.cells[index].width = inches
            drawn += 1

    return drawn


def _render_section(
    document,
    section: str,
    profile: CourtFormProfile,
    context: Mapping[str, str],
    container=None,
) -> str:
    """Draw one fixed section, preferring a Word-authored override.

    Returns "docx", "yaml" or "" so callers can report which template was used.
    """
    fragment = find_docx_fragment(profile.id, section, profile.fragment_dirs)
    if fragment:
        # A header or footer is its own block container; the body splice has to
        # respect the trailing sectPr, which is why it is handled separately.
        target_element = getattr(container, "_element", None) if container else None
        splice_docx_fragment(document, fragment, container=target_element)
        return "docx"
    blocks = profile.section_blocks(section)
    if not blocks:
        return ""
    target = container if container is not None else document
    render_blocks(document, target, blocks, profile, context)
    return "yaml"


def _field_expression(field_spec) -> str:
    """The Jinja for one discovered field, guarded against being undefined."""
    if getattr(field_spec, "datatype", "") in ("yesno", "boolean"):
        return (
            "{% if showifdef('"
            + field_spec.var_name
            + "') %}Yes{% else %}No{% endif %}"
        )
    return "{{ showifdef('" + field_spec.var_name + "') }}"


def _add_list_table(document, container, list_spec, profile: CourtFormProfile) -> None:
    """A repeating table for one of the interview's lists."""
    attributes = [
        attribute
        for attribute in list_spec.attributes
        if attribute.name not in _SKIP_BODY_ATTRS
    ]
    if not attributes:
        return

    table = container.add_table(rows=0, cols=len(attributes) + 1)
    _set_table_borders(table, "grid")
    label_style = _style_or_none(document, profile.style_name("label_style", "Normal"))
    text_style = _style_or_none(document, profile.style_name("text_style", "Normal"))

    header = table.add_row()
    header.cells[0].text = "#"
    for index, attribute in enumerate(attributes, start=1):
        header.cells[index].text = attribute.description or attribute.name
    if label_style:
        for cell in header.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = document.styles[label_style]

    # Standalone loop rows keep docxtpl's row repetition working.
    loop_start = table.add_row()
    loop_start.cells[0].text = "{%tr for item in showifdef('" + list_spec.name + "') %}"

    data = table.add_row()
    data.cells[0].text = "{{ loop.index }}"
    for index, attribute in enumerate(attributes, start=1):
        if attribute.name == "item":
            data.cells[index].text = "{{ item }}"
        else:
            data.cells[index].text = (
                "{{ showifdef(item.attr_name('" + attribute.name + "')) }}"
            )
    if text_style:
        for cell in data.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = document.styles[text_style]

    loop_end = table.add_row()
    loop_end.cells[0].text = "{%tr endfor %}"


def _add_numbered_body(
    document,
    groups: Sequence[Any],
    profile: CourtFormProfile,
    *,
    numbered: bool = True,
    show_headings: bool = True,
) -> int:
    """The middle of the document: the interview's screens, in order.

    Returns the number of numbered paragraphs written, so a caller can tell
    whether the draft has any substance to it.
    """
    heading_style = profile.style_name("heading_style", "Heading 2")
    text_style = profile.style_name("text_style", "Normal")
    counter = 0

    for group in groups:
        fields = list(getattr(group, "fields", []) or [])
        lists = list(getattr(group, "lists", []) or [])
        if not fields and not lists:
            continue

        if show_headings:
            _add_paragraph(document, document, group.title, heading_style)

        for field_spec in fields:
            counter += 1
            prefix = f"{counter}. " if numbered else ""
            text = f"{prefix}{field_spec.label}: {_field_expression(field_spec)}"
            _add_paragraph(document, document, text, text_style)

        for list_spec in lists:
            _add_paragraph(
                document,
                document,
                list_spec.label or list_spec.name,
                heading_style,
            )
            _add_list_table(document, document, list_spec, profile)
            _add_paragraph(document, document, "", text_style)

    return counter


def generate_court_form_docx(
    groups: Sequence[Any],
    variables: Mapping[str, Any],
    *,
    shape: str = "court_form",
    profile: Optional[CourtFormProfile] = None,
    profile_id: Optional[str] = None,
    document_title: Optional[str] = None,
    output_path: Optional[str] = None,
    include_certificate_of_service: Optional[bool] = None,
    numbered_paragraphs: Optional[bool] = None,
    fragment_dirs: Optional[Sequence[str]] = None,
    context_extra: Optional[Mapping[str, Any]] = None,
) -> Dict[str, Any]:
    """Draft one court document and return it with a description of what it is.

    The result's ``sections`` map says, for each fixed section, whether it came
    from the profile YAML or from a Word-authored ``.docx`` override -- which is
    the question people ask first when a caption does not look the way they
    expected.
    """
    _require_docx()

    shape_name = str(shape or "court_form").strip().lower()
    if shape_name not in COURT_FORM_SHAPES:
        raise ValueError(
            f"Unknown court form shape '{shape}'. "
            f"Choose one of: {', '.join(COURT_FORM_SHAPES)}."
        )

    if profile is None:
        profile = load_court_form_profile(profile_id, fragment_dirs=fragment_dirs)
    elif fragment_dirs:
        profile.fragment_dirs = [str(d) for d in fragment_dirs]

    title = str(document_title or "").strip() or "Court Document Draft"
    context = build_court_form_context(profile, title, context_extra)

    document = docx.Document()
    apply_profile_styles(document, profile)

    section = document.sections[0]
    section.top_margin = Inches(profile.page.top_margin)
    section.bottom_margin = Inches(profile.page.bottom_margin)
    section.left_margin = Inches(profile.page.left_margin)
    section.right_margin = Inches(profile.page.right_margin)

    used: Dict[str, str] = {}

    if numbered_paragraphs is None:
        numbered_paragraphs = bool(profile.body.get("numbered_paragraphs", True))

    if shape_name == "letter":
        used["letterhead"] = _render_section(document, "letterhead", profile, context)
        _build_letter_body(document, groups, profile, context, title)
    else:
        used["caption"] = _render_section(document, "caption", profile, context)
        _add_paragraph(
            document, document, title, profile.style_name("title_style", "Heading 1")
        )

        text_style = profile.style_name("text_style", "Normal")
        if shape_name == "motion" and profile.body.get("motion_intro"):
            _add_paragraph(
                document,
                document,
                _substitute(profile.body["motion_intro"], context),
                text_style,
            )
        elif shape_name == "affidavit" and profile.body.get("affidavit_intro"):
            _add_paragraph(
                document,
                document,
                _substitute(profile.body["affidavit_intro"], context),
                text_style,
            )

        _add_numbered_body(
            document, groups, profile, numbered=bool(numbered_paragraphs)
        )

        if shape_name == "motion" and profile.body.get("motion_prayer"):
            _add_paragraph(
                document,
                document,
                _substitute(profile.body["motion_prayer"], context),
                text_style,
            )

        if shape_name == "affidavit":
            closing = profile.body.get("affidavit_closing")
            if closing:
                _add_paragraph(
                    document, document, _substitute(closing, context), text_style
                )
            used["jurat"] = _render_section(document, "jurat", profile, context)
        else:
            used["signature"] = _render_section(document, "signature", profile, context)

        if include_certificate_of_service is None:
            include_certificate_of_service = shape_name == "motion"
        if include_certificate_of_service:
            used["certificate_of_service"] = _render_section(
                document, "certificate_of_service", profile, context
            )

    # Running header and footer belong to the section, not the body flow. A
    # letter is correspondence rather than a filing, so the court's form number
    # and "Page 1 of 2" footer would be wrong on it.
    if shape_name != "letter":
        for name, container in (("header", section.header), ("footer", section.footer)):
            origin = _render_section(
                document, name, profile, context, container=container
            )
            if origin:
                used[name] = origin

    path = output_path
    if not path:
        import os
        import tempfile

        path = os.path.join(tempfile.mkdtemp(prefix="court_form_"), "court_form.docx")
    document.save(path)

    return {
        "docx_path": path,
        "shape": shape_name,
        "profile_id": profile.id,
        "profile_name": profile.name,
        "document_title": title,
        "sections": {key: value for key, value in used.items() if value},
        "styles": sorted(profile.styles.keys()),
    }


def _build_letter_body(
    document,
    groups: Sequence[Any],
    profile: CourtFormProfile,
    context: Mapping[str, str],
    title: str,
) -> None:
    """A letter has no caption; it has correspondents."""
    text_style = profile.style_name("text_style", "Normal")
    signature_style = profile.style_name("signature_style", text_style)

    for line in (
        "{{ users[0] }}",
        "{{ users[0].address.block() }}",
        "",
        "{{ letter_date }}",
        "",
        "{{ recipient }}",
        "{{ recipient.address.block() }}",
        "",
    ):
        _add_paragraph(document, document, line, text_style)

    _add_paragraph(
        document,
        document,
        f"RE: {title}",
        profile.style_name("label_style", text_style),
    )
    _add_paragraph(document, document, "", text_style)
    _add_paragraph(document, document, "Dear {{ recipient }}:", text_style)

    _add_numbered_body(document, groups, profile, numbered=False, show_headings=True)

    for line in ("", "Sincerely,", "", "{{ users[0].signature }}", "{{ users[0] }}"):
        _add_paragraph(document, document, line, signature_style)


def generate_court_form_markdown(
    groups: Sequence[Any],
    variables: Mapping[str, Any],
    *,
    shape: str = "court_form",
    profile: Optional[CourtFormProfile] = None,
    profile_id: Optional[str] = None,
    document_title: Optional[str] = None,
    context_extra: Optional[Mapping[str, Any]] = None,
) -> str:
    """A Mako + Markdown rendering of the same draft, for preview and for
    authors who assemble to Markdown rather than to DOCX."""
    if profile is None:
        profile = load_court_form_profile(profile_id)
    shape_name = str(shape or "court_form").strip().lower()
    title = str(document_title or "").strip() or "Court Document Draft"
    context = build_court_form_context(profile, title, context_extra)

    lines: List[str] = []

    def block_lines(section: str) -> None:
        for block in profile.section_blocks(section):
            block_type = str(block.get("type") or "paragraph").strip().lower()
            if block_type == "paragraph":
                text = _substitute(block.get("text", ""), context).strip()
                if text:
                    lines.append(text)
                    lines.append("")
            elif block_type == "table":
                for row_spec in block.get("rows") or []:
                    cells = (
                        row_spec.get("cells")
                        if isinstance(row_spec, Mapping)
                        else row_spec
                    ) or []
                    rendered = [
                        _substitute(cell, context).replace("\n", " ").strip()
                        for cell in cells
                    ]
                    if any(rendered):
                        lines.append(" | ".join(rendered))
                lines.append("")
            elif block_type == "rule":
                lines.append("---")
                lines.append("")

    if shape_name == "letter":
        lines.append("{{ users[0] }}")
        lines.append("")
        lines.append("{{ letter_date }}")
        lines.append("")
        lines.append("{{ recipient }}")
        lines.append("")
        lines.append(f"**RE: {title}**")
        lines.append("")
        lines.append("Dear {{ recipient }}:")
        lines.append("")
    else:
        block_lines("caption")
        lines.append(f"# {title}")
        lines.append("")
        if shape_name == "motion" and profile.body.get("motion_intro"):
            lines.append(_substitute(profile.body["motion_intro"], context))
            lines.append("")
        elif shape_name == "affidavit" and profile.body.get("affidavit_intro"):
            lines.append(_substitute(profile.body["affidavit_intro"], context))
            lines.append("")

    counter = 0
    for group in groups:
        fields = list(getattr(group, "fields", []) or [])
        lists = list(getattr(group, "lists", []) or [])
        if not fields and not lists:
            continue
        lines.append(f"## {group.title}")
        lines.append("")
        for field_spec in fields:
            counter += 1
            if getattr(field_spec, "datatype", "") in ("yesno", "boolean"):
                value = (
                    "% if showifdef('"
                    + field_spec.var_name
                    + "'): Yes % else: No % endif"
                )
            else:
                value = "${ showifdef('" + field_spec.var_name + "') }"
            prefix = f"{counter}. " if shape_name != "letter" else ""
            lines.append(f"{prefix}**{field_spec.label}:** {value}")
            lines.append("")
        for list_spec in lists:
            lines.append(f"### {list_spec.label or list_spec.name}")
            lines.append("")
            attributes = [
                attribute
                for attribute in list_spec.attributes
                if attribute.name not in _SKIP_BODY_ATTRS
            ]
            if not attributes:
                continue
            lines.append(
                "| # | "
                + " | ".join(a.description or a.name for a in attributes)
                + " |"
            )
            lines.append("| --- | " + " | ".join(["---"] * len(attributes)) + " |")
            lines.append(
                "% for i, item in enumerate(showifdef('" + list_spec.name + "') or []):"
            )
            cells = []
            for attribute in attributes:
                if attribute.name == "item":
                    cells.append("${ item }")
                else:
                    cells.append(
                        "${ showifdef(item.attr_name('" + attribute.name + "')) }"
                    )
            lines.append("| ${ i + 1 } | " + " | ".join(cells) + " |")
            lines.append("% endfor")
            lines.append("")

    if shape_name == "motion" and profile.body.get("motion_prayer"):
        lines.append(_substitute(profile.body["motion_prayer"], context))
        lines.append("")
    if shape_name == "affidavit":
        closing = profile.body.get("affidavit_closing")
        if closing:
            lines.append(_substitute(closing, context))
            lines.append("")
        block_lines("jurat")
    elif shape_name == "letter":
        lines.append("Sincerely,")
        lines.append("")
        lines.append("{{ users[0].signature }}")
        lines.append("")
    else:
        block_lines("signature")
        if shape_name == "motion":
            block_lines("certificate_of_service")

    return "\n".join(lines)
