from __future__ import annotations

import io
import os
import posixpath
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional
from xml.etree import ElementTree as ET

import docx

MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

NS = {
    "mc": MC_NS,
    "pr": REL_NS,
    "w": W_NS,
    "a": A_NS,
    "r": R_NS,
}

XML_EXTS = (".xml", ".rels")


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def is_xml_part(name: str) -> bool:
    return name.endswith(XML_EXTS)


def norm_part_path(base_part: str, target: str) -> str:
    if base_part.endswith(".rels"):
        rels_dir = posixpath.dirname(base_part)
        source_dir = posixpath.dirname(rels_dir)
    else:
        source_dir = posixpath.dirname(base_part)

    joined = posixpath.normpath(posixpath.join(source_dir, target))
    return joined.lstrip("/")


def choose_alternate_content(root: ET.Element) -> ET.Element:
    parent_map = {child: parent for parent in root.iter() for child in parent}

    for ac in list(root.findall(".//mc:AlternateContent", NS)):
        parent = parent_map.get(ac)
        if parent is None:
            continue

        replacement_children: List[ET.Element] = []
        fallback = ac.find("mc:Fallback", NS)
        if fallback is not None:
            replacement_children = list(fallback)
        else:
            choice = ac.find("mc:Choice", NS)
            if choice is not None:
                replacement_children = list(choice)

        idx = list(parent).index(ac)
        parent.remove(ac)
        for child in reversed(replacement_children):
            parent.insert(idx, child)

    return root


def repair_docx_xml_conservatively(src: str, dst: str) -> Dict[str, Any]:
    report: Dict[str, Any] = {
        "xml_parse_errors": [],
        "removed_relationships": [],
        "alternatecontent_rewritten": [],
        "copied_parts": [],
        "output_filename": Path(dst).name,
    }

    src_path = Path(src)
    dst_path = Path(dst)

    with tempfile.TemporaryDirectory() as td:
        work = Path(td) / "unzipped"
        work.mkdir()

        with zipfile.ZipFile(src_path, "r") as zf:
            zf.extractall(work)

        for path in work.rglob("*"):
            if not path.is_file():
                continue
            rel_name = path.relative_to(work).as_posix()
            if not is_xml_part(rel_name):
                continue

            try:
                tree = ET.parse(path)
                root = tree.getroot()
            except ET.ParseError as err:
                report["xml_parse_errors"].append(
                    {"part": rel_name, "error": str(err)}
                )
                continue

            before = ET.tostring(root, encoding="utf-8")
            root = choose_alternate_content(root)
            after = ET.tostring(root, encoding="utf-8")
            if before != after:
                tree = ET.ElementTree(root)
                tree.write(path, encoding="utf-8", xml_declaration=True)
                report["alternatecontent_rewritten"].append(rel_name)

        for rels_path in work.rglob("*.rels"):
            rel_name = rels_path.relative_to(work).as_posix()
            try:
                tree = ET.parse(rels_path)
                root = tree.getroot()
            except ET.ParseError:
                continue

            changed = False
            for rel in list(root):
                if _local_name(rel.tag) != "Relationship":
                    continue
                target = rel.attrib.get("Target")
                mode = rel.attrib.get("TargetMode")
                if not target or mode == "External":
                    continue

                resolved = norm_part_path(rel_name, target)
                if not (work / resolved).exists():
                    root.remove(rel)
                    changed = True
                    report["removed_relationships"].append(
                        {
                            "rels_part": rel_name,
                            "target": target,
                            "resolved": resolved,
                            "id": rel.attrib.get("Id"),
                        }
                    )

            if changed:
                tree.write(rels_path, encoding="utf-8", xml_declaration=True)

        if dst_path.exists():
            dst_path.unlink()

        with zipfile.ZipFile(dst_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for path in sorted(work.rglob("*")):
                if path.is_file():
                    arcname = path.relative_to(work).as_posix()
                    zf.write(path, arcname)
                    report["copied_parts"].append(arcname)

    return report


def roundtrip_docx_via_soffice(src: str, dst: str) -> Dict[str, Any]:
    soffice_path = shutil.which("soffice")
    if not soffice_path:
        return {
            "available": False,
            "message": "LibreOffice 'soffice' is not installed on this server.",
            "output_filename": Path(dst).name,
        }

    src_path = Path(src)
    dst_path = Path(dst)
    with tempfile.TemporaryDirectory() as td:
        temp_dir = Path(td)
        subprocess.run(
            [
                soffice_path,
                "--headless",
                "--convert-to",
                "rtf",
                "--outdir",
                str(temp_dir),
                str(src_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        rtf_path = temp_dir / (src_path.stem + ".rtf")
        if not rtf_path.exists():
            raise RuntimeError("LibreOffice did not produce an intermediate RTF file.")

        subprocess.run(
            [
                soffice_path,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(temp_dir),
                str(rtf_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        converted_path = temp_dir / (src_path.stem + ".docx")
        if not converted_path.exists():
            raise RuntimeError("LibreOffice did not produce a round-tripped DOCX file.")
        shutil.copyfile(converted_path, dst_path)

    return {
        "available": True,
        "message": "Round-tripped the DOCX through LibreOffice (DOCX -> RTF -> DOCX).",
        "output_filename": dst_path.name,
    }


def _read_relationship_targets(archive: zipfile.ZipFile) -> Dict[str, str]:
    try:
        rels_xml = archive.read("word/_rels/document.xml.rels")
    except KeyError:
        return {}
    root = ET.fromstring(rels_xml)
    mapping: Dict[str, str] = {}
    for rel in root:
        if _local_name(rel.tag) != "Relationship":
            continue
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target")
        if rel_id and target:
            mapping[rel_id] = norm_part_path("word/_rels/document.xml.rels", target)
    return mapping


def _apply_run_formatting(new_run: Any, run_element: ET.Element) -> None:
    properties = run_element.find("w:rPr", NS)
    if properties is None:
        return
    new_run.bold = properties.find("w:b", NS) is not None
    new_run.italic = properties.find("w:i", NS) is not None
    new_run.underline = properties.find("w:u", NS) is not None


def _extract_run_text(run_element: ET.Element) -> str:
    chunks: List[str] = []
    for child in run_element:
        name = _local_name(child.tag)
        if name == "t":
            chunks.append(child.text or "")
        elif name == "tab":
            chunks.append("\t")
        elif name == "br":
            chunks.append("\n")
    return "".join(chunks)


def _append_images_from_run(
    new_run: Any,
    run_element: ET.Element,
    archive: zipfile.ZipFile,
    rel_targets: Dict[str, str],
    report: Dict[str, Any],
) -> None:
    for blip in run_element.findall(".//a:blip", NS):
        rel_id = blip.attrib.get(f"{{{R_NS}}}embed")
        if not rel_id:
            continue
        target = rel_targets.get(rel_id)
        if not target:
            continue
        try:
            image_bytes = archive.read(target)
        except KeyError:
            report["errors"].append(
                {"type": "missing_image", "relationship_id": rel_id, "target": target}
            )
            continue
        try:
            new_run.add_picture(io.BytesIO(image_bytes))
            report["images"] += 1
        except Exception as err:
            report["errors"].append(
                {"type": "image_copy_failed", "target": target, "error": str(err)}
            )


def _append_paragraph_from_xml(
    target_doc: docx.document.Document,
    paragraph_element: ET.Element,
    archive: zipfile.ZipFile,
    rel_targets: Dict[str, str],
    report: Dict[str, Any],
) -> None:
    paragraph = target_doc.add_paragraph()
    for run_element in paragraph_element.findall(".//w:r", NS):
        text = _extract_run_text(run_element)
        run = paragraph.add_run(text)
        _apply_run_formatting(run, run_element)
        _append_images_from_run(run, run_element, archive, rel_targets, report)
    report["paragraphs"] += 1


def _paragraph_text_from_xml(paragraph_element: ET.Element) -> str:
    return "".join(_extract_run_text(run) for run in paragraph_element.findall(".//w:r", NS))


def _append_table_from_xml(target_doc: docx.document.Document, table_element: ET.Element, report: Dict[str, Any]) -> None:
    row_elements = table_element.findall("w:tr", NS)
    if not row_elements:
        return
    max_cells = max(len(row.findall("w:tc", NS)) for row in row_elements) or 1
    table = target_doc.add_table(rows=len(row_elements), cols=max_cells)
    for row_index, row_element in enumerate(row_elements):
        cell_elements = row_element.findall("w:tc", NS)
        for col_index, cell_element in enumerate(cell_elements):
            paragraphs = [
                _paragraph_text_from_xml(paragraph)
                for paragraph in cell_element.findall(".//w:p", NS)
            ]
            table.cell(row_index, col_index).text = "\n".join(
                paragraph for paragraph in paragraphs if paragraph
            )
    report["tables"] += 1


def rescue_docx_to_shell(src: str, dst: str) -> Dict[str, Any]:
    report: Dict[str, Any] = {
        "paragraphs": 0,
        "tables": 0,
        "images": 0,
        "errors": [],
        "output_filename": Path(dst).name,
    }

    with zipfile.ZipFile(src, "r") as archive:
        document_xml = archive.read("word/document.xml")
        root = ET.fromstring(document_xml)
        rel_targets = _read_relationship_targets(archive)
        body = root.find("w:body", NS)
        if body is None:
            raise RuntimeError("DOCX body is missing from word/document.xml.")

        rescued = docx.Document()
        if len(rescued.paragraphs) == 1 and not rescued.paragraphs[0].text:
            paragraph_element = rescued.paragraphs[0]._element
            parent = paragraph_element.getparent()
            if parent is not None:
                parent.remove(paragraph_element)

        for child in body:
            name = _local_name(child.tag)
            if name == "p":
                _append_paragraph_from_xml(rescued, child, archive, rel_targets, report)
            elif name == "tbl":
                _append_table_from_xml(rescued, child, report)

        rescued.save(dst)

    return report
