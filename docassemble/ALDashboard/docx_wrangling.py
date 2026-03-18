import copy
import docx
import io
import sys
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor
from urllib.parse import urlsplit

import tiktoken
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from docassemble.ALToolbox.llms import chat_completion

from typing import Any, Dict, List, Tuple, Optional, Union, Sequence

from .labeler_config import get_docx_prompt_profile
from .validate_docx import analyze_docx_template_markup, get_jinja_template_validation

__all__ = [
    "apply_docx_label_renames",
    "aggregate_docx_label_suggestion_runs",
    "defragment_docx_runs",
    "get_labeled_docx_runs",
    "get_docx_run_text",
    "get_docx_run_items",
    "get_voted_docx_label_suggestions",
    "review_flagged_docx_label_suggestions",
    "update_docx",
    "validate_docx_label_suggestions",
    "validate_docx_template_syntax",
    "modify_docx_with_openai_guesses",
]


DEFAULT_DOCX_PROMPT_PROFILE = "standard"
DEFAULT_TEMPLATE_HIGHLIGHT_FILL = "C7E1DD"
NESTED_IF_TEMPLATE_HIGHLIGHT_FILLS = [
    DEFAULT_TEMPLATE_HIGHLIGHT_FILL,
    "B8D7D3",
    "A8CDCB",
    "D7E7E4",
    "C5DEDA",
]
_JINJA_TAG_PATTERN = re.compile(r"(\{\{[\s\S]*?\}\}|\{%[\s\S]*?%\})")
_IF_OPEN_PATTERN = re.compile(r"^(?:p\s+)?if\b", re.IGNORECASE)
_IF_BRANCH_PATTERN = re.compile(r"^(?:p\s+)?(?:elif\b|else\b)", re.IGNORECASE)
_IF_CLOSE_PATTERN = re.compile(r"^(?:p\s+)?endif\b", re.IGNORECASE)


def _get_docx_label_role_description(
    *,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
    custom_prompt: Optional[str] = None,
    prompt_library_path: Optional[str] = None,
) -> str:
    """Resolve the DOCX labeler role description for the active prompt profile.

    Args:
        prompt_profile: Prompt profile name to resolve.
        custom_prompt: Optional caller-supplied prompt override.
        prompt_library_path: Optional prompt library override path.

    Returns:
        str: The role description sent to the labeling model.
    """
    if custom_prompt:
        return custom_prompt
    profile_config = get_docx_prompt_profile(
        prompt_profile,
        prompt_library_path=prompt_library_path,
    )
    return str(profile_config.get("role_description") or "")


def _get_docx_label_rules_addendum(
    *,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
    prompt_library_path: Optional[str] = None,
) -> str:
    """Resolve extra prompt rules for the active DOCX prompt profile.

    Args:
        prompt_profile: Prompt profile name to resolve.
        prompt_library_path: Optional prompt library override path.

    Returns:
        str: Additional labeling rules appended to the base prompt.
    """
    profile_config = get_docx_prompt_profile(
        prompt_profile,
        prompt_library_path=prompt_library_path,
    )
    return str(profile_config.get("rules_addendum") or "")


def _get_docx_label_temperature(
    *,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
    prompt_library_path: Optional[str] = None,
) -> float:
    """Resolve the model temperature for the active DOCX prompt profile.

    Args:
        prompt_profile: Prompt profile name to resolve.
        prompt_library_path: Optional prompt library override path.

    Returns:
        float: Temperature value for model calls.
    """
    profile_config = get_docx_prompt_profile(
        prompt_profile,
        prompt_library_path=prompt_library_path,
    )
    try:
        return float(profile_config.get("temperature", 0.5))
    except (TypeError, ValueError):
        return 0.5


def _coerce_modified_run_item(
    item: Any,
) -> Optional[Tuple[int, int, str, int]]:
    """Normalize one model result into (paragraph, run, text, paragraph_delta)."""
    if isinstance(item, dict):
        paragraph_number = item.get("paragraph")
        run_number = item.get("run")
        modified_text = item.get("text")
        new_paragraph = item.get("new_paragraph", 0)
    elif isinstance(item, (list, tuple)) and len(item) >= 4:
        paragraph_number, run_number, modified_text, new_paragraph = item[:4]
    else:
        return None

    if paragraph_number is None:
        return None
    try:
        paragraph_number = int(paragraph_number)
    except (TypeError, ValueError):
        return None
    if run_number is None:
        return None
    try:
        run_number = int(run_number)
    except (TypeError, ValueError):
        # Some models emit [paragraph, original_text, replacement_text, ...].
        if (
            isinstance(item, (list, tuple))
            and len(item) >= 3
            and isinstance(item[1], str)
            and item[2] is not None
        ):
            run_number = 0
            modified_text = item[2]
            new_paragraph = 0
        else:
            return None

    if paragraph_number < 0:
        return None
    if run_number < 0:
        run_number = 0

    if isinstance(new_paragraph, bool):
        new_paragraph = 0
    else:
        try:
            new_paragraph = int(new_paragraph)
        except (TypeError, ValueError):
            new_paragraph = 0
    if new_paragraph not in (-1, 0, 1):
        new_paragraph = 0

    if modified_text is None:
        return None

    return (
        paragraph_number,
        run_number,
        _normalize_paragraph_insert_tag(str(modified_text), new_paragraph),
        new_paragraph,
    )


def _normalize_modified_runs(
    modified_runs: Sequence[Any],
) -> List[Tuple[int, int, str, int]]:
    """Normalize a heterogeneous suggestion list into docx update tuples.

    Args:
        modified_runs: Raw model or API suggestion items.

    Returns:
        List[Tuple[int, int, str, int]]: Valid normalized run modifications.
    """
    normalized: List[Tuple[int, int, str, int]] = []
    for item in modified_runs:
        coerced = _coerce_modified_run_item(item)
        if coerced is not None:
            normalized.append(coerced)
    return normalized


def _normalize_paragraph_insert_tag(text: str, new_paragraph: int) -> str:
    """Convert paragraph-level control tags into Docassemble {%p ... %} tags."""
    if new_paragraph == 0:
        return text

    stripped = text.strip()
    replacements = (
        ("{% if ", "{%p if "),
        ("{% elif ", "{%p elif "),
        ("{% else %}", "{%p else %}"),
        ("{% endif %}", "{%p endif %}"),
        ("{% for ", "{%p for "),
        ("{% endfor %}", "{%p endfor %}"),
    )
    for old, new in replacements:
        if stripped.startswith(old):
            return text.replace(old, new, 1)
    return text


def _normalize_openai_base_url(openai_base_url: Optional[str]) -> Optional[str]:
    """Ensure Azure resource URLs become SDK-compatible `/openai/v1/` bases."""
    base_url = str(openai_base_url or "").strip()
    if not base_url:
        return None

    parsed = urlsplit(base_url)
    hostname = (parsed.hostname or "").lower()
    path = parsed.path.rstrip("/")
    is_azure_host = hostname.endswith(".openai.azure.com") or hostname.endswith(
        ".cognitiveservices.azure.com"
    )

    if is_azure_host and "/openai/deployments/" in path:
        rebuilt = parsed._replace(path="/openai/v1/", query="", fragment="")
        return rebuilt.geturl()

    if is_azure_host and not path.startswith("/openai/"):
        rebuilt = parsed._replace(path="/openai/v1/", query="", fragment="")
        return rebuilt.geturl()

    if is_azure_host and path == "/openai/v1":
        rebuilt = parsed._replace(path="/openai/v1/")
        return rebuilt.geturl()

    return base_url


def _extract_model_results(response: Any) -> List[Any]:
    """Extract a best-effort list of run updates from varied model JSON shapes."""
    if isinstance(response, list):
        return response
    if not isinstance(response, dict):
        return []

    results = response.get("results")
    if isinstance(results, list):
        return results

    for alt_key in ("suggestions", "items", "changes", "labels"):
        alt = response.get(alt_key)
        if isinstance(alt, list):
            return alt

    # Some lightweight models return {"p,r": "replacement text"} maps.
    mapped_results: List[Any] = []
    for key, value in response.items():
        if not isinstance(key, str):
            continue
        match = re.match(r"^\s*(\d+)\s*,\s*(\d+)\s*$", key)
        if not match:
            continue
        if value is None:
            continue
        paragraph_number = int(match.group(1))
        run_number = int(match.group(2))
        if isinstance(value, dict):
            text_value = value.get("text")
            new_paragraph = value.get("new_paragraph", 0)
        else:
            text_value = value
            new_paragraph = 0
        mapped_results.append(
            [paragraph_number, run_number, str(text_value), new_paragraph]
        )
    return mapped_results


def _append_text_content(run_element: Any, text: str) -> None:
    """Append text to a w:r element, preserving tabs/newlines in WordprocessingML."""
    parts = re.split(r"(\t|\n)", text)
    for part in parts:
        if part == "\t":
            run_element.append(OxmlElement("w:tab"))
            continue
        if part == "\n":
            run_element.append(OxmlElement("w:br"))
            continue
        if not part:
            continue

        text_element = OxmlElement("w:t")
        # Preserve leading/trailing spaces exactly when present.
        if part[:1].isspace() or part[-1:].isspace():
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = part
        run_element.append(text_element)


def _copy_run_properties(source_run: Any) -> Optional[Any]:
    """Clone run properties while clearing existing highlight/shading markup."""
    source_properties = source_run._element.rPr
    if source_properties is None:
        return None

    copied_properties = copy.deepcopy(source_properties)
    for child in list(copied_properties):
        if child.tag in {qn("w:highlight"), qn("w:shd")}:
            copied_properties.remove(child)
    return copied_properties


def _build_run_element(
    source_run: Any,
    text: str,
    *,
    shading_fill: Optional[str] = None,
) -> Any:
    """Create a ``w:r`` element mirroring a source run's formatting."""
    run_element = OxmlElement("w:r")
    copied_properties = _copy_run_properties(source_run)

    if copied_properties is not None:
        if shading_fill:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), shading_fill)
            copied_properties.append(shading)
        run_element.append(copied_properties)
    elif shading_fill:
        properties = OxmlElement("w:rPr")
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), shading_fill)
        properties.append(shading)
        run_element.append(properties)

    _append_text_content(run_element, text)
    return run_element


def _if_highlight_fill_for_depth(depth: int) -> str:
    """Return a low-saturation teal-family fill for the nesting depth."""
    return NESTED_IF_TEMPLATE_HIGHLIGHT_FILLS[
        depth % len(NESTED_IF_TEMPLATE_HIGHLIGHT_FILLS)
    ]


def _template_tag_highlight_fill(tag_text: str, if_stack: List[str]) -> str:
    """Choose the highlight color for one Jinja tag and update nesting state."""
    if tag_text.startswith("{{"):
        return DEFAULT_TEMPLATE_HIGHLIGHT_FILL

    control_body = tag_text[2:-2].strip()
    if _IF_OPEN_PATTERN.match(control_body):
        fill = _if_highlight_fill_for_depth(len(if_stack))
        if_stack.append(fill)
        return fill
    if _IF_BRANCH_PATTERN.match(control_body):
        return if_stack[-1] if if_stack else DEFAULT_TEMPLATE_HIGHLIGHT_FILL
    if _IF_CLOSE_PATTERN.match(control_body):
        fill = if_stack[-1] if if_stack else DEFAULT_TEMPLATE_HIGHLIGHT_FILL
        if if_stack:
            if_stack.pop()
        return fill
    return DEFAULT_TEMPLATE_HIGHLIGHT_FILL


def _split_template_run_segments(
    text: str, if_stack: List[str]
) -> List[Tuple[str, Optional[str]]]:
    """Split one run into literal text and separately highlightable Jinja segments."""
    segments: List[Tuple[str, Optional[str]]] = []
    last_index = 0

    for match in _JINJA_TAG_PATTERN.finditer(text):
        if match.start() > last_index:
            literal = text[last_index : match.start()]
            if literal:
                segments.append((literal, None))

        tag_text = match.group(0)
        opening = tag_text[:2]
        closing = tag_text[-2:]
        body = tag_text[2:-2]
        fill = _template_tag_highlight_fill(tag_text, if_stack)

        segments.append((opening, None))
        if body:
            segments.append((body, fill))
        segments.append((closing, None))
        last_index = match.end()

    if last_index < len(text):
        trailing = text[last_index:]
        if trailing:
            segments.append((trailing, None))

    return segments or [(text, None)]


def _replace_run_with_segments(
    run: Any, segments: Sequence[Tuple[str, Optional[str]]]
) -> None:
    """Replace one run with multiple runs that preserve formatting and split tags."""
    parent = run._element.getparent()
    if parent is None:
        return

    non_empty_segments = [segment for segment in segments if segment[0]]
    if not non_empty_segments:
        parent.remove(run._element)
        return

    for segment_text, shading_fill in non_empty_segments:
        parent.insert(
            parent.index(run._element),
            _build_run_element(run, segment_text, shading_fill=shading_fill),
        )

    parent.remove(run._element)


def apply_jinja2_highlights(
    document: Union[docx.document.Document, str],
) -> docx.document.Document:
    """Split Jinja tags into separate runs and shade the tag body text.

    The inner Jinja code is placed in its own run between the opening and closing
    delimiters so python-docx-template can still consume the rendered text safely.
    Nested ``if`` and ``%p if`` control blocks receive depth-based colors.
    """
    if isinstance(document, str):
        document = docx.Document(document)

    initial_paragraphs = _collect_target_paragraphs(document)
    target_paragraph_numbers = [
        paragraph_number
        for paragraph_number, paragraph in enumerate(initial_paragraphs)
        if _contains_template_markup(paragraph.text)
    ]
    if not target_paragraph_numbers:
        return document

    document, _ = defragment_docx_runs(
        document, paragraph_numbers=target_paragraph_numbers
    )
    paragraphs = _collect_target_paragraphs(document)
    target_numbers_set = set(target_paragraph_numbers)
    if_stack: List[str] = []

    for paragraph_number, paragraph in enumerate(paragraphs):
        if paragraph_number not in target_numbers_set:
            continue
        for run in list(paragraph.runs):
            if run._element.getparent() is None:
                continue
            if not _contains_template_markup(run.text):
                continue
            segments = _split_template_run_segments(run.text, if_stack)
            if len(segments) == 1 and segments[0] == (run.text, None):
                continue
            _replace_run_with_segments(run, segments)

    return document


_SAFE_RUN_CHILD_TAGS = {
    qn("w:rPr"),
    qn("w:t"),
    qn("w:tab"),
    qn("w:br"),
    qn("w:cr"),
    qn("w:noBreakHyphen"),
    qn("w:softHyphen"),
}


def _run_is_safe_to_defragment(run: Any) -> bool:
    """Only merge plain-text runs that contain no fields, drawings, or other complex XML."""
    for child in run._element:
        if child.tag not in _SAFE_RUN_CHILD_TAGS:
            return False
    return True


def _collect_paragraphs_from_table(
    table: Any, collected: List[Any], seen_elements: set
) -> None:
    for row in table.rows:
        for cell in row.cells:
            _collect_paragraphs_from_container(cell, collected, seen_elements)


def _collect_paragraphs_from_container(
    container: Any, collected: List[Any], seen_elements: set
) -> None:
    """Collect unique paragraphs from a container and any nested tables.

    Args:
        container: Document, cell, header, or footer-like object to inspect.
        collected: Output list receiving paragraph objects.
        seen_elements: Set of paragraph element IDs already collected.
    """
    for paragraph in getattr(container, "paragraphs", []):
        paragraph_element_id = id(paragraph._element)
        if paragraph_element_id not in seen_elements:
            seen_elements.add(paragraph_element_id)
            collected.append(paragraph)

    for table in getattr(container, "tables", []):
        _collect_paragraphs_from_table(table, collected, seen_elements)


def _collect_target_paragraphs(document: Any) -> List[Any]:
    """Collect paragraphs from body, tables, headers, and footers."""
    collected: List[Any] = []
    seen_elements: set = set()

    _collect_paragraphs_from_container(document, collected, seen_elements)

    for section in document.sections:
        section_parts = [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]
        for part in section_parts:
            _collect_paragraphs_from_container(part, collected, seen_elements)

    return collected


def defragment_docx_runs(
    document: Union[docx.document.Document, str],
    paragraph_numbers: Optional[Sequence[int]] = None,
) -> Tuple[docx.document.Document, dict]:
    """Merge text-only runs within target paragraphs.

    Args:
        document: A loaded ``python-docx`` document or a path to a DOCX file.
        paragraph_numbers: Optional paragraph indexes to limit defragmentation.

    Returns:
        Tuple[docx.document.Document, dict]: The updated document and summary stats.
    """
    if isinstance(document, str):
        document = docx.Document(document)

    target_paragraphs = (
        {int(paragraph_number) for paragraph_number in paragraph_numbers}
        if paragraph_numbers is not None
        else None
    )

    stats = {
        "paragraphs_defragmented": 0,
        "runs_removed": 0,
    }

    for paragraph_number, paragraph in enumerate(_collect_target_paragraphs(document)):
        if target_paragraphs is not None and paragraph_number not in target_paragraphs:
            continue

        runs = list(paragraph.runs)
        if len(runs) <= 1:
            continue
        if not all(_run_is_safe_to_defragment(run) for run in runs):
            continue

        combined_text = "".join(run.text for run in runs)
        runs[0].text = combined_text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)

        stats["paragraphs_defragmented"] += 1
        stats["runs_removed"] += len(runs) - 1

    return document, stats


def _clone_document(document: docx.document.Document) -> docx.document.Document:
    """Create an in-memory clone of a ``python-docx`` document.

    Args:
        document: Source document to copy.

    Returns:
        docx.document.Document: A detached copy of the input document.
    """
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return docx.Document(buffer)


def _contains_template_markup(text: str) -> bool:
    """Check whether text contains Jinja-style template delimiters.

    Args:
        text: Text to inspect.

    Returns:
        bool: ``True`` when template delimiters are present.
    """
    return "{{" in text or "{%" in text


def _has_balanced_template_delimiters(text: str) -> bool:
    """Check whether Jinja delimiter pairs are balanced in a text snippet.

    Args:
        text: Text to inspect.

    Returns:
        bool: ``True`` when opening and closing delimiters are balanced.
    """
    return text.count("{{") == text.count("}}") and text.count("{%") == text.count("%}")


def _normalize_docassemble_template_source(text: str) -> str:
    """Convert docassemble paragraph tags into plain Jinja tags for AST parsing."""
    return re.sub(r"\{%\s*p\s+", "{% ", text)


def _classify_if_control_tag(raw_tag: str) -> Optional[Dict[str, Any]]:
    if not raw_tag.startswith("{%"):
        return None

    statement = raw_tag[2:-2].strip()
    is_paragraph = False
    if statement.lower().startswith("p "):
        is_paragraph = True
        statement = statement[1:].strip()

    lowered = statement.lower()
    if lowered.startswith("if ") or lowered == "if":
        kind = "open"
    elif lowered.startswith("elif ") or lowered == "else":
        kind = "branch"
    elif lowered == "endif":
        kind = "close"
    else:
        return None

    return {
        "kind": kind,
        "is_paragraph": is_paragraph,
        "raw_tag": raw_tag,
    }


def _append_validation_flag(
    flags_by_index: Dict[int, List[Dict[str, str]]],
    index: int,
    code: str,
    message: str,
) -> None:
    existing = flags_by_index.setdefault(index, [])
    if any(flag.get("code") == code and flag.get("message") == message for flag in existing):
        return
    existing.append({"code": code, "message": message})


def _validate_if_control_tag_parity(
    suggestions: Sequence[Tuple[int, int, str, int]],
) -> Dict[int, List[Dict[str, str]]]:
    flags_by_index: Dict[int, List[Dict[str, str]]] = {}
    ordered = sorted(
        enumerate(suggestions),
        key=lambda item: (
            item[1][0],
            {-1: 0, 0: 1, 1: 2}.get(item[1][3], 1),
            item[1][1],
            item[0],
        ),
    )
    stack: List[Dict[str, Any]] = []

    for index, (_paragraph_number, _run_number, text, _new_paragraph) in ordered:
        for raw_tag in _JINJA_TAG_PATTERN.findall(text):
            control = _classify_if_control_tag(raw_tag)
            if control is None:
                continue

            if control["kind"] == "open":
                stack.append({"index": index, **control})
                continue

            if control["kind"] == "branch":
                if not stack:
                    _append_validation_flag(
                        flags_by_index,
                        index,
                        "unmatched_conditional_branch",
                        "Conditional branch tag must follow a matching if tag.",
                    )
                    continue
                opener = stack[-1]
                if opener["is_paragraph"] != control["is_paragraph"]:
                    _append_validation_flag(
                        flags_by_index,
                        index,
                        "mismatched_conditional_control_tag_style",
                        "Paragraph conditional tags must keep using %p for elif/else branches.",
                    )
                    _append_validation_flag(
                        flags_by_index,
                        opener["index"],
                        "mismatched_conditional_control_tag_style",
                        "Paragraph conditional tags must keep using %p for elif/else branches.",
                    )
                continue

            if not stack:
                _append_validation_flag(
                    flags_by_index,
                    index,
                    "unmatched_endif_control_tag",
                    "Conditional tags must be balanced: found endif without a matching if.",
                )
                continue

            opener = stack.pop()
            if opener["is_paragraph"] != control["is_paragraph"]:
                message = (
                    "Paragraph conditional tags must close with %p endif, not plain endif."
                    if opener["is_paragraph"]
                    else "Plain if tags must close with plain endif, not %p endif."
                )
                _append_validation_flag(
                    flags_by_index,
                    opener["index"],
                    "mismatched_conditional_control_tag_style",
                    message,
                )
                _append_validation_flag(
                    flags_by_index,
                    index,
                    "mismatched_conditional_control_tag_style",
                    message,
                )

    for opener in stack:
        _append_validation_flag(
            flags_by_index,
            opener["index"],
            "unmatched_if_control_tag",
            "Conditional tags must be balanced: found if without a matching endif.",
        )

    return flags_by_index


def _should_validate_snippet_as_template(text: str) -> bool:
    controls = []
    for raw_tag in _JINJA_TAG_PATTERN.findall(text):
        control = _classify_if_control_tag(raw_tag)
        if control is not None:
            controls.append(control)

    if not controls:
        return True

    if any(control["kind"] == "branch" for control in controls):
        return False

    opens = sum(1 for control in controls if control["kind"] == "open")
    closes = sum(1 for control in controls if control["kind"] == "close")
    return opens == closes


def apply_docx_label_renames(
    document: docx.document.Document,
    renames: Sequence[Dict[str, Any]],
) -> int:
    """Apply label find/replace operations across all paragraphs in the DOCX."""
    rename_count = 0
    for rename in renames:
        original = str(rename.get("original", ""))
        replacement = str(rename.get("replacement", ""))
        if not original or not replacement or original == replacement:
            continue

        for paragraph in _collect_target_paragraphs(document):
            for run in paragraph.runs:
                if original in run.text:
                    run.text = run.text.replace(original, replacement)
                    rename_count += 1

    return rename_count


def validate_docx_template_syntax(
    document: Union[docx.document.Document, str],
    *,
    suggestions: Sequence[Any] = (),
    renames: Sequence[Dict[str, Any]] = (),
    defragment_runs: bool = False,
) -> Dict[str, Any]:
    """Validate the Jinja syntax of a DOCX after simulated edits are applied."""
    if isinstance(document, str):
        document = docx.Document(document)

    working_document = _clone_document(document)
    normalized_suggestions = _normalize_modified_runs(suggestions)

    if renames:
        apply_docx_label_renames(working_document, renames)

    if normalized_suggestions:
        working_document = update_docx(
            working_document,
            normalized_suggestions,
            defragment_runs=defragment_runs,
        )

    paragraphs = _collect_target_paragraphs(working_document)
    template_source = "\n".join(
        _normalize_docassemble_template_source(paragraph.text)
        for paragraph in paragraphs
    )
    validation = get_jinja_template_validation(template_source)
    validation["warnings"].extend(analyze_docx_template_markup(working_document))

    for issue_group in (validation["errors"], validation["warnings"]):
        for issue in issue_group:
            line_number = issue.get("line")
            if isinstance(line_number, int) and 1 <= line_number <= len(paragraphs):
                issue["paragraph"] = line_number - 1
                issue["paragraph_text"] = paragraphs[line_number - 1].text

    validation["error_count"] = len(validation["errors"])
    validation["warning_count"] = len(validation["warnings"])
    return validation


def _has_placeholder_markers(text: str) -> bool:
    """Detect obvious placeholder markers such as long underscores or tab runs.

    Args:
        text: Text to inspect.

    Returns:
        bool: ``True`` when placeholder markers are present.
    """
    return bool(re.search(r"_{3,}", text) or re.search(r"\t{2,}", text))


def _has_adjacent_word_fragments(text: str) -> bool:
    """Detect template markup that is glued directly to surrounding letters.

    Args:
        text: Text to inspect.

    Returns:
        bool: ``True`` when markup appears attached to word fragments.
    """
    return bool(
        re.search(r"(\}\}|\%\})[A-Za-z]", text)
        or re.search(r"[A-Za-z](\{\{|\{%)", text)
    )


def _looks_inline_placeholder(source_run_text: str, source_paragraph_text: str) -> bool:
    """Decide whether the source text looks like an inline placeholder target.

    Args:
        source_run_text: Original text from the targeted run.
        source_paragraph_text: Original paragraph text containing the run.

    Returns:
        bool: ``True`` when the text should likely be replaced in place.
    """
    stripped_run = source_run_text.strip()
    if not stripped_run:
        return True
    if _has_placeholder_markers(source_run_text):
        return True
    return _has_placeholder_markers(source_paragraph_text)


def _filter_noop_suggestions(
    document: docx.document.Document,
    suggestions: Sequence[Tuple[int, int, str, int]],
) -> List[Tuple[int, int, str, int]]:
    """Drop suggestions that leave the targeted run or paragraph unchanged."""
    paragraphs = _collect_target_paragraphs(document)
    filtered: List[Tuple[int, int, str, int]] = []

    for paragraph_number, run_number, text, new_paragraph in suggestions:
        if paragraph_number < 0 or paragraph_number >= len(paragraphs):
            filtered.append((paragraph_number, run_number, text, new_paragraph))
            continue

        paragraph = paragraphs[paragraph_number]
        source_text = (
            paragraph.runs[run_number].text
            if 0 <= run_number < len(paragraph.runs)
            else paragraph.text
        )
        if new_paragraph == 0 and text == source_text:
            continue
        filtered.append((paragraph_number, run_number, text, new_paragraph))

    return filtered


def _run_has_fragmented_word_boundary(paragraph: Any, run_number: int) -> bool:
    """Check whether a run is embedded inside a split alphabetical word.

    Args:
        paragraph: Paragraph containing the run.
        run_number: Index of the run to inspect.

    Returns:
        bool: ``True`` when neighboring runs split a word around the target run.
    """
    if run_number < 0 or run_number >= len(paragraph.runs):
        return False

    run_text = paragraph.runs[run_number].text
    if not run_text:
        return False

    left_fragment = (
        run_number > 0
        and paragraph.runs[run_number - 1].text[-1:].isalpha()
        and run_text[:1].isalpha()
    )
    right_fragment = (
        run_number + 1 < len(paragraph.runs)
        and run_text[-1:].isalpha()
        and paragraph.runs[run_number + 1].text[:1].isalpha()
    )
    return left_fragment or right_fragment


def validate_docx_label_suggestions(
    document: Union[docx.document.Document, str],
    suggestions: Sequence[Any],
) -> Dict[str, Any]:
    """Run deterministic checks over model suggestions and the simulated output.

    Args:
        document: A loaded ``python-docx`` document or a path to a DOCX file.
        suggestions: Raw suggestion items returned by labeling helpers or models.

    Returns:
        Dict[str, Any]: Per-suggestion validation results and aggregate counts.
    """
    if isinstance(document, str):
        document = docx.Document(document)

    normalized = _normalize_modified_runs(suggestions)
    paragraphs = _collect_target_paragraphs(document)
    simulated_document = update_docx(_clone_document(document), normalized)
    simulated_paragraphs = _collect_target_paragraphs(simulated_document)
    parity_flags = _validate_if_control_tag_parity(normalized)
    syntax_validation = validate_docx_template_syntax(simulated_document)

    results: List[Dict[str, Any]] = []
    flagged_count = 0

    for index, (paragraph_number, run_number, text, new_paragraph) in enumerate(
        normalized
    ):
        paragraph = (
            paragraphs[paragraph_number]
            if 0 <= paragraph_number < len(paragraphs)
            else None
        )
        source_paragraph_text = paragraph.text if paragraph is not None else ""
        source_run_text = (
            paragraph.runs[run_number].text
            if paragraph is not None and 0 <= run_number < len(paragraph.runs)
            else source_paragraph_text
        )
        simulated_paragraph_text = (
            simulated_paragraphs[paragraph_number].text
            if new_paragraph == 0 and 0 <= paragraph_number < len(simulated_paragraphs)
            else ""
        )

        flags: List[Dict[str, str]] = []
        flags.extend(parity_flags.get(index, []))

        if not _contains_template_markup(text):
            flags.append(
                {
                    "code": "missing_template_markup",
                    "message": "Suggestion does not contain Jinja2 markup.",
                }
            )

        if not _has_balanced_template_delimiters(text):
            flags.append(
                {
                    "code": "unbalanced_template_delimiters",
                    "message": "Suggestion has unbalanced Jinja2 delimiters.",
                }
            )

        if new_paragraph != 0 and not text.strip().startswith("{%p "):
            flags.append(
                {
                    "code": "paragraph_insert_without_control_tag",
                    "message": "Suggestion inserts a new paragraph without a Docassemble paragraph control tag.",
                }
            )

        if new_paragraph != 0 and _looks_inline_placeholder(
            source_run_text, source_paragraph_text
        ):
            flags.append(
                {
                    "code": "inline_placeholder_emitted_as_paragraph",
                    "message": "Inline placeholder was modeled as a paragraph insertion instead of an in-place replacement.",
                }
            )

        if paragraph is not None and _run_has_fragmented_word_boundary(
            paragraph, run_number
        ):
            flags.append(
                {
                    "code": "fragmented_word_boundary",
                    "message": "Target run sits inside a split word boundary; replacement may leave word fragments behind without run consolidation.",
                }
            )

        if (
            new_paragraph == 0
            and simulated_paragraph_text
            and _contains_template_markup(text)
        ):
            if _has_placeholder_markers(simulated_paragraph_text):
                flags.append(
                    {
                        "code": "leftover_placeholder_markers",
                        "message": "Simulated output still contains tabs or underscore placeholders near the inserted label.",
                    }
                )
            if _has_adjacent_word_fragments(simulated_paragraph_text):
                flags.append(
                    {
                        "code": "leftover_word_fragments",
                        "message": "Simulated output leaves word fragments attached to template markup.",
                    }
                )

        if _should_validate_snippet_as_template(text):
            snippet_validation = get_jinja_template_validation(
                _normalize_docassemble_template_source(text)
            )
            for issue in snippet_validation["errors"]:
                flags.append(
                    {
                        "code": "invalid_jinja_syntax",
                        "message": issue.get("message")
                        or "Suggestion does not parse as valid Jinja syntax.",
                    }
                )

        if flags:
            flagged_count += 1

        results.append(
            {
                "index": index,
                "paragraph": paragraph_number,
                "run": run_number,
                "new_paragraph": new_paragraph,
                "source_paragraph_text": source_paragraph_text,
                "source_run_text": source_run_text,
                "suggested_text": text,
                "simulated_paragraph_text": simulated_paragraph_text,
                "flags": flags,
            }
        )

    return {
        "results": results,
        "flagged_count": flagged_count,
        "ai_review_recommended": flagged_count > 0,
        "syntax_validation": syntax_validation,
    }


def review_flagged_docx_label_suggestions(
    document: Union[docx.document.Document, str],
    suggestions: Sequence[Any],
    deterministic_validation: Dict[str, Any],
    *,
    openai_client: Optional[Any] = None,
    openai_api: Optional[str] = None,
    openai_base_url: Optional[str] = None,
    model: str = "gpt-5-mini",
    max_output_tokens: Optional[int] = None,
) -> Dict[str, Any]:
    """Ask an LLM to review only deterministic-validator flagged suggestions.

    Args:
        document: A loaded ``python-docx`` document or a path to a DOCX file.
        suggestions: Raw suggestion items returned by a model or API.
        deterministic_validation: Validator output describing flagged suggestions.
        openai_client: Optional initialized OpenAI client.
        openai_api: Optional API key override.
        openai_base_url: Optional OpenAI-compatible base URL override.
        model: Model name to use for the review step.
        max_output_tokens: Optional token limit for the review call.

    Returns:
        Dict[str, Any]: Review status, any error text, and normalized review items.
    """
    openai_base_url = _normalize_openai_base_url(openai_base_url)

    if isinstance(document, str):
        document = docx.Document(document)

    flagged_items = [
        item
        for item in deterministic_validation.get("results", [])
        if item.get("flags")
    ]
    if not flagged_items:
        return {"performed": False, "reviews": []}

    review_payload = []
    for item in flagged_items:
        review_payload.append(
            {
                "index": item["index"],
                "paragraph": item["paragraph"],
                "run": item["run"],
                "new_paragraph": item["new_paragraph"],
                "source_paragraph_text": item["source_paragraph_text"],
                "source_run_text": item["source_run_text"],
                "suggested_text": item["suggested_text"],
                "simulated_paragraph_text": item["simulated_paragraph_text"],
                "deterministic_flags": item["flags"],
            }
        )

    messages = [
        {
            "role": "system",
            "content": (
                "Review flagged DOCX labeling suggestions. "
                "Return JSON with a 'reviews' array. For each flagged item, include: "
                "'index', 'verdict' (one of 'ok', 'revise', 'reject'), and 'reason'. "
                "Use 'ok' only when the deterministic flag looks like a false positive."
            ),
        },
        {"role": "user", "content": json.dumps(review_payload, ensure_ascii=False)},
    ]

    try:
        response = chat_completion(
            model=model,
            messages=messages,
            json_mode=True,
            temperature=0,
            max_output_tokens=max_output_tokens,
            openai_client=openai_client,
            openai_api=openai_api,
            openai_base_url=openai_base_url,
        )
    except Exception as exc:
        return {
            "performed": False,
            "error": str(exc),
            "reviews": [],
        }

    reviews = response.get("reviews", []) if isinstance(response, dict) else []
    normalized_reviews: List[Dict[str, Any]] = []
    for review in reviews:
        if not isinstance(review, dict):
            continue
        try:
            review_index = int(review.get("index"))  # type: ignore[arg-type]
        except (TypeError, ValueError):
            continue
        verdict = str(review.get("verdict") or "revise").strip().lower()
        if verdict not in {"ok", "revise", "reject"}:
            verdict = "revise"
        normalized_reviews.append(
            {
                "index": review_index,
                "verdict": verdict,
                "reason": str(review.get("reason") or "").strip(),
            }
        )

    return {"performed": True, "reviews": normalized_reviews}


def _suggestion_confidence_tier(clean_vote_count: int, total_generations: int) -> str:
    """Classify suggestion confidence from clean votes across generations.

    Args:
        clean_vote_count: Number of generations that agreed without validation flags.
        total_generations: Total number of candidate generations considered.

    Returns:
        str: ``high``, ``medium``, or ``low`` confidence.
    """
    if clean_vote_count >= min(3, total_generations):
        return "high"
    if clean_vote_count >= 2:
        return "medium"
    return "low"


def _litigation_template_paragraph_likely_templated(text: str) -> bool:
    """Heuristically detect pleading text that likely needs templating.

    Args:
        text: Paragraph text to inspect.

    Returns:
        bool: ``True`` when the paragraph looks like drafting or placeholder text.
    """
    paragraph_text = str(text or "")
    stripped = paragraph_text.strip()
    lowered = stripped.lower()
    if not stripped:
        return False
    if re.search(r"_{3,}", paragraph_text):
        return True
    if re.search(r"\[[^\]]+\]", paragraph_text):
        return True
    if "case no." in lowered:
        return True
    if stripped.startswith("#.") and any(
        phrase in lowered
        for phrase in (
            "short introduction",
            "if helpful",
            "if applicable",
            "consider adding",
            "add numbered paragraphs",
            "add the legal background",
            "allege constitutional violation",
            "allege statutory",
        )
    ):
        return True
    return False


def _litigation_alignment_heavy_line(text: str) -> bool:
    """Detect caption-style lines where tabs and parentheticals drive layout.

    Args:
        text: Paragraph text to inspect.

    Returns:
        bool: ``True`` when the line is alignment-heavy.
    """
    line_text = str(text or "")
    return "\t" in line_text or line_text.count(")") >= 2


def _candidate_flag_codes(candidate: Dict[str, Any]) -> set[str]:
    """Extract validation flag codes from an aggregated candidate record.

    Args:
        candidate: Candidate suggestion metadata.

    Returns:
        set[str]: Unique flag codes present on the candidate.
    """
    return {
        str(flag.get("code") or "").strip()
        for flag in candidate.get("validation_flags", [])
        if isinstance(flag, dict)
    }


def _effective_candidate_flags(
    candidate: Dict[str, Any],
    group: Dict[str, Any],
    *,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
) -> List[Dict[str, Any]]:
    """Adjust candidate validation flags for prompt-profile-specific heuristics.

    Args:
        candidate: Candidate suggestion metadata.
        group: Aggregated source-group metadata.
        prompt_profile: Prompt profile name driving validation behavior.

    Returns:
        List[Dict[str, Any]]: The effective validation flags to use for ranking.
    """
    flags = list(candidate.get("validation_flags") or [])
    normalized_profile = (
        str(prompt_profile or DEFAULT_DOCX_PROMPT_PROFILE).strip().lower()
    )
    if normalized_profile != "litigation_template":
        return flags

    flag_codes = _candidate_flag_codes(candidate)
    source_paragraph_text = str(group.get("source_paragraph_text") or "")
    if flag_codes == {
        "leftover_placeholder_markers"
    } and _litigation_alignment_heavy_line(source_paragraph_text):
        return []
    return flags


def _effective_clean_vote_count(
    candidate: Dict[str, Any],
    group: Dict[str, Any],
    *,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
) -> int:
    """Compute the effective clean-vote count after profile-specific adjustments.

    Args:
        candidate: Candidate suggestion metadata.
        group: Aggregated source-group metadata.
        prompt_profile: Prompt profile name driving validation behavior.

    Returns:
        int: Effective count of clean votes for candidate ranking.
    """
    if not _effective_candidate_flags(candidate, group, prompt_profile=prompt_profile):
        raw_clean_count = int(candidate.get("clean_vote_count", 0))
        if raw_clean_count > 0:
            return raw_clean_count
        return int(candidate.get("vote_count", 0))
    return int(candidate.get("clean_vote_count", 0))


def _candidate_priority_key(candidate: Dict[str, Any]) -> Tuple[int, int, int, int]:
    """Build a sort key that favors consensus, cleanliness, and concise labels.

    Args:
        candidate: Candidate suggestion metadata.

    Returns:
        Tuple[int, int, int, int]: Sort key used to rank competing candidates.
    """
    return (
        int(
            candidate.get(
                "effective_clean_vote_count", candidate.get("clean_vote_count", 0)
            )
        ),
        int(candidate.get("vote_count", 0)),
        -len(
            candidate.get(
                "effective_validation_flags", candidate.get("validation_flags", [])
            )
        ),
        -len(str(candidate.get("text") or "")),
    )


def review_docx_label_candidate_groups(
    candidate_groups: Sequence[Dict[str, Any]],
    *,
    openai_client: Optional[Any] = None,
    openai_api: Optional[str] = None,
    openai_base_url: Optional[str] = None,
    model: str = "gpt-5-mini",
    max_output_tokens: Optional[int] = None,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
) -> Dict[str, Any]:
    """Ask an LLM judge to choose the best candidate per ambiguous position.

    Args:
        candidate_groups: Candidate groups requiring judge review.
        openai_client: Optional initialized OpenAI client.
        openai_api: Optional API key override.
        openai_base_url: Optional OpenAI-compatible base URL override.
        model: Model name to use for judge review.
        max_output_tokens: Optional token limit for the judge call.
        prompt_profile: Prompt profile name used for litigation-specific heuristics.

    Returns:
        Dict[str, Any]: Judge review status and normalized review items.
    """
    openai_base_url = _normalize_openai_base_url(openai_base_url)
    if not candidate_groups:
        return {"performed": False, "reviews": []}

    review_payload: List[Dict[str, Any]] = []
    for group in candidate_groups:
        normalized_profile = (
            str(prompt_profile or DEFAULT_DOCX_PROMPT_PROFILE).strip().lower()
        )
        instruction = (
            "Choose the best candidate when it safely improves the DOCX template. "
            "Prefer clean candidates with 2+ votes. Single-vote candidates are low confidence. "
            "Reject all candidates if they look unsafe, redundant, or malformed."
        )
        if normalized_profile == "litigation_template":
            instruction = (
                "Choose the best candidate when it safely improves a litigation-style DOCX template. "
                "Prefer clean candidates with 2+ votes, but a clean single-vote candidate can still be acceptable "
                "when the source paragraph obviously contains blanks, bracketed drafting notes, caption placeholders, "
                "or other authoring artifacts. Do not reject solely because vote count is 1 if the candidate cleanly "
                "replaces an obvious placeholder paragraph."
            )
        review_payload.append(
            {
                "group_index": group["group_index"],
                "paragraph": group["paragraph"],
                "run": group["run"],
                "new_paragraph": group["new_paragraph"],
                "source_paragraph_text": group.get("source_paragraph_text", ""),
                "source_run_text": group.get("source_run_text", ""),
                "instruction": instruction,
                "candidates": [
                    {
                        "candidate_index": candidate["candidate_index"],
                        "text": candidate["text"],
                        "vote_count": candidate["vote_count"],
                        "clean_vote_count": candidate.get(
                            "effective_clean_vote_count",
                            candidate["clean_vote_count"],
                        ),
                        "confidence": candidate["confidence"],
                        "validation_flags": candidate.get(
                            "effective_validation_flags",
                            candidate["validation_flags"],
                        ),
                        "sources": candidate["sources"],
                    }
                    for candidate in group["candidates"]
                ],
            }
        )

    messages = [
        {
            "role": "system",
            "content": (
                "Adjudicate competing DOCX labeling candidates. "
                "Return JSON with a 'reviews' array. For each group include: "
                "'group_index', 'decision' (choose or reject), optional 'candidate_index', and 'reason'."
            ),
        },
        {"role": "user", "content": json.dumps(review_payload, ensure_ascii=False)},
    ]

    try:
        response = chat_completion(
            model=model,
            messages=messages,
            json_mode=True,
            temperature=0,
            max_output_tokens=max_output_tokens,
            openai_client=openai_client,
            openai_api=openai_api,
            openai_base_url=openai_base_url,
        )
    except Exception as exc:
        return {
            "performed": False,
            "error": str(exc),
            "reviews": [],
        }

    reviews = response.get("reviews", []) if isinstance(response, dict) else []
    normalized_reviews: List[Dict[str, Any]] = []
    for review in reviews:
        if not isinstance(review, dict):
            continue
        try:
            group_index = int(review.get("group_index"))  # type: ignore[arg-type]
        except (TypeError, ValueError):
            continue
        decision = str(review.get("decision") or "reject").strip().lower()
        if decision not in {"choose", "reject"}:
            decision = "reject"
        candidate_index = review.get("candidate_index")
        if decision == "choose":
            try:
                candidate_index = int(candidate_index)  # type: ignore[arg-type]
            except (TypeError, ValueError):
                decision = "reject"
                candidate_index = None
        else:
            candidate_index = None
        normalized_reviews.append(
            {
                "group_index": group_index,
                "decision": decision,
                "candidate_index": candidate_index,
                "reason": str(review.get("reason") or "").strip(),
            }
        )

    return {"performed": True, "reviews": normalized_reviews}


def aggregate_docx_label_suggestion_runs(
    document: Union[docx.document.Document, str],
    suggestion_runs: Sequence[Dict[str, Any]],
    *,
    judge_model: Optional[str] = None,
    openai_client: Optional[Any] = None,
    openai_api: Optional[str] = None,
    openai_base_url: Optional[str] = None,
    judge_max_output_tokens: Optional[int] = None,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
) -> Dict[str, Any]:
    """Combine repeated suggestion runs into one ranked set with alternates.

    Args:
        document: A loaded ``python-docx`` document or a path to a DOCX file.
        suggestion_runs: Multiple labeling result sets to combine.
        judge_model: Optional model override for ambiguous-group judging.
        openai_client: Optional initialized OpenAI client for judge review.
        openai_api: Optional API key override.
        openai_base_url: Optional OpenAI-compatible base URL override.
        judge_max_output_tokens: Optional token limit for judge calls.
        prompt_profile: Prompt profile name used for ranking heuristics.

    Returns:
        Dict[str, Any]: Aggregated suggestions, vote data, and judge metadata.
    """
    if isinstance(document, str):
        document = docx.Document(document)

    total_generations = len(suggestion_runs)
    position_groups: Dict[Tuple[int, int, int], Dict[str, Any]] = {}

    for run_info in suggestion_runs:
        model_name = str(run_info.get("model") or "")
        generation_index = int(run_info.get("generation_index") or 0)
        suggestions = _normalize_modified_runs(run_info.get("suggestions") or [])
        validation = run_info.get("validation") or {}
        validation_results = validation.get("results", [])
        for suggestion_index, suggestion in enumerate(suggestions):
            paragraph_number, run_number, text, new_paragraph = suggestion
            validation_item = (
                validation_results[suggestion_index]
                if suggestion_index < len(validation_results)
                else {
                    "flags": [],
                    "source_paragraph_text": "",
                    "source_run_text": "",
                    "simulated_paragraph_text": "",
                }
            )
            position_key = (paragraph_number, run_number, new_paragraph)
            group = position_groups.setdefault(
                position_key,
                {
                    "paragraph": paragraph_number,
                    "run": run_number,
                    "new_paragraph": new_paragraph,
                    "source_paragraph_text": validation_item.get(
                        "source_paragraph_text", ""
                    ),
                    "source_run_text": validation_item.get("source_run_text", ""),
                    "candidates": {},
                },
            )
            candidate = group["candidates"].setdefault(
                text,
                {
                    "text": text,
                    "paragraph": paragraph_number,
                    "run": run_number,
                    "new_paragraph": new_paragraph,
                    "source_paragraph_text": validation_item.get(
                        "source_paragraph_text", ""
                    ),
                    "source_run_text": validation_item.get("source_run_text", ""),
                    "simulated_paragraph_text": validation_item.get(
                        "simulated_paragraph_text", ""
                    ),
                    "validation_flags": validation_item.get("flags", []),
                    "sources": [],
                    "vote_count": 0,
                    "clean_vote_count": 0,
                },
            )
            candidate["sources"].append(
                {
                    "model": model_name,
                    "generation_index": generation_index,
                    "suggestion_index": suggestion_index,
                }
            )
            candidate["vote_count"] += 1
            if not candidate["validation_flags"]:
                candidate["clean_vote_count"] += 1

    ambiguous_groups: List[Dict[str, Any]] = []
    selected_suggestions: List[Dict[str, Any]] = []
    discarded_groups = 0

    for group_index, position_key in enumerate(sorted(position_groups.keys())):
        group = position_groups[position_key]
        candidates = list(group["candidates"].values())
        for candidate_index, candidate in enumerate(candidates):
            effective_flags = _effective_candidate_flags(
                candidate, group, prompt_profile=prompt_profile
            )
            effective_clean_vote_count = _effective_clean_vote_count(
                candidate, group, prompt_profile=prompt_profile
            )
            candidate["effective_validation_flags"] = effective_flags
            candidate["effective_clean_vote_count"] = effective_clean_vote_count
            candidate["candidate_index"] = candidate_index
            candidate["confidence"] = _suggestion_confidence_tier(
                effective_clean_vote_count, total_generations
            )
        candidates.sort(key=_candidate_priority_key, reverse=True)
        for candidate_index, candidate in enumerate(candidates):
            candidate["candidate_index"] = candidate_index

        clean_consensus = [
            candidate
            for candidate in candidates
            if not candidate["effective_validation_flags"]
            and candidate["effective_clean_vote_count"] >= 2
        ]
        chosen_candidate: Optional[Dict[str, Any]] = None
        judge_needed = False

        if len(clean_consensus) == 1:
            chosen_candidate = clean_consensus[0]
        elif len(clean_consensus) > 1:
            judge_needed = True
        else:
            normalized_profile = (
                str(prompt_profile or DEFAULT_DOCX_PROMPT_PROFILE).strip().lower()
            )
            if normalized_profile == "litigation_template":
                clean_singletons = [
                    candidate
                    for candidate in candidates
                    if not candidate["effective_validation_flags"]
                    and candidate["effective_clean_vote_count"] == 1
                ]
                if len(
                    clean_singletons
                ) == 1 and _litigation_template_paragraph_likely_templated(
                    group.get("source_paragraph_text", "")
                ):
                    chosen_candidate = clean_singletons[0]
                else:
                    judge_needed = True
            else:
                judge_needed = True

        if judge_needed:
            ambiguous_groups.append(
                {
                    "group_index": group_index,
                    "paragraph": group["paragraph"],
                    "run": group["run"],
                    "new_paragraph": group["new_paragraph"],
                    "source_paragraph_text": group["source_paragraph_text"],
                    "source_run_text": group["source_run_text"],
                    "candidates": candidates,
                }
            )
        else:
            alternates = [
                candidate
                for candidate in candidates
                if candidate is not chosen_candidate
            ]
            assert chosen_candidate is not None
            selected_suggestions.append(
                {
                    "paragraph": chosen_candidate["paragraph"],
                    "run": chosen_candidate["run"],
                    "text": chosen_candidate["text"],
                    "new_paragraph": chosen_candidate["new_paragraph"],
                    "validation_flags": chosen_candidate["effective_validation_flags"],
                    "judge_review": None,
                    "confidence": chosen_candidate["confidence"],
                    "vote_count": chosen_candidate["vote_count"],
                    "clean_vote_count": chosen_candidate["effective_clean_vote_count"],
                    "vote_total": total_generations,
                    "sources": chosen_candidate["sources"],
                    "alternates": alternates,
                }
            )

    judge_review = review_docx_label_candidate_groups(
        ambiguous_groups,
        openai_client=openai_client,
        openai_api=openai_api,
        openai_base_url=openai_base_url,
        model=judge_model or "gpt-5-mini",
        max_output_tokens=judge_max_output_tokens,
        prompt_profile=prompt_profile,
    )
    judge_reviews_by_index = {
        int(item["group_index"]): item for item in judge_review.get("reviews", [])
    }

    for group in ambiguous_groups:
        review = judge_reviews_by_index.get(group["group_index"], {})
        judge_chosen: Optional[Dict[str, Any]] = None
        if review.get("decision") == "choose":
            candidate_index = review.get("candidate_index")
            for candidate in group["candidates"]:
                if candidate["candidate_index"] == candidate_index:
                    judge_chosen = candidate
                    break
        if judge_chosen is None:
            discarded_groups += 1
            continue
        alternates = [
            candidate
            for candidate in group["candidates"]
            if candidate["candidate_index"] != judge_chosen["candidate_index"]
        ]
        selected_suggestions.append(
            {
                "paragraph": judge_chosen["paragraph"],
                "run": judge_chosen["run"],
                "text": judge_chosen["text"],
                "new_paragraph": judge_chosen["new_paragraph"],
                "validation_flags": judge_chosen["effective_validation_flags"],
                "judge_review": review,
                "confidence": judge_chosen["confidence"],
                "vote_count": judge_chosen["vote_count"],
                "clean_vote_count": judge_chosen["effective_clean_vote_count"],
                "vote_total": total_generations,
                "sources": judge_chosen["sources"],
                "alternates": alternates,
            }
        )

    confidence_counts: Dict[str, int] = {"high": 0, "medium": 0, "low": 0}
    for suggestion in selected_suggestions:  # type: ignore[assignment]
        key = suggestion["confidence"]  # type: ignore[call-overload]
        confidence_counts[key] = confidence_counts.get(key, 0) + 1  # type: ignore[arg-type,index]

    selected_suggestions.sort(
        key=lambda item: (
            {"high": 0, "medium": 1, "low": 2}.get(item["confidence"], 3),
            item["paragraph"],
            item["run"],
            item["new_paragraph"],
        )
    )

    return {
        "suggestions": selected_suggestions,
        "aggregation": {
            "generator_runs": total_generations,
            "generator_models": [
                str(run.get("model") or "") for run in suggestion_runs
            ],
            "judge_model": judge_model or None,
            "total_candidate_groups": len(position_groups),
            "ambiguous_group_count": len(ambiguous_groups),
            "discarded_group_count": discarded_groups,
            "confidence_counts": confidence_counts,
        },
        "judge_review": judge_review,
    }


def get_voted_docx_label_suggestions(
    docx_path: str,
    custom_people_names: Optional[List[Tuple[str, str]]] = None,
    preferred_variable_names: Optional[Sequence[str]] = None,
    openai_client: Optional[Any] = None,
    openai_api: Optional[str] = None,
    openai_base_url: Optional[str] = None,
    model: str = "gpt-5-mini",
    generator_models: Optional[Sequence[str]] = None,
    judge_model: Optional[str] = None,
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
    prompt_library_path: Optional[str] = None,
    optional_context: Optional[str] = None,
    custom_prompt: Optional[str] = None,
    additional_instructions: Optional[str] = None,
    max_output_tokens: Optional[int] = None,
    judge_max_output_tokens: Optional[int] = None,
    defragment_runs: bool = False,
) -> Dict[str, Any]:
    """Run repeated generations and aggregate them into one ranked suggestion set.

    Args:
        docx_path: Path to the DOCX file to label.
        custom_people_names: Optional list of custom people-variable descriptions.
        preferred_variable_names: Optional preferred variable names to bias prompts.
        openai_client: Optional initialized OpenAI client.
        openai_api: Optional API key override.
        openai_base_url: Optional OpenAI-compatible base URL override.
        model: Default generator model to use.
        generator_models: Optional explicit sequence of generator models.
        judge_model: Optional model override for ambiguous-group judging.
        prompt_profile: Prompt profile name used for prompt/ranking heuristics.
        prompt_library_path: Optional prompt library override path.
        optional_context: Optional extra source context for the prompt.
        custom_prompt: Optional full prompt override.
        additional_instructions: Optional prompt suffix for extra guidance.
        max_output_tokens: Optional token limit for generation calls.
        judge_max_output_tokens: Optional token limit for judge calls.
        defragment_runs: Whether to merge safe split runs before labeling.

    Returns:
        Dict[str, Any]: Aggregated suggestions and generation metadata.
    """
    if generator_models:
        generation_models = [
            str(item).strip() for item in generator_models if str(item).strip()
        ]
    else:
        generation_models = [str(model)] * 3

    review_document = docx.Document(docx_path)
    if defragment_runs:
        review_document, _ = defragment_docx_runs(review_document)

    def _run_generation_pass(
        generation_index: int, generation_model: str
    ) -> Dict[str, Any]:
        suggestions = get_labeled_docx_runs(
            docx_path=docx_path,
            custom_people_names=custom_people_names,
            preferred_variable_names=preferred_variable_names,
            openai_client=openai_client,
            openai_api=openai_api,
            openai_base_url=openai_base_url,
            model=generation_model,
            prompt_profile=prompt_profile,
            prompt_library_path=prompt_library_path,
            optional_context=optional_context,
            custom_prompt=custom_prompt,
            additional_instructions=additional_instructions,
            max_output_tokens=max_output_tokens,
            defragment_runs=defragment_runs,
        )
        return {
            "model": generation_model,
            "generation_index": generation_index,
            "suggestions": suggestions,
        }

    if len(generation_models) == 1:
        generated_runs = [_run_generation_pass(0, generation_models[0])]
    else:
        with ThreadPoolExecutor(max_workers=len(generation_models)) as executor:
            generated_runs = list(
                executor.map(
                    lambda item: _run_generation_pass(item[0], item[1]),
                    enumerate(generation_models),
                )
            )

    generation_runs: List[Dict[str, Any]] = []
    for run_info in generated_runs:
        suggestions = run_info["suggestions"]
        validation = validate_docx_label_suggestions(review_document, suggestions)
        generation_runs.append(
            {
                "model": run_info["model"],
                "generation_index": run_info["generation_index"],
                "suggestions": suggestions,
                "validation": validation,
            }
        )

    aggregated = aggregate_docx_label_suggestion_runs(
        review_document,
        generation_runs,
        judge_model=judge_model or model,
        openai_client=openai_client,
        openai_api=openai_api,
        openai_base_url=openai_base_url,
        judge_max_output_tokens=judge_max_output_tokens,
        prompt_profile=prompt_profile,
    )
    aggregated["generation_runs"] = generation_runs
    return aggregated


def _build_paragraph_with_text(source_paragraph: Any, text: str) -> Any:
    """Build a new paragraph element that mirrors a source paragraph's style.

    Args:
        source_paragraph: Existing paragraph used as a formatting reference.
        text: Text content for the new paragraph.

    Returns:
        Any: A ``w:p`` XML element ready to insert into the document.
    """
    paragraph_element = OxmlElement("w:p")

    # Carry paragraph-level style/formatting so inserted tags don't look out of place.
    if source_paragraph is not None and source_paragraph._p.pPr is not None:
        paragraph_element.append(copy.deepcopy(source_paragraph._p.pPr))

    run_element = OxmlElement("w:r")
    _append_text_content(run_element, text)
    paragraph_element.append(run_element)
    return paragraph_element


def add_paragraph_after(paragraph: Any, text: str) -> None:
    """Insert a new paragraph after an existing paragraph.

    Args:
        paragraph: Existing paragraph that should receive a successor.
        text: Text content for the inserted paragraph.
    """
    paragraph._element.addnext(_build_paragraph_with_text(paragraph, text))


def add_paragraph_before(paragraph: Any, text: str) -> None:
    """Insert a new paragraph before an existing paragraph.

    Args:
        paragraph: Existing paragraph that should receive a predecessor.
        text: Text content for the inserted paragraph.
    """
    paragraph._element.addprevious(_build_paragraph_with_text(paragraph, text))


def get_docx_run_text(
    document: Union[docx.document.Document, str], paragraph_number: int, run_number: int
) -> str:
    """Get run text by unified paragraph index across body, tables, headers, and footers.

    Args:
        document: A loaded ``python-docx`` document or a path to a DOCX file.
        paragraph_number: Unified paragraph index in the flattened traversal.
        run_number: Run index within the selected paragraph.

    Returns:
        str: The run text, or an empty string when the coordinates are invalid.
    """
    if isinstance(document, str):
        document = docx.Document(document)

    paragraphs = _collect_target_paragraphs(document)
    if paragraph_number < 0 or paragraph_number >= len(paragraphs):
        return ""

    paragraph = paragraphs[paragraph_number]
    if 0 <= run_number < len(paragraph.runs):
        return paragraph.runs[run_number].text
    return paragraph.text


def get_docx_run_items(
    document: Union[docx.document.Document, str],
    defragment_runs: bool = False,
) -> List[List[Any]]:
    """Return ``[paragraph_index, run_index, run_text]`` across all document parts.

    Args:
        document: A loaded ``python-docx`` document or a path to a DOCX file.
        defragment_runs: Whether to merge safe split runs before traversal.

    Returns:
        List[List[Any]]: Flattened run coordinates and text for the document.
    """
    if isinstance(document, str):
        document = docx.Document(document)
    if defragment_runs:
        document, _ = defragment_docx_runs(document)
    paragraphs = _collect_target_paragraphs(document)
    items: List[List[Any]] = []
    for pnum, paragraph in enumerate(paragraphs):
        for rnum, run in enumerate(paragraph.runs):
            items.append([pnum, rnum, run.text])
    return items


def update_docx(
    document: Union[docx.document.Document, str],
    modified_runs: List[Tuple[int, int, str, int]],
    defragment_runs: bool = False,
    apply_jinja_highlights: bool = False,
) -> docx.document.Document:
    """Update the document with modified runs.

    Args:
        document: The ``python-docx`` document object, or the path to the DOCX file.
        modified_runs: Tuples of paragraph number, run number, modified text, and
            paragraph insertion indicator.
        defragment_runs: Whether to merge safe split runs before applying edits.
        apply_jinja_highlights: Whether to split and shade all Jinja tags after
            edits are applied.

    Returns:
        docx.document.Document: The modified document.
    """
    normalized_runs = _normalize_modified_runs(modified_runs)
    normalized_runs.sort(key=lambda x: (x[0], x[1]), reverse=True)

    if isinstance(document, str):
        document = docx.Document(document)
    if defragment_runs:
        target_paragraph_numbers = sorted(
            {
                paragraph_number
                for paragraph_number, _run_number, _modified_text, new_paragraph in normalized_runs
                if new_paragraph == 0 and paragraph_number >= 0
            }
        )
        if target_paragraph_numbers:
            document, _ = defragment_docx_runs(
                document, paragraph_numbers=target_paragraph_numbers
            )

    paragraphs = _collect_target_paragraphs(document)
    for paragraph_number, run_number, modified_text, new_paragraph in normalized_runs:
        if paragraph_number >= len(paragraphs):
            continue  # Skip invalid paragraph index

        paragraph = paragraphs[paragraph_number]

        if new_paragraph == 1:
            add_paragraph_after(paragraph, modified_text)
            continue
        if new_paragraph == -1:
            add_paragraph_before(paragraph, modified_text)
            continue

        if run_number < len(paragraph.runs):
            paragraph.runs[run_number].text = modified_text
        else:
            # Empty or run-mismatched paragraphs are common in legal forms.
            # Fall back to appending a run so we do not silently drop a valid label.
            paragraph.add_run(modified_text)

    if apply_jinja_highlights:
        document = apply_jinja2_highlights(document)

    return document


def get_labeled_docx_runs(
    docx_path: str,
    custom_people_names: Optional[List[Tuple[str, str]]] = None,
    preferred_variable_names: Optional[Sequence[str]] = None,
    openai_client: Optional[Any] = None,
    openai_api: Optional[str] = None,
    openai_base_url: Optional[str] = None,
    model: str = "gpt-5-mini",
    prompt_profile: str = DEFAULT_DOCX_PROMPT_PROFILE,
    prompt_library_path: Optional[str] = None,
    optional_context: Optional[str] = None,
    custom_prompt: Optional[str] = None,
    additional_instructions: Optional[str] = None,
    max_output_tokens: Optional[int] = None,
    defragment_runs: bool = False,
) -> List[Tuple[int, int, str, int]]:
    """Scan the DOCX and return a list of modified text with Jinja2 variable names inserted.

    Args:
        docx_path: Path to the DOCX file.
        custom_people_names: Optional list of custom ``(name, description)`` pairs.
        preferred_variable_names: Optional preferred variable names to bias prompts.
        openai_client: Optional preconfigured OpenAI client.
        openai_api: Optional API key override.
        openai_base_url: Optional OpenAI-compatible base URL override.
        model: OpenAI model to use.
        prompt_profile: Prompt profile name used to select prompt text and rules.
        prompt_library_path: Optional prompt library override path.
        optional_context: Optional extra source context included in the prompt.
        custom_prompt: Optional full prompt override.
        additional_instructions: Optional extra instructions appended to the prompt.
        max_output_tokens: Optional token limit passed to the chat completion helper.
        defragment_runs: Whether to merge safe split runs before labeling.

    Returns:
        List[Tuple[int, int, str, int]]: Suggested DOCX run replacements.
    """
    openai_base_url = _normalize_openai_base_url(openai_base_url)

    role_description = _get_docx_label_role_description(
        prompt_profile=prompt_profile,
        custom_prompt=custom_prompt,
        prompt_library_path=prompt_library_path,
    )

    custom_name_text = ""
    if custom_people_names is not None:
        if not isinstance(custom_people_names, list):
            raise ValueError(
                "custom_people_names must be a list of [name, description] pairs."
            )
        for item in custom_people_names:
            if not isinstance(item, (list, tuple)) or len(item) != 2:
                raise ValueError(
                    "Each custom_people_names item must be a [name, description] pair."
                )
            name, description = item
            custom_name_text += f"    {name} ({description}), \n"

    preferred_name_text = ""
    if preferred_variable_names:
        normalized_preferred_names = sorted(
            {
                str(name).strip()
                for name in preferred_variable_names
                if str(name).strip()
            }
        )
        if normalized_preferred_names:
            top_level_names = sorted(
                {
                    name.split(".", 1)[0].split("[", 1)[0]
                    for name in normalized_preferred_names
                    if name
                }
            )
            preferred_name_text = (
                "\n\nExisting variable names from the selected Playground interview:\n"
                "Use these names when they fit the document instead of inventing new top-level names.\n"
                "Prefer these top-level objects/lists for people and case data when appropriate:\n    "
                + ", ".join(top_level_names[:80])
            )
            if len(normalized_preferred_names) <= 120:
                preferred_name_text += (
                    "\nSpecific interview names already in use:\n    "
                    + ", ".join(normalized_preferred_names)
                )
            else:
                preferred_name_text += (
                    "\nSpecific interview names already in use (sample):\n    "
                    + ", ".join(normalized_preferred_names[:120])
                )

    rules = f"""
    Rules for variable names:
        1. Variables usually refer to people or their attributes.
        2. People are stored in lists.
        3. We use Docassemble objects and conventions.
        4. Use variable names and patterns from the list below. Invent new variable names when it is appropriate.
        5. Bracketed drafting notes and blank placeholders are good candidates for replacement.
        6. Keep legal role titles like Plaintiff, Defendant, Petitioner, Respondent, Warden,
           Attorney General, and similar caption titles as literal text unless the document clearly
           asks for a specific person's or entity's actual name.

    List names for people:
{custom_name_text}
        users (for the person benefiting from the form, especially when for a pro se filer)
        other_parties (the opposing party in a lawsuit or transactional party)
        plaintiffs
        defendants
        petitioners
        respondents
        children
        spouses
        parents
        caregivers
        attorneys
        translators
        debt_collectors
        creditors
        witnesses
        guardians_ad_litem
        guardians
        decedents
        interested_parties

        Name Forms:
            users (full name of all users)
            users[0] (full name of first user)
            users[0].name.full() (Alternate full name of first user)
            users[0].name.first (First name only)
            users[0].name.middle (Middle name only)
            users[0].name.middle_initial() (First letter of middle name)
            users[0].name.last (Last name only)
            users[0].name.suffix (Suffix of user's name only)

    Attribute names (replace `users` with the appropriate list name):
        Demographic Data:
            users[0].birthdate (Birthdate)
            users[0].age_in_years() (Calculated age based on birthdate)
            users[0].gender (Gender)
            users[0].gender_female (User is female, for checkbox field)
            users[0].gender_male (User is male, for checkbox field)
            users[0].gender_other (User is not male or female, for checkbox field)
            users[0].gender_nonbinary (User identifies as nonbinary, for checkbox field)
            users[0].gender_undisclosed (User chose not to disclose gender, for checkbox field)
            users[0].gender_self_described (User chose to self-describe gender, for checkbox field)
            user_needs_interpreter (User needs an interpreter, for checkbox field)
            user_preferred_language (User's preferred language)

        Addresses:
            users[0].address.block() (Full address, on multiple lines)
            users[0].address.on_one_line() (Full address on one line)
            users[0].address.line_one() (Line one of the address, including unit or apartment number)
            users[0].address.line_two() (Line two of the address, usually city, state, and Zip/postal code)
            users[0].address.address (Street address)
            users[0].address.unit (Apartment, unit, or suite)
            users[0].address.city (City or town)
            users[0].address.state (State, province, or sub-locality)
            users[0].address.zip (Zip or postal code)
            users[0].address.county (County or parish)
            users[0].address.country (Country)

        Other Contact Information:
            users[0].phone_number (Phone number)
            users[0].mobile_number (A phone number explicitly labeled as the "mobile" number)
            users[0].phone_numbers() (A list of both mobile and other phone numbers)
            users[0].email (Email)

        Signatures:
            users[0].signature (Signature)
            signature_date (Date the form is completed)

        Information about Court and Court Processes:
            trial_court (Court's full name)
            trial_court.address.county (County where court is located)
            trial_court.division (Division of court)
            trial_court.department (Department of court)
            docket_number (Case or docket number)
            docket_numbers (A comma-separated list of docket numbers)
            
    When No Existing Variable Name Exists:
        1. Craft short, readable variable names in python snake_case.
        2. Represent people with lists, even if only one person.
        3. Use valid Python variable names within complete Jinja2 tags, like: {{ new_variable_name }}.

        Special endings:
            Suffix _date for date values.
            Suffix _value or _amount for currency values.

        Examples: 
        "(State the reason for eviction)" transforms into `{{ eviction_reason }}`.
    """
    rules += preferred_name_text
    rules += _get_docx_label_rules_addendum(
        prompt_profile=prompt_profile,
        prompt_library_path=prompt_library_path,
    )
    if optional_context and optional_context.strip():
        role_description += (
            "\n\nOptional context for understanding this document:\n"
            + optional_context.strip()
            + "\n\nUse this context only as background to interpret the template. "
            "Do not copy context text into the output unless the document itself calls for it. "
            "Prefer concise placeholders and final-form template text over explanatory material."
        )
    if additional_instructions and additional_instructions.strip():
        role_description += (
            "\n\nAdditional instructions:\n" + additional_instructions.strip()
        )

    encoding = tiktoken.encoding_for_model("gpt-4")

    doc = docx.Document(docx_path)
    if defragment_runs:
        doc, _ = defragment_docx_runs(doc)
    paragraphs = _collect_target_paragraphs(doc)

    items = []
    for pnum, para in enumerate(paragraphs):
        for rnum, run in enumerate(para.runs):
            items.append([pnum, rnum, run.text])

    encoding = tiktoken.encoding_for_model("gpt-4")
    token_count = len(encoding.encode(role_description + rules + repr(items)))
    if token_count > 128000:
        raise Exception(
            f"Input to OpenAI is too long ({token_count} tokens). Maximum is 128000 tokens."
        )

    messages = [
        {"role": "system", "content": role_description + rules},
        {"role": "user", "content": repr(items)},
    ]
    response = chat_completion(
        model=model,
        messages=messages,
        json_mode=True,
        temperature=_get_docx_label_temperature(
            prompt_profile=prompt_profile,
            prompt_library_path=prompt_library_path,
        ),
        max_output_tokens=max_output_tokens,
        openai_client=openai_client,
        openai_api=openai_api,
        openai_base_url=openai_base_url,
    )

    if isinstance(response, str):
        try:
            response = json.loads(response)
        except json.JSONDecodeError as exc:
            raise ValueError("chat_completion returned non-JSON output") from exc
    results = _extract_model_results(response)
    guesses = _normalize_modified_runs(results)
    guesses = _filter_noop_suggestions(doc, guesses)
    return guesses


def modify_docx_with_openai_guesses(docx_path: str) -> docx.document.Document:
    """Uses OpenAI to guess the variable names for a document and then modifies the document with the guesses.

    Args:
        docx_path (str): Path to the DOCX file to modify.

    Returns:
        docx.Document: The modified document, ready to be saved to the same or a new path
    """
    guesses = get_labeled_docx_runs(docx_path)

    return update_docx(docx.Document(docx_path), guesses)


if __name__ == "__main__":
    new_doc = modify_docx_with_openai_guesses(sys.argv[1])
    new_doc.save(sys.argv[1] + ".output.docx")
