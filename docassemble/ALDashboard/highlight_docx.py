import sys
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
import re
from typing import Any, List, Tuple

__all__ = ['highlight_jinja2_content', 'process_document']

def highlight_jinja2_content(paragraph:Paragraph, control_color:str, variable_color:str, comment_color:str) -> None:
    """Highlights text within specific Jinja2 delimiters, ensuring correct scope and avoiding duplicates.
    
    Args:
        paragraph (Paragraph): The paragraph to process
        control_color (str): The color to use for control structures
        variable_color (str): The color to use for variables
        comment_color (str): The color to use for comments
    """
    patterns = [
        (r'\{\{\s*(p|r)\s+(.*?)\s*\}\}', variable_color),  # {{p ... }} or {{r ... }}
        (r'\{%\s*(p|tr|tc|r)\s+(.*?)\s*%\}', control_color),  # {%p ... %}, {%tr ... %}, etc.
        (r'\{\{(.*?)\}\}', variable_color),  # General {{ ... }}
        (r'\{%(.*?)%\}', control_color),  # General {% ... %}
        (r'\{#(.*?)#\}', comment_color)  # Jinja2 comments {# ... #}
    ]

    text = paragraph.text
    cursor = 0
    paragraph.clear()

    # Find all matches and sort them by their start positions
    matches: List[Tuple[int, int, Tuple[str, ...], str]] = []
    for pattern, color in patterns:
        for match in re.finditer(pattern, text):
            matches.append((match.start(), match.end(), match.groups(), color))
    matches.sort()

    # Remove overlapping matches, prioritizing more specific patterns
    filtered_matches: List[Tuple[int, int, Tuple[str, ...], str]] = []
    for match in matches: # type: ignore
        # Unpack the start and end positions from the current match
        match_start = match[0]
        # match_end = match[1]

        # Initialize a flag to check if the current match overlaps with any existing match
        is_overlapping = False

        # Check against all matches already in filtered_matches
        for existing_match in filtered_matches:
            existing_start = existing_match[0]
            existing_end = existing_match[1]

            # Check if the current match overlaps with this existing match
            if existing_start <= match_start < existing_end: # type: ignore
                is_overlapping = True
                break  # Exit the loop early if an overlap is found

        # If no overlaps were found, add the current match to filtered_matches
        if not is_overlapping:
            filtered_matches.append(match) # type: ignore

    # Rebuild the paragraph
    for start, end, groups, color in filtered_matches:
        if start > cursor:
            paragraph.add_run(text[cursor:start])

        full_match = text[start:end]
        if len(groups) > 1:  # Matches with prefixes
            prefix, content = groups
            delimiter_start, delimiter_end = full_match.split(content)[0], full_match.split(content)[-1]
            
            # Add opening delimiter and prefix as plain text
            paragraph.add_run(delimiter_start)
            
            # Highlight only the content
            highlighted_run = paragraph.add_run(content)
            highlighted_run.font.highlight_color = color
            
            # Add closing delimiter as plain text
            paragraph.add_run(delimiter_end)
        else:  # General matches without prefixes
            content = groups[0]
            delimiter_start, delimiter_end = full_match.split(content)[0], full_match.split(content)[-1]
            
            # Add opening delimiter as plain text
            paragraph.add_run(delimiter_start)
            
            # Highlight the content
            highlighted_run = paragraph.add_run(content)
            highlighted_run.font.highlight_color = color
            
            # Add closing delimiter as plain text
            paragraph.add_run(delimiter_end)

        cursor = end

    if cursor < len(text):
        paragraph.add_run(text[cursor:])

def highlight_jinja2_document(doc_path:str, save_path:str) -> None:
    """Highlights Jinja2 content in a Word document and saves the result.
    
    Args:
        doc_path (str): The path to the Word document to process
        save_path (str): The path to save the processed document

    Returns:
        None
    """
    doc = Document(doc_path)
    control_color = WD_COLOR_INDEX.TURQUOISE  # Turquoise for control structures
    variable_color = WD_COLOR_INDEX.YELLOW  # Yellow for variables
    comment_color = WD_COLOR_INDEX.GRAY_25  # Gray for comments

    # Collect all paragraphs from all parts of the document
    elements = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                elements.extend(cell.paragraphs)
    for section in doc.sections:
        elements.extend(section.header.paragraphs if section.header else [])
        elements.extend(section.footer.paragraphs if section.footer else [])

    for paragraph in elements:
        highlight_jinja2_content(paragraph, control_color, variable_color, comment_color)

    doc.save(save_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python highlight_docx.py <path_to_document> <path_to_save_location>")
    else:
        document_path = sys.argv[1]
        save_path = sys.argv[2]
        highlight_jinja2_document(document_path, save_path)
