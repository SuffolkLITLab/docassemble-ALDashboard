# do not pre-load
import os
import tempfile
import unittest

import docx

from docassemble.ALDashboard.court_form_generator import (
    COURT_FORM_SHAPES,
    build_court_form_context,
    generate_court_form_docx,
    generate_court_form_markdown,
)
from docassemble.ALDashboard.court_form_profiles import (
    apply_profile_styles,
    find_docx_fragment,
    list_court_form_profiles,
    load_court_form_profile,
    splice_docx_fragment,
)
from docassemble.ALDashboard.variable_report_generator import (
    extract_interview_questions_and_variables,
    generate_variable_report,
    list_court_form_profile_choices,
    list_court_form_shapes,
)

SAMPLE_YAML = """
---
objects:
  - users: ALPeopleList
  - other_parties: ALPeopleList
---
id: case information
question: |
  Tell us about your case
fields:
  - Docket number: docket_number
  - Were you served?: was_served
    datatype: yesno
---
id: grounds
question: |
  Why are you filing?
fields:
  - Reason: motion_reason
  - Date of the order: order_date
    datatype: date
"""

# Every jurisdiction profile this package ships.
SHIPPED_PROFILES = (
    "generic",
    "ma_trial_court",
    "il_circuit",
    "vt_superior",
    "mi_scao",
    "mn_district",
    "dmass_federal",
)


def _document_text(document) -> str:
    """All the text in a document, body and tables and header and footer."""
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            parts.extend(paragraph.text for paragraph in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class TestCourtFormProfiles(unittest.TestCase):
    def test_every_shipped_profile_loads(self):
        listed = {profile["id"] for profile in list_court_form_profiles()}
        for profile_id in SHIPPED_PROFILES:
            self.assertIn(profile_id, listed)
            profile = load_court_form_profile(profile_id)
            self.assertEqual(profile.id, profile_id)
            self.assertTrue(profile.name)
            self.assertTrue(profile.styles, f"{profile_id} defines no Word styles")

    def test_unknown_profile_falls_back_to_generic(self):
        profile = load_court_form_profile("no_such_court")
        self.assertEqual(profile.id, "generic")

    def test_extends_merges_styles_and_keeps_own_caption(self):
        profile = load_court_form_profile("dmass_federal")
        # Inherited from generic...
        self.assertIn("CourtDocumentTitle", profile.styles)
        self.assertTrue(profile.body.get("motion_intro"))
        # ...while its own overrides win.
        self.assertEqual(profile.styles["CourtBodyText"].line_spacing, 2.0)
        self.assertEqual(profile.label("docket"), "Civil Action No.")

    def test_style_name_accepts_either_spelling_of_a_role(self):
        profile = load_court_form_profile("generic")
        self.assertEqual(profile.style_name("text_style"), "CourtBodyText")
        self.assertEqual(profile.style_name("text"), "CourtBodyText")
        self.assertEqual(profile.style_name("nonexistent", "Normal"), "Normal")

    def test_apply_profile_styles_writes_font_and_spacing(self):
        profile = load_court_form_profile("mi_scao")
        document = docx.Document()
        applied = apply_profile_styles(document, profile)
        self.assertIn("CourtBodyText", applied)
        style = document.styles["CourtBodyText"]
        self.assertEqual(style.font.name, "Arial")
        self.assertEqual(style.font.size.pt, 10.0)

    def test_changing_the_font_in_a_profile_changes_the_document(self):
        """The point of using Word styles: one edit repoints a whole court."""
        profile = load_court_form_profile("ma_trial_court")
        profile.styles["CourtBodyText"].font = "Comic Sans MS"
        document = docx.Document()
        apply_profile_styles(document, profile)
        self.assertEqual(document.styles["CourtBodyText"].font.name, "Comic Sans MS")


class TestCourtFormGeneration(unittest.TestCase):
    def setUp(self):
        self.groups, self.variables = extract_interview_questions_and_variables(
            [SAMPLE_YAML]
        )
        self.tempdir = tempfile.mkdtemp(prefix="court_form_test_")

    def _generate(self, **kwargs):
        kwargs.setdefault("document_title", "Motion to Vacate Default Judgment")
        shape = kwargs.get("shape", "court_form")
        profile_id = kwargs.get("profile_id", "generic")
        kwargs["output_path"] = os.path.join(self.tempdir, f"{profile_id}_{shape}.docx")
        return generate_court_form_docx(self.groups, self.variables, **kwargs)

    def test_every_shape_and_profile_produces_a_document(self):
        for profile_id in SHIPPED_PROFILES:
            for shape in COURT_FORM_SHAPES:
                with self.subTest(profile=profile_id, shape=shape):
                    result = self._generate(shape=shape, profile_id=profile_id)
                    self.assertTrue(os.path.getsize(result["docx_path"]) > 0)
                    self.assertEqual(result["shape"], shape)
                    self.assertEqual(result["profile_id"], profile_id)

    def test_unknown_shape_is_rejected(self):
        with self.assertRaises(ValueError):
            self._generate(shape="subpoena")

    def test_caption_carries_the_jurisdiction_boilerplate(self):
        expectations = {
            "ma_trial_court": "COMMONWEALTH OF MASSACHUSETTS",
            "vt_superior": "STATE OF VERMONT",
            "mi_scao": "STATE OF MICHIGAN",
            "mn_district": "STATE OF MINNESOTA",
            "dmass_federal": "UNITED STATES DISTRICT COURT",
            "il_circuit": "STATE OF ILLINOIS",
        }
        for profile_id, expected in expectations.items():
            with self.subTest(profile=profile_id):
                result = self._generate(profile_id=profile_id)
                text = _document_text(docx.Document(result["docx_path"]))
                self.assertIn(expected, text)

    def test_docket_label_follows_the_jurisdiction(self):
        for profile_id, expected in (
            ("dmass_federal", "Civil Action No."),
            ("mn_district", "Court File Number:"),
            ("ma_trial_court", "DOCKET NO."),
        ):
            with self.subTest(profile=profile_id):
                result = self._generate(profile_id=profile_id)
                text = _document_text(docx.Document(result["docx_path"]))
                self.assertIn(expected, text)

    def test_body_uses_the_profiles_word_styles(self):
        result = self._generate(profile_id="ma_trial_court")
        document = docx.Document(result["docx_path"])
        styles_used = {paragraph.style.name for paragraph in document.paragraphs}
        self.assertIn("CourtBodyText", styles_used)
        self.assertIn("CourtDocumentTitle", styles_used)
        self.assertIn("CourtCaptionHeading", styles_used)

    def test_interview_fields_become_guarded_body_paragraphs(self):
        result = self._generate()
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("{{ showifdef('motion_reason') }}", text)
        # A yes/no field renders as a conditional, not a bare value.
        self.assertIn("showifdef('was_served')", text)
        self.assertIn("1. Docket number:", text)

    def test_caption_boilerplate_is_left_as_plain_jinja(self):
        """Profile-named variables are the profile's contract; do not wrap them."""
        result = self._generate(profile_id="ma_trial_court")
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("{{ trial_court }}", text)
        self.assertIn("{{ docket_number }}", text)

    def test_draft_time_placeholders_are_resolved(self):
        result = self._generate(profile_id="ma_trial_court")
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("Motion to Vacate Default Judgment", text)
        self.assertNotIn("{{ document_title }}", text)
        self.assertNotIn("{{ docket_label }}", text)

    def test_motion_adds_intro_prayer_and_certificate_of_service(self):
        result = self._generate(shape="motion", profile_id="generic")
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("NOW COMES", text)
        self.assertIn("WHEREFORE", text)
        self.assertIn("CERTIFICATE OF SERVICE", text)
        self.assertEqual(result["sections"].get("certificate_of_service"), "yaml")

    def test_affidavit_closes_with_the_jurisdictions_perjury_language(self):
        result = self._generate(shape="affidavit", profile_id="ma_trial_court")
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("pains and penalties of perjury", text)
        self.assertEqual(result["sections"].get("jurat"), "yaml")
        self.assertNotIn("signature", result["sections"])

    def test_court_form_has_no_certificate_of_service_by_default(self):
        result = self._generate(shape="court_form")
        self.assertNotIn("certificate_of_service", result["sections"])

    def test_certificate_of_service_can_be_forced_on(self):
        result = self._generate(shape="court_form", include_certificate_of_service=True)
        self.assertEqual(result["sections"].get("certificate_of_service"), "yaml")

    def test_letter_has_no_caption_header_or_footer(self):
        result = self._generate(shape="letter", profile_id="il_circuit")
        self.assertNotIn("caption", result["sections"])
        self.assertNotIn("header", result["sections"])
        self.assertNotIn("footer", result["sections"])
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("Dear {{ recipient }}", text)
        self.assertIn("Sincerely,", text)
        self.assertNotIn("STATE OF ILLINOIS", text)

    def test_illinois_running_header_names_the_case_number(self):
        result = self._generate(profile_id="il_circuit")
        document = docx.Document(result["docx_path"])
        header_text = "\n".join(
            paragraph.text for paragraph in document.sections[0].header.paragraphs
        )
        self.assertIn("Enter the Case Number", header_text)
        self.assertEqual(result["sections"].get("header"), "yaml")

    def test_page_setup_follows_the_profile(self):
        result = self._generate(profile_id="mi_scao")
        section = docx.Document(result["docx_path"]).sections[0]
        self.assertAlmostEqual(section.top_margin.inches, 0.5, places=2)
        self.assertAlmostEqual(section.left_margin.inches, 0.75, places=2)

    def test_list_variables_become_repeating_tables(self):
        result = self._generate()
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("{%tr for item in showifdef('users') %}", text)
        self.assertIn("{%tr endfor %}", text)

    def test_numbering_can_be_turned_off(self):
        result = self._generate(numbered_paragraphs=False)
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("Docket number:", text)
        self.assertNotIn("1. Docket number:", text)

    def test_unsupplied_form_metadata_shows_as_a_fill_in(self):
        result = self._generate(profile_id="mi_scao")
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("Approved, SCAO", text)
        self.assertIn("[form number]", text)

    def test_form_metadata_can_be_supplied(self):
        result = self._generate(
            profile_id="mi_scao",
            context_extra={"form_code": "Form MC 02", "form_revision": "12/21"},
        )
        text = _document_text(docx.Document(result["docx_path"]))
        self.assertIn("Form MC 02, Rev. 12/21", text)
        self.assertNotIn("[form number]", text)

    def test_build_context_exposes_profile_labels(self):
        profile = load_court_form_profile("mi_scao")
        context = build_court_form_context(profile, "Appearance")
        self.assertEqual(context["document_title"], "Appearance")
        self.assertEqual(context["docket_label"], "CASE NO.")
        self.assertEqual(context["labels.defendant"], "Defendant(s)/Respondent(s)")


class TestDocxFragmentOverride(unittest.TestCase):
    """A section may be replaced by a caption somebody drew in Word."""

    def setUp(self):
        self.groups, self.variables = extract_interview_questions_and_variables(
            [SAMPLE_YAML]
        )
        self.tempdir = tempfile.mkdtemp(prefix="court_form_fragment_")
        self.fragment_root = os.path.join(self.tempdir, "overrides")
        profile_dir = os.path.join(self.fragment_root, "ma_trial_court")
        os.makedirs(profile_dir)
        fragment = docx.Document()
        fragment.add_paragraph("CAPTION AUTHORED IN WORD")
        table = fragment.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Suffolk, ss"
        table.rows[0].cells[1].text = "DOCKET NO. {{ docket_number }}"
        fragment.save(os.path.join(profile_dir, "caption.docx"))

    def test_fragment_is_found(self):
        path = find_docx_fragment("ma_trial_court", "caption", [self.fragment_root])
        self.assertTrue(path and os.path.isfile(path))
        self.assertIsNone(
            find_docx_fragment("ma_trial_court", "signature", [self.fragment_root])
        )

    def test_fragment_replaces_only_that_section(self):
        result = generate_court_form_docx(
            self.groups,
            self.variables,
            shape="court_form",
            profile_id="ma_trial_court",
            document_title="Motion to Vacate",
            output_path=os.path.join(self.tempdir, "out.docx"),
            fragment_dirs=[self.fragment_root],
        )
        self.assertEqual(result["sections"]["caption"], "docx")
        # The sections the override does not cover still come from the profile.
        self.assertEqual(result["sections"]["signature"], "yaml")
        self.assertEqual(result["sections"]["footer"], "yaml")

        document = docx.Document(result["docx_path"])
        text = _document_text(document)
        self.assertIn("CAPTION AUTHORED IN WORD", text)
        self.assertNotIn("COMMONWEALTH OF MASSACHUSETTS", text)
        # The profile's styles are still defined for the rest of the document.
        self.assertIn("CourtBodyText", [style.name for style in document.styles])

    def test_splice_keeps_the_documents_own_section_properties(self):
        target = docx.Document()
        target.sections[0].left_margin = docx.shared.Inches(2.0)
        fragment_path = os.path.join(
            self.fragment_root, "ma_trial_court", "caption.docx"
        )
        copied = splice_docx_fragment(target, fragment_path)
        self.assertGreater(copied, 0)
        out = os.path.join(self.tempdir, "spliced.docx")
        target.save(out)
        self.assertAlmostEqual(
            docx.Document(out).sections[0].left_margin.inches, 2.0, places=2
        )


class TestCourtFormMarkdown(unittest.TestCase):
    def setUp(self):
        self.groups, self.variables = extract_interview_questions_and_variables(
            [SAMPLE_YAML]
        )

    def test_markdown_carries_caption_and_body(self):
        text = generate_court_form_markdown(
            self.groups,
            self.variables,
            shape="motion",
            profile_id="vt_superior",
            document_title="Motion to Enforce",
        )
        self.assertIn("STATE OF VERMONT", text)
        self.assertIn("# Motion to Enforce", text)
        self.assertIn("NOW COMES", text)
        self.assertIn("${ showifdef('motion_reason') }", text)

    def test_letter_markdown_has_no_caption(self):
        text = generate_court_form_markdown(
            self.groups,
            self.variables,
            shape="letter",
            profile_id="ma_trial_court",
            document_title="Request for Records",
        )
        self.assertNotIn("COMMONWEALTH OF MASSACHUSETTS", text)
        self.assertIn("RE: Request for Records", text)
        self.assertIn("Sincerely,", text)


class TestVariableReportIntegration(unittest.TestCase):
    """The court shapes reach callers through the variable report's own API."""

    def test_intake_remains_the_default_and_is_unchanged(self):
        result = generate_variable_report([SAMPLE_YAML], report_title="Intake")
        self.assertEqual(result["shape"], "intake")
        self.assertIn("| Field | Value |", result["mako_markdown"])
        self.assertNotIn("profile_id", result)

    def test_court_shape_reports_which_templates_it_used(self):
        result = generate_variable_report(
            [SAMPLE_YAML],
            report_title="Motion to Vacate",
            shape="motion",
            court_profile="mn_district",
        )
        self.assertEqual(result["shape"], "motion")
        self.assertEqual(result["profile_id"], "mn_district")
        self.assertEqual(result["profile_name"], "Minnesota District Court")
        self.assertEqual(result["sections"]["caption"], "yaml")
        self.assertIn("CourtBodyText", result["styles"])
        self.assertTrue(os.path.isfile(result["docx_path"]))

    def test_output_path_is_honoured_for_court_shapes(self):
        with tempfile.TemporaryDirectory() as tempdir:
            target = os.path.join(tempdir, "drafted.docx")
            result = generate_variable_report(
                [SAMPLE_YAML],
                report_title="Affidavit of Indigency",
                shape="affidavit",
                court_profile="ma_trial_court",
                output_docx_path=target,
            )
            self.assertEqual(result["docx_path"], target)
            self.assertTrue(os.path.isfile(target))

    def test_choice_helpers_list_shapes_and_profiles(self):
        shapes = {item["value"] for item in list_court_form_shapes()}
        self.assertEqual(shapes, {"intake"} | set(COURT_FORM_SHAPES))
        profiles = {item["value"] for item in list_court_form_profile_choices()}
        for profile_id in SHIPPED_PROFILES:
            self.assertIn(profile_id, profiles)


class TestCourtFormAPI(unittest.TestCase):
    """The same shapes over the HTTP API the Weaver and other clients use."""

    def setUp(self):
        from docassemble.ALDashboard.api_dashboard_utils import (
            court_form_profiles_payload_from_options,
            variable_report_payload_from_options,
        )

        self.report = variable_report_payload_from_options
        self.profiles = court_form_profiles_payload_from_options

    def test_court_shape_returns_the_docx_by_default(self):
        payload = self.report(
            {
                "yaml_text": SAMPLE_YAML,
                "report_title": "Motion to Dismiss",
                "shape": "motion",
                "court_profile": "il_circuit",
            }
        )
        self.assertEqual(payload["shape"], "motion")
        self.assertEqual(payload["profile_id"], "il_circuit")
        self.assertIn("caption", payload["sections"])
        self.assertTrue(payload["docx_base64"])

    def test_intake_response_stays_lean(self):
        payload = self.report({"yaml_text": SAMPLE_YAML})
        self.assertEqual(payload["shape"], "intake")
        self.assertNotIn("docx_base64", payload)
        self.assertNotIn("profile_id", payload)

    def test_docx_can_be_requested_for_intake_too(self):
        payload = self.report({"yaml_text": SAMPLE_YAML, "include_docx_base64": True})
        self.assertTrue(payload["docx_base64"])

    def test_unknown_shape_is_a_validation_error(self):
        from docassemble.ALDashboard.api_dashboard_utils import (
            DashboardAPIValidationError,
        )

        with self.assertRaises(DashboardAPIValidationError):
            self.report({"yaml_text": SAMPLE_YAML, "shape": "subpoena"})

    def test_profiles_endpoint_lists_shapes_profiles_and_sections(self):
        payload = self.profiles({})
        listed = {profile["id"] for profile in payload["profiles"]}
        for profile_id in SHIPPED_PROFILES:
            self.assertIn(profile_id, listed)
        self.assertIn("court_form", {item["value"] for item in payload["shapes"]})
        self.assertIn("caption", payload["sections"])


if __name__ == "__main__":
    unittest.main()
