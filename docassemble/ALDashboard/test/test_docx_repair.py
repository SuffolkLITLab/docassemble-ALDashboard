# do not pre-load
import os
import tempfile
import unittest
import zipfile
from unittest.mock import patch

import docx

from docassemble.ALDashboard.docx_repair import (
    repair_docx_xml_conservatively,
    rescue_docx_to_shell,
    roundtrip_docx_via_soffice,
)


class TestDocxRepair(unittest.TestCase):
    def _build_docx(self, parts):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
            docx_path = temp_docx.name
        with zipfile.ZipFile(docx_path, "w") as archive:
            for name, content in parts.items():
                archive.writestr(name, content)
        return docx_path

    def test_conservative_repair_removes_dangling_relationships(self):
        source = self._build_docx(
            {
                "[Content_Types].xml": "<Types/>",
                "_rels/.rels": "<Relationships/>",
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p><w:r><w:t>Hello</w:t></w:r></w:p></w:body>
</w:document>""",
                "word/_rels/document.xml.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Target="media/missing.png" />
</Relationships>""",
            }
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_out:
            output = temp_out.name

        try:
            report = repair_docx_xml_conservatively(source, output)
            self.assertEqual(len(report["removed_relationships"]), 1)
            with zipfile.ZipFile(output, "r") as archive:
                rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
            self.assertNotIn("media/missing.png", rels_xml)
        finally:
            if os.path.exists(source):
                os.remove(source)
            if os.path.exists(output):
                os.remove(output)

    def test_rescue_docx_to_shell_rebuilds_document_text(self):
        source = self._build_docx(
            {
                "[Content_Types].xml": "<Types/>",
                "_rels/.rels": "<Relationships/>",
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>First paragraph</w:t></w:r></w:p>
    <w:tbl>
      <w:tr>
        <w:tc><w:p><w:r><w:t>Cell A</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>Cell B</w:t></w:r></w:p></w:tc>
      </w:tr>
    </w:tbl>
  </w:body>
</w:document>""",
                "word/_rels/document.xml.rels": '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>',
            }
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_out:
            output = temp_out.name

        try:
            report = rescue_docx_to_shell(source, output)
            rescued = docx.Document(output)
            self.assertEqual(report["paragraphs"], 1)
            self.assertEqual(report["tables"], 1)
            self.assertIn("First paragraph", rescued.paragraphs[0].text)
            self.assertEqual(rescued.tables[0].cell(0, 0).text, "Cell A")
        finally:
            if os.path.exists(source):
                os.remove(source)
            if os.path.exists(output):
                os.remove(output)

    @patch("docassemble.ALDashboard.docx_repair.shutil.which", return_value=None)
    def test_roundtrip_reports_when_soffice_is_missing(self, _mock_which):
        report = roundtrip_docx_via_soffice("input.docx", "output.docx")

        self.assertFalse(report["available"])
        self.assertIn("soffice", report["message"])


if __name__ == "__main__":
    unittest.main()
