import unittest
from pathlib import Path
import tempfile
from unittest.mock import patch

import docx

from docassemble.ALDashboard.docx_wrangling import (
    _get_docx_label_role_description,
    _normalize_openai_base_url,
    aggregate_docx_label_suggestion_runs,
    defragment_docx_runs,
    get_labeled_docx_runs,
    get_voted_docx_label_suggestions,
    review_flagged_docx_label_suggestions,
    update_docx,
    validate_docx_label_suggestions,
)


class TestDocxWranglingUpdateDocx(unittest.TestCase):
    def test_normalize_openai_base_url_appends_azure_api_path(self):
        self.assertEqual(
            _normalize_openai_base_url("https://workflowdocs.openai.azure.com"),
            "https://workflowdocs.openai.azure.com/openai/v1/",
        )
        self.assertEqual(
            _normalize_openai_base_url("https://workflowdocs.openai.azure.com/"),
            "https://workflowdocs.openai.azure.com/openai/v1/",
        )
        self.assertEqual(
            _normalize_openai_base_url(
                "https://workflowdocs.openai.azure.com/openai/v1"
            ),
            "https://workflowdocs.openai.azure.com/openai/v1/",
        )
        self.assertEqual(
            _normalize_openai_base_url("https://api.openai.com/v1/"),
            "https://api.openai.com/v1/",
        )
        self.assertEqual(
            _normalize_openai_base_url(
                "https://quint-mln02sj6-eastus2.cognitiveservices.azure.com"
            ),
            "https://quint-mln02sj6-eastus2.cognitiveservices.azure.com/openai/v1/",
        )
        self.assertEqual(
            _normalize_openai_base_url(
                "https://quint-mln02sj6-eastus2.cognitiveservices.azure.com/openai/deployments/gpt-5/chat/completions?api-version=2025-01-01-preview"
            ),
            "https://quint-mln02sj6-eastus2.cognitiveservices.azure.com/openai/v1/",
        )

    def test_update_docx_replaces_existing_run(self):
        document = docx.Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Name: ____")

        updated = update_docx(document, [(0, 0, "Name: {{ users[0] }}", 0)])

        self.assertEqual(updated.paragraphs[0].runs[0].text, "Name: {{ users[0] }}")

    def test_update_docx_inserts_wordprocessingml_safe_paragraphs(self):
        document = docx.Document()
        paragraph = document.add_paragraph("Anchor")

        updated = update_docx(
            document,
            [
                (0, 0, "{%p if has_value %}\t", -1),
                (0, 0, "{%p endif %}\n", 1),
            ],
        )

        self.assertEqual(updated.paragraphs[0].text, "{%p if has_value %}\t")
        self.assertEqual(updated.paragraphs[1].text, "Anchor")
        self.assertEqual(updated.paragraphs[2].text, "{%p endif %}\n")

        before_xml = updated.paragraphs[0]._p.xml
        after_xml = updated.paragraphs[2]._p.xml

        # New paragraphs should contain proper run/text elements, not raw text directly under <w:p>.
        self.assertIn("<w:r>", before_xml)
        self.assertIn("<w:t", before_xml)
        self.assertIn("<w:tab/>", before_xml)
        self.assertIn("<w:br/>", after_xml)

    def test_update_docx_appends_run_when_run_index_is_out_of_bounds(self):
        document = docx.Document()
        document.add_paragraph("Only one run")

        updated = update_docx(document, [(0, 99, "Fallback run", 0)])

        self.assertEqual(updated.paragraphs[0].runs[-1].text, "Fallback run")

    def test_update_docx_ignores_invalid_items_and_accepts_dict_items(self):
        document = docx.Document()
        document.add_paragraph("Original")

        updated = update_docx(
            document,
            [
                {"paragraph": 0, "run": 0, "text": "From dict", "new_paragraph": 0},
                ["bad", "item"],
                None,
            ],
        )

        self.assertEqual(updated.paragraphs[0].runs[0].text, "From dict")

    def test_defragment_docx_runs_merges_plain_text_runs(self):
        document = docx.Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Property A")
        paragraph.add_run("ddress:")
        paragraph.add_run(" ")

        updated, stats = defragment_docx_runs(document)

        self.assertEqual(updated.paragraphs[0].text, "Property Address: ")
        self.assertEqual(len(updated.paragraphs[0].runs), 1)
        self.assertEqual(stats["paragraphs_defragmented"], 1)
        self.assertEqual(stats["runs_removed"], 2)

    def test_update_docx_with_defragment_runs_handles_fragmented_fixture(self):
        fixture_path = Path(__file__).with_name("condo_deed.docx")
        document = docx.Document(str(fixture_path))

        updated = update_docx(
            document,
            [(28, 0, "Property Address: {{ property.address.on_one_line() }}", 0)],
            defragment_runs=True,
        )

        self.assertEqual(
            updated.paragraphs[28].text,
            "Property Address: {{ property.address.on_one_line() }}",
        )

    def test_prompt_profile_adds_litigation_specific_guidance(self):
        standard = _get_docx_label_role_description(prompt_profile="standard")
        litigation = _get_docx_label_role_description(
            prompt_profile="litigation_template"
        )

        self.assertNotIn("This document may be a pleading", standard)
        self.assertIn("This document may be a pleading", litigation)
        self.assertIn("repeated underscores", litigation)
        self.assertIn("bracketed drafting notes", litigation)
        self.assertIn("Keep titles like Petitioner", litigation)
        self.assertNotIn("immigration habeas petition practice advisory", litigation)

    @patch("docassemble.ALDashboard.docx_wrangling.chat_completion")
    def test_get_labeled_docx_runs_uses_custom_prompt_library(self, mock_chat_completion):
        document = docx.Document()
        document.add_paragraph("Name: ____")
        mock_chat_completion.return_value = {"results": []}

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "sample.docx"
            prompt_path = Path(tmpdir) / "custom_labeler.yml"
            document.save(str(docx_path))
            prompt_path.write_text(
                """
docx:
  default_prompt_profile: alternate
  prompt_profiles:
    alternate:
      label: Alternate
      role_description: Alternate role description
      rules_addendum: Alternate rules addendum
""".strip(),
                encoding="utf-8",
            )

            get_labeled_docx_runs(
                str(docx_path),
                prompt_profile="alternate",
                prompt_library_path=str(prompt_path),
            )

        system_message = mock_chat_completion.call_args.kwargs["messages"][0]["content"]
        self.assertIn("Alternate role description", system_message)
        self.assertIn("Alternate rules addendum", system_message)

    def test_validator_flags_inline_placeholder_rendered_as_new_paragraph(self):
        document = docx.Document()
        paragraph = document.add_paragraph("Grants to:")
        paragraph.add_run(" ")

        validation = validate_docx_label_suggestions(
            document,
            [(0, 1, "{{ other_parties[0].name.full() }}", 1)],
        )

        self.assertEqual(validation["flagged_count"], 1)
        codes = {flag["code"] for flag in validation["results"][0]["flags"]}
        self.assertIn("paragraph_insert_without_control_tag", codes)
        self.assertIn("inline_placeholder_emitted_as_paragraph", codes)

    def test_validator_flags_leftover_placeholder_markers_on_fragmented_run(self):
        fixture_path = Path(__file__).with_name("condo_deed.docx")
        document = docx.Document(str(fixture_path))

        validation = validate_docx_label_suggestions(
            document,
            [(28, 0, "Property Address: {{ property.address.on_one_line() }}", 0)],
        )

        self.assertEqual(validation["flagged_count"], 1)
        codes = {flag["code"] for flag in validation["results"][0]["flags"]}
        self.assertIn("fragmented_word_boundary", codes)
        self.assertIn("leftover_word_fragments", codes)

    @patch("docassemble.ALDashboard.docx_wrangling.chat_completion")
    def test_ai_review_only_reviews_flagged_suggestions(self, mock_chat_completion):
        document = docx.Document()
        paragraph = document.add_paragraph("Grants to:")
        paragraph.add_run(" ")
        deterministic = validate_docx_label_suggestions(
            document,
            [(0, 1, "{{ other_parties[0].name.full() }}", 1)],
        )
        mock_chat_completion.return_value = {
            "reviews": [
                {"index": 0, "verdict": "reject", "reason": "Should be inline."}
            ]
        }

        reviewed = review_flagged_docx_label_suggestions(
            document,
            [(0, 1, "{{ other_parties[0].name.full() }}", 1)],
            deterministic,
        )

        self.assertTrue(reviewed["performed"])
        self.assertEqual(reviewed["reviews"][0]["verdict"], "reject")
        review_payload = mock_chat_completion.call_args.kwargs["messages"][1]["content"]
        self.assertIn("inline_placeholder_emitted_as_paragraph", review_payload)

    def test_validator_does_not_flag_single_tab_as_leftover_placeholder(self):
        document = docx.Document()
        document.add_paragraph("One-year deadline:\t____________________")

        validation = validate_docx_label_suggestions(
            document,
            [(0, 0, "One-year deadline:\t{{ asylum_date.format('MM/dd/yyyy') }}", 0)],
        )

        codes = {flag["code"] for flag in validation["results"][0]["flags"]}
        self.assertNotIn("leftover_placeholder_markers", codes)

    @patch("docassemble.ALDashboard.docx_wrangling.chat_completion")
    def test_get_labeled_docx_runs_filters_unchanged_and_normalizes_paragraph_tags(
        self, mock_chat_completion
    ):
        document = docx.Document()
        document.add_paragraph("Attorney heading")
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            mock_chat_completion.return_value = {
                "results": [
                    [0, 0, "Attorney heading", 0],
                    [0, 0, "{% if attorneys %}", -1],
                    [0, 0, "{% endif %}", 1],
                ]
            }

            results = get_labeled_docx_runs(tmp.name, model="gpt-5-mini")

        self.assertEqual(
            results,
            [
                (0, 0, "{%p if attorneys %}", -1),
                (0, 0, "{%p endif %}", 1),
            ],
        )

    @patch("docassemble.ALDashboard.docx_wrangling.chat_completion")
    def test_get_labeled_docx_runs_normalizes_azure_base_url(self, mock_chat_completion):
        document = docx.Document()
        document.add_paragraph("Name: ____")
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            mock_chat_completion.return_value = {"results": []}

            get_labeled_docx_runs(
                tmp.name,
                model="gpt-5-mini",
                openai_api="test-key",
                openai_base_url="https://workflowdocs.openai.azure.com",
            )

        self.assertEqual(
            mock_chat_completion.call_args.kwargs["openai_base_url"],
            "https://workflowdocs.openai.azure.com/openai/v1/",
        )

    @patch("docassemble.ALDashboard.docx_wrangling.chat_completion")
    def test_get_labeled_docx_runs_uses_selected_prompt_profile(
        self, mock_chat_completion
    ):
        document = docx.Document()
        document.add_paragraph("[NAME]")
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            mock_chat_completion.return_value = {"results": []}

            get_labeled_docx_runs(
                tmp.name,
                model="gpt-5-mini",
                prompt_profile="litigation_template",
            )

        system_prompt = mock_chat_completion.call_args.kwargs["messages"][0]["content"]
        self.assertIn("This document may be a pleading", system_prompt)
        self.assertIn("bracketed drafting notes", system_prompt)
        self.assertEqual(mock_chat_completion.call_args.kwargs["temperature"], 0.5)

    @patch("docassemble.ALDashboard.docx_wrangling.chat_completion")
    def test_get_labeled_docx_runs_includes_optional_context_only_when_provided(
        self, mock_chat_completion
    ):
        document = docx.Document()
        document.add_paragraph("Case No. ____")
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            mock_chat_completion.return_value = {"results": []}

            get_labeled_docx_runs(
                tmp.name,
                model="gpt-5-mini",
                prompt_profile="litigation_template",
                optional_context="This template is based on an immigration habeas guide.",
            )
            with_context = mock_chat_completion.call_args.kwargs["messages"][0]["content"]

            get_labeled_docx_runs(
                tmp.name,
                model="gpt-5-mini",
                prompt_profile="litigation_template",
            )
            without_context = mock_chat_completion.call_args.kwargs["messages"][0]["content"]

        self.assertIn("Optional context for understanding this document", with_context)
        self.assertIn("immigration habeas guide", with_context)
        self.assertNotIn("Optional context for understanding this document", without_context)

    @patch("docassemble.ALDashboard.docx_wrangling.chat_completion")
    def test_get_labeled_docx_runs_uses_default_temperature_for_standard_profile(
        self, mock_chat_completion
    ):
        document = docx.Document()
        document.add_paragraph("Name: ____")
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            mock_chat_completion.return_value = {"results": []}

            get_labeled_docx_runs(
                tmp.name,
                model="gpt-5-mini",
                prompt_profile="standard",
            )

        self.assertEqual(mock_chat_completion.call_args.kwargs["temperature"], 0.5)

    def test_aggregate_docx_label_suggestion_runs_prefers_two_clean_votes(self):
        document = docx.Document()
        document.add_paragraph("Name: ____")

        run_a = [(0, 0, "Name: {{ users[0].name.full() }}", 0)]
        run_b = [(0, 0, "Name: {{ users[0].name.full() }}", 0)]
        run_c = [(0, 0, "Name: {{ users[0] }}", 0)]

        aggregated = aggregate_docx_label_suggestion_runs(
            document,
            [
                {
                    "model": "gpt-5-mini",
                    "generation_index": 0,
                    "suggestions": run_a,
                    "validation": validate_docx_label_suggestions(document, run_a),
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 1,
                    "suggestions": run_b,
                    "validation": validate_docx_label_suggestions(document, run_b),
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 2,
                    "suggestions": run_c,
                    "validation": validate_docx_label_suggestions(document, run_c),
                },
            ],
            judge_model="gpt-5-mini",
        )

        self.assertEqual(len(aggregated["suggestions"]), 1)
        suggestion = aggregated["suggestions"][0]
        self.assertEqual(suggestion["text"], "Name: {{ users[0].name.full() }}")
        self.assertEqual(suggestion["confidence"], "medium")
        self.assertEqual(suggestion["clean_vote_count"], 2)
        self.assertEqual(len(suggestion["alternates"]), 1)
        self.assertFalse(aggregated["judge_review"]["performed"])

    def test_aggregate_docx_label_suggestion_runs_litigation_accepts_clean_singleton(
        self,
    ):
        document = docx.Document()
        document.add_paragraph("Case No. _______________")

        singleton = [(0, 0, "Case No. {{ docket_number }}", 0)]

        aggregated = aggregate_docx_label_suggestion_runs(
            document,
            [
                {
                    "model": "gpt-5-mini",
                    "generation_index": 0,
                    "suggestions": singleton,
                    "validation": validate_docx_label_suggestions(document, singleton),
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 1,
                    "suggestions": [],
                    "validation": {"results": [], "flagged_count": 0},
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 2,
                    "suggestions": [],
                    "validation": {"results": [], "flagged_count": 0},
                },
            ],
            judge_model="gpt-5-mini",
            prompt_profile="litigation_template",
        )

        self.assertEqual(len(aggregated["suggestions"]), 1)
        self.assertEqual(aggregated["suggestions"][0]["text"], "Case No. {{ docket_number }}")
        self.assertEqual(aggregated["suggestions"][0]["confidence"], "low")
        self.assertFalse(aggregated["judge_review"]["performed"])

    def test_aggregate_docx_label_suggestion_runs_litigation_relaxes_caption_tab_flag(
        self,
    ):
        document = docx.Document()
        document.add_paragraph("[NAME],\t\t)")

        caption_candidate = [(0, 0, "{{ users[0].name.full() }},\t\t)", 0)]
        caption_validation = {
            "results": [
                {
                    "index": 0,
                    "paragraph": 0,
                    "run": 0,
                    "new_paragraph": 0,
                    "source_paragraph_text": "[NAME],\t\t)",
                    "source_run_text": "[NAME],\t\t)",
                    "suggested_text": "{{ users[0].name.full() }},\t\t)",
                    "simulated_paragraph_text": "{{ users[0].name.full() }},\t\t)",
                    "flags": [
                        {
                            "code": "leftover_placeholder_markers",
                            "message": "Tabs remain for alignment.",
                        }
                    ],
                }
            ],
            "flagged_count": 1,
        }

        aggregated = aggregate_docx_label_suggestion_runs(
            document,
            [
                {
                    "model": "gpt-5-mini",
                    "generation_index": 0,
                    "suggestions": caption_candidate,
                    "validation": caption_validation,
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 1,
                    "suggestions": caption_candidate,
                    "validation": caption_validation,
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 2,
                    "suggestions": [],
                    "validation": {"results": [], "flagged_count": 0},
                },
            ],
            judge_model="gpt-5-mini",
            prompt_profile="litigation_template",
        )

        self.assertEqual(len(aggregated["suggestions"]), 1)
        self.assertEqual(aggregated["suggestions"][0]["clean_vote_count"], 2)
        self.assertEqual(aggregated["suggestions"][0]["validation_flags"], [])
        self.assertEqual(aggregated["suggestions"][0]["confidence"], "medium")
        self.assertFalse(aggregated["judge_review"]["performed"])

    @patch("docassemble.ALDashboard.docx_wrangling.review_docx_label_candidate_groups")
    def test_aggregate_docx_label_suggestion_runs_uses_judge_for_singletons(
        self, mock_chat_completion
    ):
        document = docx.Document()
        document.add_paragraph("Name: ____")

        run_a = [(0, 0, "Name: {{ users[0].name.full() }}", 0)]
        run_b = [(0, 0, "Name: {{ users[0].name.first }}", 0)]
        run_c = [(0, 0, "Name: {{ users[0].name.last }}", 0)]

        def fake_review(candidate_groups, **kwargs):
            candidate_index = next(
                candidate["candidate_index"]
                for candidate in candidate_groups[0]["candidates"]
                if candidate["text"] == "Name: {{ users[0].name.full() }}"
            )
            return {
                "performed": True,
                "reviews": [
                    {
                        "group_index": 0,
                        "decision": "choose",
                        "candidate_index": candidate_index,
                        "reason": "Best reusable full-name candidate.",
                    }
                ],
            }

        mock_chat_completion.side_effect = fake_review

        aggregated = aggregate_docx_label_suggestion_runs(
            document,
            [
                {
                    "model": "gpt-5-mini",
                    "generation_index": 0,
                    "suggestions": run_a,
                    "validation": validate_docx_label_suggestions(document, run_a),
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 1,
                    "suggestions": run_b,
                    "validation": validate_docx_label_suggestions(document, run_b),
                },
                {
                    "model": "gpt-5-mini",
                    "generation_index": 2,
                    "suggestions": run_c,
                    "validation": validate_docx_label_suggestions(document, run_c),
                },
            ],
            judge_model="gpt-5-mini",
        )

        self.assertTrue(aggregated["judge_review"]["performed"])
        self.assertEqual(len(aggregated["suggestions"]), 1)
        self.assertEqual(
            aggregated["suggestions"][0]["text"], "Name: {{ users[0].name.full() }}"
        )
        self.assertEqual(aggregated["suggestions"][0]["confidence"], "low")

    @patch("docassemble.ALDashboard.docx_wrangling.get_labeled_docx_runs")
    def test_get_voted_docx_label_suggestions_passes_prompt_profile(
        self, mock_get_labeled_docx_runs
    ):
        document = docx.Document()
        document.add_paragraph("Name: ____")
        mock_get_labeled_docx_runs.return_value = []

        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            get_voted_docx_label_suggestions(
                tmp.name,
                model="gpt-5-mini",
                prompt_profile="litigation_template",
                generator_models=["gpt-5-mini"],
            )

        self.assertEqual(mock_get_labeled_docx_runs.call_count, 1)
        self.assertEqual(
            mock_get_labeled_docx_runs.call_args.kwargs["prompt_profile"],
            "litigation_template",
        )

    @patch("docassemble.ALDashboard.docx_wrangling.get_labeled_docx_runs")
    def test_get_voted_docx_label_suggestions_passes_optional_context(
        self, mock_get_labeled_docx_runs
    ):
        document = docx.Document()
        document.add_paragraph("Case No. ____")
        mock_get_labeled_docx_runs.return_value = []

        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            get_voted_docx_label_suggestions(
                tmp.name,
                model="gpt-5-mini",
                optional_context="Helpful background about this template.",
                generator_models=["gpt-5-mini"],
            )

        self.assertEqual(
            mock_get_labeled_docx_runs.call_args.kwargs["optional_context"],
            "Helpful background about this template.",
        )


if __name__ == "__main__":
    unittest.main()
