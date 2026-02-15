import unittest

import docx

from docassemble.ALDashboard.docx_wrangling import (
    _collect_target_paragraphs,
    get_docx_run_text,
    update_docx,
)


class TestDocxWranglingScopes(unittest.TestCase):
    def _find_paragraph_index(self, document, text):
        paragraphs = _collect_target_paragraphs(document)
        for idx, paragraph in enumerate(paragraphs):
            if paragraph.text == text:
                return idx
        raise AssertionError(f"Paragraph text not found: {text}")

    def test_update_docx_updates_table_header_and_footer_runs(self):
        document = docx.Document()
        document.add_paragraph("Body text")

        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Cell placeholder"

        section = document.sections[0]
        section.header.paragraphs[0].text = "Header placeholder"
        section.footer.paragraphs[0].text = "Footer placeholder"

        cell_idx = self._find_paragraph_index(document, "Cell placeholder")
        header_idx = self._find_paragraph_index(document, "Header placeholder")
        footer_idx = self._find_paragraph_index(document, "Footer placeholder")

        updated = update_docx(
            document,
            [
                (cell_idx, 0, "{{ table_value }}", 0),
                (header_idx, 0, "{{ header_value }}", 0),
                (footer_idx, 0, "{{ footer_value }}", 0),
            ],
        )

        self.assertEqual(get_docx_run_text(updated, cell_idx, 0), "{{ table_value }}")
        self.assertEqual(
            get_docx_run_text(updated, header_idx, 0), "{{ header_value }}"
        )
        self.assertEqual(
            get_docx_run_text(updated, footer_idx, 0), "{{ footer_value }}"
        )


if __name__ == "__main__":
    unittest.main()
