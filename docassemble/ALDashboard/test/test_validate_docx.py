# do not pre-load

import unittest
from typing import Optional
import docx
from unittest.mock import patch
from docassemble.ALDashboard.validate_docx import (
    analyze_docx_template_markup,
    detect_docx_automation_features,
    get_jinja_errors,
    get_jinja_template_validation,
    strip_docx_problem_controls,
    validate_docx_ooxml_schema,
)
from pathlib import Path
import tempfile
import zipfile
import os
import xml.etree.ElementTree as ET


class TestGetJinjaErrors(unittest.TestCase):
    def _build_docx(self, parts):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
            docx_path = temp_docx.name
        with zipfile.ZipFile(docx_path, "w") as archive:
            for name, content in parts.items():
                archive.writestr(name, content)
        return docx_path

    def _read_part(self, docx_path: str, part_name: str) -> str:
        with zipfile.ZipFile(docx_path, "r") as archive:
            return archive.read(part_name).decode("utf-8", errors="ignore")

    def test_working_template(self):
        working_template = Path(__file__).parent / "made_up_variables.docx"
        result: Optional[str] = get_jinja_errors(working_template)
        self.assertIsNone(result)

    def test_failing_template(self):
        failing_template = Path(__file__).parent / "valid_word_invalid_jinja.docx"
        result: Optional[str] = get_jinja_errors(failing_template)
        self.assertIsNotNone(result)
        self.assertIsInstance(result, str)

    def test_detects_field_controls_in_body(self):
        docx_path = self._build_docx(
            {
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p><w:fldChar w:fldCharType="begin"/><w:instrText>MERGEFIELD ClientName</w:instrText></w:p></w:body>
</w:document>""",
            }
        )
        try:
            findings = detect_docx_automation_features(docx_path)
            codes = {item["code"] for item in findings["warning_details"]}
            self.assertIn("classic_fields", codes)
            self.assertIn("field_instructions", codes)
        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)

    def test_detects_sdt_controls(self):
        docx_path = self._build_docx(
            {
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:sdt><w:sdtPr><w:text/><w:tag w:val="x"/></w:sdtPr><w:sdtContent/></w:sdt></w:body>
</w:document>""",
            }
        )
        try:
            findings = detect_docx_automation_features(docx_path)
            codes = {item["code"] for item in findings["warning_details"]}
            self.assertIn("structured_document_tags", codes)
            self.assertIn("sdt_plain_text_control", codes)
            self.assertIn("sdt_metadata", codes)
        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)

    def test_detects_fragmented_runs(self):
        runs = "".join(f"<w:r><w:t>fragment{i:02d}</w:t></w:r>" for i in range(12))
        docx_path = self._build_docx(
            {
                "word/document.xml": f"""<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p>{runs}</w:p></w:body>
</w:document>""",
            }
        )
        try:
            findings = detect_docx_automation_features(docx_path)
            codes = {item["code"] for item in findings["warning_details"]}
            self.assertIn("fragmented_runs", codes)
        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)

    def test_ignores_footer_page_number_controls(self):
        docx_path = self._build_docx(
            {
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:t>x</w:t></w:p></w:body></w:document>""",
                "word/footer1.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:sdt><w:sdtPr><w:docPartObj><w:docPartGallery w:val="Page Numbers (Bottom of Page)"/></w:docPartObj></w:sdtPr>
    <w:sdtContent><w:p><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> PAGE </w:instrText></w:r></w:p></w:sdtContent>
  </w:sdt>
</w:ftr>""",
            }
        )
        try:
            findings = detect_docx_automation_features(docx_path)
            codes = {item["code"] for item in findings["warning_details"]}
            self.assertNotIn("classic_fields", codes)
            self.assertNotIn("structured_document_tags", codes)
        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)

    def test_strip_docx_problem_controls_removes_sdt_and_non_whitelisted_simple_fields(
        self,
    ):
        input_path = self._build_docx(
            {
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:sdt><w:sdtPr><w:text/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>inside sdt</w:t></w:r></w:p></w:sdtContent></w:sdt>
    <w:p><w:fldSimple w:instr=" MERGEFIELD ClientName "><w:r><w:t>Client Name</w:t></w:r></w:fldSimple></w:p>
    <w:p><w:fldSimple w:instr=" PAGE "><w:r><w:t>2</w:t></w:r></w:fldSimple></w:p>
  </w:body>
</w:document>""",
                "[Content_Types].xml": "<Types/>",
                "_rels/.rels": "<Relationships/>",
                "word/_rels/document.xml.rels": "<Relationships/>",
            }
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_out:
            output_path = temp_out.name

        try:
            stats = strip_docx_problem_controls(input_path, output_path)
            self.assertTrue(stats["modified"])
            self.assertEqual(stats["removed_sdt"], 1)
            self.assertEqual(stats["removed_fldSimple"], 1)

            xml = self._read_part(output_path, "word/document.xml")
            root = ET.fromstring(xml)
            local_names = {el.tag.rsplit("}", 1)[-1] for el in root.iter()}
            self.assertNotIn("sdt", local_names)
            self.assertNotIn("MERGEFIELD ClientName", xml)
            self.assertIn("fldSimple", xml)
            self.assertIn('instr=" PAGE "', xml)
        finally:
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)


class TestGetJinjaTemplateValidation(unittest.TestCase):
    def test_accepts_valid_template_source(self):
        result = get_jinja_template_validation("Hello {{ user.name.first }}")

        self.assertTrue(result["valid"])
        self.assertEqual(result["errors"], [])

    def test_reports_template_syntax_errors(self):
        result = get_jinja_template_validation("{% if user.name %}Hello")

        self.assertFalse(result["valid"])
        self.assertEqual(result["errors"][0]["code"], "template_syntax_error")

    def test_unknown_filter_is_warning_not_error(self):
        result = get_jinja_template_validation("{{ user.name | custom_filter }}")

        self.assertTrue(result["valid"])
        self.assertEqual(result["errors"], [])
        self.assertEqual(result["warnings"][0]["code"], "template_assertion_error")


class TestDocxTemplateMarkupWarnings(unittest.TestCase):
    def test_warns_when_special_docxtpl_tag_is_missing_space(self):
        document = docx.Document()
        document.add_paragraph("Hello {{p.user_name }}")

        warnings = analyze_docx_template_markup(document)

        self.assertTrue(
            any(
                item["code"] == "docxtpl_special_tag_missing_space" for item in warnings
            )
        )

    def test_warns_when_paragraph_tag_shares_paragraph_with_other_text(self):
        document = docx.Document()
        document.add_paragraph("Prefix {{p user_name }} suffix")

        warnings = analyze_docx_template_markup(document)

        self.assertTrue(
            any(
                item["code"] == "docxtpl_paragraph_tag_with_surrounding_content"
                for item in warnings
            )
        )

    def test_strip_docx_problem_controls_cleans_track_changes_and_hidden_run_properties(
        self,
    ):
        input_path = TestGetJinjaErrors()._build_docx(
            {
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:ins><w:r><w:t>Inserted</w:t></w:r></w:ins>
      <w:del><w:r><w:t>Deleted</w:t></w:r></w:del>
      <w:r><w:rPr><w:vanish/></w:rPr><w:t>Hidden</w:t></w:r>
    </w:p>
  </w:body>
</w:document>""",
                "[Content_Types].xml": "<Types/>",
                "_rels/.rels": "<Relationships/>",
                "word/_rels/document.xml.rels": "<Relationships/>",
            }
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_out:
            output_path = temp_out.name

        try:
            stats = strip_docx_problem_controls(input_path, output_path)
            xml = TestGetJinjaErrors()._read_part(output_path, "word/document.xml")

            self.assertEqual(stats["unwrapped_track_changes"], 1)
            self.assertEqual(stats["removed_track_changes"], 1)
            self.assertEqual(stats["removed_hidden_run_properties"], 1)
            self.assertIn("Inserted", xml)
            self.assertNotIn("Deleted", xml)
            self.assertNotIn("vanish", xml)
        finally:
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_validate_docx_ooxml_schema_uses_mapped_schemas(self):
        docx_path = TestGetJinjaErrors()._build_docx(
            {
                "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>""",
                "_rels/.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>""",
                "word/document.xml": """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p><w:r><w:t>Hello</w:t></w:r></w:p></w:body>
</w:document>""",
                "word/_rels/document.xml.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>""",
            }
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            transitional = os.path.join(temp_dir, "transitional")
            opc = os.path.join(temp_dir, "opc")
            strict = os.path.join(temp_dir, "strict")
            os.makedirs(transitional)
            os.makedirs(opc)
            os.makedirs(strict)
            for path in [
                os.path.join(transitional, "wml.xsd"),
                os.path.join(opc, "opc-contentTypes.xsd"),
                os.path.join(opc, "opc-relationships.xsd"),
            ]:
                with open(path, "w", encoding="utf-8") as handle:
                    handle.write(
                        "<xsd:schema xmlns:xsd='http://www.w3.org/2001/XMLSchema'/>"
                    )

            class _FakeSchema:
                def validate(self, _xml_bytes):
                    return None

                def iter_errors(self, _xml_bytes):
                    return iter([])

            with patch(
                "docassemble.ALDashboard.validate_docx.ensure_ooxml_schema_cache",
                return_value={
                    "transitional": transitional,
                    "strict": strict,
                    "opc": opc,
                },
            ), patch(
                "docassemble.ALDashboard.validate_docx._load_xmlschema",
                return_value=_FakeSchema(),
            ):
                report = validate_docx_ooxml_schema(docx_path)

        try:
            self.assertTrue(report["available"])
            self.assertEqual(report["xml_parse_errors"], [])
            self.assertEqual(report["schema_errors"], [])
            self.assertIn("word/document.xml", report["validated_parts"])
            self.assertIn("[Content_Types].xml", report["validated_parts"])
            self.assertIn("word/_rels/document.xml.rels", report["validated_parts"])
        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)


if __name__ == "__main__":
    unittest.main()
