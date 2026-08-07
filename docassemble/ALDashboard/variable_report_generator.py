import html
import os
import re
from dataclasses import dataclass, field
from typing import Any, Dict, Iterable, List, Mapping, Optional, Set, Tuple

from ruamel.yaml import YAML
from ruamel.yaml.compat import StringIO

try:
    import docx
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    docx = None  # type: ignore[assignment]

# System variables to skip in user-facing intake outputs
SKIP_SYSTEM_VARS: Set[str] = {
    "al_form_type",
    "interview_short_title",
    "al_intro_screen",
    "acknowledged_information_use",
    "interview_metadata",
    "form_approved_for_email_filing",
    "user_started_case",
    "nav",
    "menu_items",
    "AL_ORGANIZATION_TITLE",
    "session_local",
    "device_local",
    "user_local",
    "url_args",
    "cron_hourly",
    "cron_daily",
    "cron_weekly",
    "cron_monthly",
    "incoming_email",
    "role_needed",
    "speak_text",
    "track_location",
    "multi_user",
    "allow_cron",
    "basic_questions_intro_screen",
}

# Internal noise list attributes to filter out
NOISE_LIST_ATTRS: Set[str] = {
    "add_action",
    "revisit",
    "table",
    "there_is_another",
    "there_are_any",
    "target_number",
    "ask_number",
    "gathered",
    "auto_gather",
    "complete",
    "instanceName",
    "elements",
    "parent",
    "attr_name",
    "has_no_address",
    "impounded",
}

# AssemblyLine built-in lists
COMMON_AL_PEOPLE_LISTS: Set[str] = {
    "users",
    "children",
    "plaintiff",
    "defendants",
    "other_parties",
    "witnesses",
    "caregivers",
    "guardians_ad_litem",
    "attorneys",
    "translators",
    "debt_collectors",
    "creditors",
    "spouses",
    "parents",
    "decedents",
    "interested_parties",
    "guardians",
    "adoptees",
}

DEFAULT_INDIVIDUAL_ATTRIBUTES: List[Tuple[str, str, str]] = [
    ("name", "text", "Full Name"),
    ("address", "text", "Address"),
    ("birthdate", "date", "Birthdate"),
    ("email", "email", "Email Address"),
    ("phone_number", "text", "Phone Number"),
]


@dataclass
class AttributeSpec:
    name: str
    datatype: str = "text"
    description: str = ""


@dataclass
class VariableSpec:
    name: str
    var_type: str = "unknown"
    is_list: bool = False
    label: str = ""
    attributes: List[AttributeSpec] = field(default_factory=list)
    source: str = "yaml"


@dataclass
class FieldSpec:
    var_name: str
    label: str
    datatype: str = "text"
    is_list: bool = False
    list_name: str = ""
    attr_name: str = ""


@dataclass
class QuestionGroupSpec:
    question_id: str
    title: str
    fields: List[FieldSpec] = field(default_factory=list)
    lists: List[VariableSpec] = field(default_factory=list)


def list_variable_report_playground_projects() -> List[str]:
    from .interview_linter import list_playground_projects

    return list_playground_projects()


def list_variable_report_playground_yaml_files(
    project: str = "default",
) -> List[Dict[str, str]]:
    from .yaml_formatter import list_formatter_playground_yaml_files

    return list_formatter_playground_yaml_files(project)


def _get_playground_storage(project: str, section: str = "playground") -> Tuple[Any, str]:
    from .docassemble_compat import SavedFile, directory_for
    from .interview_linter import _resolve_current_user_id

    current_user_id = _resolve_current_user_id()
    if current_user_id is None:
        raise ValueError("Could not determine current user for playground access.")

    playground_area = SavedFile(current_user_id, fix=True, section=section)
    project_root = directory_for(playground_area, project)
    if not project_root:
        raise ValueError("Could not locate selected playground project.")
    os.makedirs(project_root, exist_ok=True)
    return playground_area, os.path.realpath(project_root)


def _load_yaml_documents(yaml_text: str) -> List[Dict[str, Any]]:
    yaml = YAML(typ="safe", pure=True)
    docs: List[Dict[str, Any]] = []

    # First attempt full load_all
    try:
        for doc in yaml.load_all(StringIO(yaml_text)):
            if isinstance(doc, dict):
                docs.append(doc)
        return docs
    except Exception:
        pass

    # Resilient fallback: split by document separators and parse blocks individually
    raw_blocks = re.split(r"^---\s*$", str(yaml_text or ""), flags=re.MULTILINE)
    for block in raw_blocks:
        if not block.strip():
            continue
        try:
            doc = yaml.load(StringIO(block))
            if isinstance(doc, dict):
                docs.append(doc)
        except Exception:
            try:
                # Replace tabs with spaces for YAML tab indentation syntax errors
                cleaned_block = block.replace("\t", "    ")
                doc = yaml.load(StringIO(cleaned_block))
                if isinstance(doc, dict):
                    docs.append(doc)
            except Exception:
                pass

    return docs


def _load_assemblyline_baseline() -> Dict[str, VariableSpec]:
    baseline_specs: Dict[str, VariableSpec] = {}

    for list_name in COMMON_AL_PEOPLE_LISTS:
        attrs = [
            AttributeSpec(name=attr_name, datatype=dtype, description=desc)
            for attr_name, dtype, desc in DEFAULT_INDIVIDUAL_ATTRIBUTES
        ]
        baseline_specs[list_name] = VariableSpec(
            name=list_name,
            var_type="ALPeopleList",
            is_list=True,
            label=list_name.replace("_", " ").title(),
            attributes=attrs,
            source="inferred_assemblyline",
        )

    # Load ql_baseline.yml if present
    baseline_path = os.path.expanduser(
        "~/docassemble-AssemblyLine/docassemble/AssemblyLine/data/questions/ql_baseline.yml"
    )
    if os.path.exists(baseline_path):
        try:
            with open(baseline_path, "r", encoding="utf-8") as f:
                docs = _load_yaml_documents(f.read())
            for doc in docs:
                objects = doc.get("objects")
                if isinstance(objects, list):
                    for item in objects:
                        if isinstance(item, Mapping):
                            for var_name, class_info in item.items():
                                name_str = str(var_name).strip()
                                class_str = str(class_info).strip()
                                if name_str in SKIP_SYSTEM_VARS:
                                    continue
                                is_list = "ALPeopleList" in class_str or "DAList" in class_str or "List" in class_str
                                if is_list and name_str not in baseline_specs:
                                    attrs = (
                                        [
                                            AttributeSpec(name=a[0], datatype=a[1], description=a[2])
                                            for a in DEFAULT_INDIVIDUAL_ATTRIBUTES
                                        ]
                                        if "ALPeopleList" in class_str
                                        else [AttributeSpec(name="item", datatype="text", description="Item")]
                                    )
                                    baseline_specs[name_str] = VariableSpec(
                                        name=name_str,
                                        var_type=class_str,
                                        is_list=True,
                                        label=name_str.replace("_", " ").title(),
                                        attributes=attrs,
                                        source="inferred_assemblyline",
                                    )
        except Exception:
            pass

    return baseline_specs


def _clean_var_name(name: str) -> str:
    cleaned = re.sub(r"\[.*?\]", "", str(name or "")).strip()
    return cleaned


def _clean_title(raw: str) -> str:
    title = str(raw or "").strip()
    title = re.sub(r"\$\{.*?\}", "", title).strip()
    title = title.strip("# \t\r\n:")
    return title or "Interview Screen"


def extract_interview_questions_and_variables(
    yaml_texts: List[str],
    infer_assemblyline: bool = True,
) -> Tuple[List[QuestionGroupSpec], Dict[str, VariableSpec]]:
    baseline_al = _load_assemblyline_baseline() if infer_assemblyline else {}
    variables: Dict[str, VariableSpec] = {}
    groups: List[QuestionGroupSpec] = []

    list_attributes_found: Dict[str, Set[Tuple[str, str]]] = {}
    referenced_var_names: Set[str] = set()

    for yaml_text in yaml_texts:
        docs = _load_yaml_documents(yaml_text)
        for doc in docs:
            # 1. Parse objects block
            objects = doc.get("objects")
            if isinstance(objects, list):
                for obj_entry in objects:
                    if isinstance(obj_entry, Mapping):
                        for raw_var, class_info in obj_entry.items():
                            var_name = _clean_var_name(str(raw_var))
                            if var_name in SKIP_SYSTEM_VARS:
                                continue
                            class_str = str(class_info)
                            is_list = any(
                                k in class_str
                                for k in ("ALPeopleList", "DAList", "ALList", "list", "DADict")
                            )
                            if var_name not in variables:
                                variables[var_name] = VariableSpec(
                                    name=var_name,
                                    var_type=class_str,
                                    is_list=is_list,
                                    label=var_name.replace("_", " ").title(),
                                    source="yaml_objects",
                                )
                            else:
                                variables[var_name].var_type = class_str
                                if is_list:
                                    variables[var_name].is_list = True

            # 2. Parse question screen
            if "question" in doc or "fields" in doc or "id" in doc:
                q_title = _clean_title(doc.get("question") or doc.get("id") or "Interview Screen")
                q_id = str(doc.get("id") or "screen").strip()
                group = QuestionGroupSpec(question_id=q_id, title=q_title)

                fields = doc.get("fields")
                if isinstance(fields, list):
                    for f in fields:
                        field_var_name: Optional[str] = None
                        label = ""
                        datatype = "text"
                        if isinstance(f, str):
                            field_var_name = f.strip()
                        elif isinstance(f, Mapping):
                            if "note" in f and "field" not in f:
                                continue
                            datatype = str(f.get("datatype") or "text").strip().lower()
                            if "field" in f and isinstance(f["field"], str):
                                field_var_name = f["field"].strip()
                                label = str(f.get("label") or f.get("note") or "").strip()
                            else:
                                reserved_keys = {
                                    "datatype", "choices", "help", "hint", "note", "label",
                                    "required", "default", "disable_others", "input type",
                                    "show if", "hide if", "rows", "code", "validation code",
                                    "address autocomplete", "js_show_if", "grid", "buttons"
                                }
                                for k, v in f.items():
                                    if str(k).strip().lower() in reserved_keys:
                                        continue
                                    if isinstance(v, str) and re.match(r"^[a-zA-Z_][a-zA-Z0-9_\[\]\.]*$", v.strip()):
                                        field_var_name = v.strip()
                                        label = str(k).strip()
                                        break
                                    elif isinstance(k, str) and re.match(r"^[a-zA-Z_][a-zA-Z0-9_\[\]\.]*$", str(k).strip()) and (v is None or isinstance(v, (str, bool, int, float))):
                                        field_var_name = str(k).strip()
                                        if isinstance(v, str):
                                            label = v.strip()
                                        break

                        if field_var_name and field_var_name not in SKIP_SYSTEM_VARS:
                            list_match = re.match(r"^([a-zA-Z_][a-zA-Z0-9_]*)(?:\[.*?\]|\.(.+))?", field_var_name)
                            if list_match:
                                base_list = list_match.group(1)
                                attr = list_match.group(2) or ""
                                has_index_or_dot = "[" in field_var_name or "." in field_var_name

                                if has_index_or_dot:
                                    if attr and attr not in NOISE_LIST_ATTRS:
                                        referenced_var_names.add(base_list)
                                        if base_list not in list_attributes_found:
                                            list_attributes_found[base_list] = set()
                                        list_attributes_found[base_list].add((attr, datatype))

                                    if base_list not in variables:
                                        variables[base_list] = VariableSpec(
                                            name=base_list,
                                            var_type="DAList",
                                            is_list=True,
                                            label=base_list.replace("_", " ").title(),
                                            source="yaml_fields",
                                        )
                                    else:
                                        variables[base_list].is_list = True
                                else:
                                    referenced_var_names.add(field_var_name)
                                    if field_var_name not in variables:
                                        variables[field_var_name] = VariableSpec(
                                            name=field_var_name,
                                            var_type=datatype,
                                            is_list=False,
                                            label=label or field_var_name.replace("_", " ").title(),
                                            source="yaml_fields",
                                        )
                                    elif label and not variables[field_var_name].label:
                                        variables[field_var_name].label = label

                                    group.fields.append(
                                        FieldSpec(
                                            var_name=field_var_name,
                                            label=label or field_var_name.replace("_", " ").title(),
                                            datatype=datatype,
                                        )
                                    )

                if group.fields:
                    groups.append(group)

            # 3. Find references in templates/code
            doc_str = str(doc)
            for var_match in re.finditer(r"\b([a-zA-Z_][a-zA-Z0-9_]*)(?:\[.*?\]|\.([a-zA-Z0-9_\.]+))?\b", doc_str):
                b_name = var_match.group(1)
                b_attr = var_match.group(2)
                if b_name not in SKIP_SYSTEM_VARS and b_name not in ("True", "False", "None", "len", "str", "int", "dict", "list", "and", "or", "not", "if", "else", "endif", "for", "in", "endfor", "include", "fields", "question", "subquestion", "id", "mandatory", "code"):
                    referenced_var_names.add(b_name)
                    if b_attr and b_attr not in NOISE_LIST_ATTRS:
                        if b_name not in list_attributes_found:
                            list_attributes_found[b_name] = set()
                        list_attributes_found[b_name].add((b_attr, "text"))

    # Infer built-ins for referenced AL objects
    if infer_assemblyline:
        for ref in referenced_var_names:
            if ref in baseline_al and ref not in SKIP_SYSTEM_VARS:
                al_spec = baseline_al[ref]
                if ref not in variables:
                    variables[ref] = VariableSpec(
                        name=al_spec.name,
                        var_type=al_spec.var_type,
                        is_list=al_spec.is_list,
                        label=al_spec.label,
                        attributes=list(al_spec.attributes),
                        source="inferred_assemblyline",
                    )
                else:
                    if variables[ref].var_type in ("unknown", "text", "DAList") and al_spec.var_type != "unknown":
                        variables[ref].var_type = al_spec.var_type
                    if al_spec.is_list:
                        variables[ref].is_list = True
                    if not variables[ref].attributes and al_spec.attributes:
                        variables[ref].attributes = list(al_spec.attributes)

    # Filter out internal list attributes and assign clean attributes
    for var_name, spec in variables.items():
        if spec.is_list:
            filtered_attrs = [a for a in spec.attributes if a.name not in NOISE_LIST_ATTRS]
            existing_attrs = {a.name for a in filtered_attrs}

            if var_name in list_attributes_found:
                for attr_name, dtype in sorted(list_attributes_found[var_name]):
                    if attr_name not in NOISE_LIST_ATTRS and attr_name not in existing_attrs:
                        filtered_attrs.append(
                            AttributeSpec(
                                name=attr_name,
                                datatype=dtype,
                                description=attr_name.replace("_", " ").replace(".", " ").title(),
                            )
                        )
                        existing_attrs.add(attr_name)

            if not filtered_attrs:
                if "People" in spec.var_type or var_name in COMMON_AL_PEOPLE_LISTS:
                    filtered_attrs = [
                        AttributeSpec(name=a[0], datatype=a[1], description=a[2])
                        for a in DEFAULT_INDIVIDUAL_ATTRIBUTES
                    ]
                else:
                    filtered_attrs = [AttributeSpec(name="item", datatype="text", description="Item")]

            spec.attributes = filtered_attrs

    # Attach list variables to groups or to an Objects & Parties group
    assigned_lists: Set[str] = set()
    for g in groups:
        for var_name, spec in variables.items():
            if spec.is_list and var_name not in assigned_lists:
                # Check if this list was referenced in this group's screen
                if any(f.var_name.startswith(var_name) for f in g.fields) or var_name in g.title.lower():
                    g.lists.append(spec)
                    assigned_lists.add(var_name)

    unassigned_lists = [v for k, v in variables.items() if v.is_list and k not in assigned_lists]
    if unassigned_lists:
        groups.append(
            QuestionGroupSpec(
                question_id="parties_objects",
                title="Parties and Lists",
                lists=unassigned_lists,
            )
        )

    return groups, variables


def generate_mako_markdown_report(
    groups: List[QuestionGroupSpec],
    variables: Dict[str, VariableSpec],
    report_title: Optional[str] = None,
    show_variable_names: bool = False,
    show_variable_types: bool = False,
    max_list_cols: int = 4,
) -> str:
    title = report_title or "Interview Document Draft"

    lines: List[str] = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append("---")
    lines.append("")

    for group in groups:
        if not group.fields and not group.lists:
            continue

        lines.append(f"## {group.title}")
        lines.append("")

        if group.fields:
            lines.append("| Field | Value |")
            lines.append("| --- | --- |")
            for fspec in group.fields:
                disp_label = fspec.label
                if show_variable_names:
                    disp_label += f" (`{fspec.var_name}`)"
                if show_variable_types:
                    disp_label += f" [{fspec.datatype}]"

                if fspec.datatype in ("yesno", "boolean"):
                    val_expr = f"% if showifdef('{fspec.var_name}'): Yes % else: No % endif"
                else:
                    val_expr = f"${{ showifdef('{fspec.var_name}') }}"
                lines.append(f"| {disp_label} | {val_expr} |")
            lines.append("")

        if group.lists:
            for lvar in group.lists:
                disp_title = lvar.label or lvar.name
                if show_variable_names:
                    disp_title += f" (`{lvar.name}`)"
                if show_variable_types:
                    disp_title += f" [{lvar.var_type}]"

                lines.append(f"### {disp_title}")
                lines.append("")

                attrs = [a for a in lvar.attributes if a.name not in NOISE_LIST_ATTRS] or [AttributeSpec(name="item", datatype="text", description="Item")]

                if len(attrs) <= max_list_cols:
                    # Horizontal table format
                    header_row = "| # | " + " | ".join([a.description or a.name for a in attrs]) + " |"
                    sep_row = "| --- | " + " | ".join(["---"] * len(attrs)) + " |"
                    lines.append(header_row)
                    lines.append(sep_row)

                    lines.append(f"% if showifdef('{lvar.name}'):")
                    lines.append(f"% for i, item in enumerate({lvar.name}):")

                    item_cells = []
                    for a in attrs:
                        if a.name == "item":
                            item_cells.append(f"${{ showifdef(f'{lvar.name}[{{i}}]') }}")
                        else:
                            item_cells.append(f"${{ showifdef(item.attr_name('{a.name}')) }}")

                    lines.append("| ${ i + 1 } | " + " | ".join(item_cells) + " |")
                    lines.append("% endfor")
                    lines.append("% else:")

                    sample_cells = [f"Sample {a.description or a.name}" for a in attrs]
                    lines.append("| 1 | " + " | ".join(sample_cells) + " |")
                    lines.append("% endif")
                    lines.append("")
                else:
                    # Vertical person-by-person format for wide lists
                    lines.append(f"% if showifdef('{lvar.name}'):")
                    lines.append(f"% for i, item in enumerate({lvar.name}):")
                    lines.append(f"#### Person ${{ i + 1 }}")
                    for a in attrs:
                        desc = a.description or a.name
                        if a.name == "item":
                            expr = f"${{ showifdef(f'{lvar.name}[{{i}}]') }}"
                        else:
                            expr = f"${{ showifdef(item.attr_name('{a.name}')) }}"
                        lines.append(f"- **{desc}:** {expr}")
                    lines.append("")
                    lines.append("% endfor")
                    lines.append("% else:")
                    lines.append("*No items listed.*")
                    lines.append("% endif")
                    lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def generate_docx_report(
    groups: List[QuestionGroupSpec],
    variables: Dict[str, VariableSpec],
    report_title: Optional[str] = None,
    output_path: Optional[str] = None,
    show_variable_names: bool = False,
    show_variable_types: bool = False,
    max_list_cols: int = 4,
) -> str:
    if docx is None:
        raise RuntimeError("python-docx is not installed in the python environment.")

    title = report_title or "Interview Document Draft"
    doc = docx.Document()

    p_title = doc.add_heading(title, level=1)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for group in groups:
        if not group.fields and not group.lists:
            continue

        doc.add_heading(group.title, level=2)

        if group.fields:
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Field"
            hdr_cells[1].text = "Value"

            for fspec in group.fields:
                row_cells = table.add_row().cells
                disp_label = fspec.label
                if show_variable_names:
                    disp_label += f" ({fspec.var_name})"
                if show_variable_types:
                    disp_label += f" [{fspec.datatype}]"

                row_cells[0].text = disp_label
                if fspec.datatype in ("yesno", "boolean"):
                    row_cells[1].text = f"{{% if showifdef('{fspec.var_name}') %}}Yes{{% else %}}No{{% endif %}}"
                else:
                    row_cells[1].text = f"{{{{ showifdef('{fspec.var_name}') }}}}"

            doc.add_paragraph()

        if group.lists:
            for lvar in group.lists:
                disp_title = lvar.label or lvar.name
                if show_variable_names:
                    disp_title += f" ({lvar.name})"
                if show_variable_types:
                    disp_title += f" [{lvar.var_type}]"

                doc.add_heading(disp_title, level=3)

                attrs = [a for a in lvar.attributes if a.name not in NOISE_LIST_ATTRS] or [AttributeSpec(name="item", datatype="text", description="Item")]

                if len(attrs) <= max_list_cols:
                    # Horizontal table format with standalone {%tr for %} and {%tr endfor %} rows
                    cols_count = len(attrs) + 1
                    table = doc.add_table(rows=4, cols=cols_count)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Row 0: Headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "#"
                    for i, a in enumerate(attrs, start=1):
                        hdr_cells[i].text = a.description or a.name

                    # Row 1: Standalone Loop Start Row
                    loop_start_cells = table.rows[1].cells
                    loop_start_cells[0].text = f"{{%tr for item in showifdef('{lvar.name}') %}}"

                    # Row 2: Data Row
                    data_cells = table.rows[2].cells
                    data_cells[0].text = "{{ loop.index }}"
                    for i, a in enumerate(attrs, start=1):
                        if a.name == "item":
                            tag = "{{ item }}"
                        else:
                            tag = f"{{{{ showifdef(item.attr_name('{a.name}')) }}}}"
                        data_cells[i].text = tag

                    # Row 3: Standalone Loop End Row
                    loop_end_cells = table.rows[3].cells
                    loop_end_cells[0].text = "{%tr endfor %}"

                    doc.add_paragraph()
                else:
                    # Vertical person-by-person format for wide lists
                    p_loop_start = doc.add_paragraph(f"{{%p for item in showifdef('{lvar.name}') %}}")
                    doc.add_heading("Person {{ loop.index }}", level=4)

                    table = doc.add_table(rows=0, cols=2)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for a in attrs:
                        row_cells = table.add_row().cells
                        row_cells[0].text = a.description or a.name
                        if a.name == "item":
                            row_cells[1].text = "{{ item }}"
                        else:
                            row_cells[1].text = f"{{{{ showifdef(item.attr_name('{a.name}')) }}}}"

                    doc.add_paragraph(f"{{%p endfor %}}")
                    doc.add_paragraph()

    if output_path:
        doc.save(output_path)
        return output_path
    else:
        temp_dir = os.path.join(os.getcwd(), "tmp")
        os.makedirs(temp_dir, exist_ok=True)
        out_file = os.path.join(temp_dir, "variable_report_draft.docx")
        doc.save(out_file)
        return out_file


def generate_variable_report(
    yaml_texts: List[str],
    report_title: Optional[str] = None,
    infer_assemblyline: bool = True,
    show_variable_names: bool = False,
    show_variable_types: bool = False,
    max_list_cols: int = 4,
    output_docx_path: Optional[str] = None,
) -> Dict[str, Any]:
    groups, variables = extract_interview_questions_and_variables(
        yaml_texts, infer_assemblyline=infer_assemblyline
    )
    mako_markdown = generate_mako_markdown_report(
        groups,
        variables,
        report_title=report_title,
        show_variable_names=show_variable_names,
        show_variable_types=show_variable_types,
        max_list_cols=max_list_cols,
    )
    docx_file_path = generate_docx_report(
        groups,
        variables,
        report_title=report_title,
        output_path=output_docx_path,
        show_variable_names=show_variable_names,
        show_variable_types=show_variable_types,
        max_list_cols=max_list_cols,
    )

    list_count = sum(1 for v in variables.values() if v.is_list)
    scalar_count = sum(1 for v in variables.values() if not v.is_list)

    return {
        "variables": variables,
        "groups": groups,
        "mako_markdown": mako_markdown,
        "docx_path": docx_file_path,
        "variables_count": len(variables),
        "list_count": list_count,
        "scalar_count": scalar_count,
    }


def save_variable_report_to_playground(
    mako_markdown: Optional[str] = None,
    docx_path: Optional[str] = None,
    *,
    selected_playground_project: str = "default",
    save_format_choice: str = "both",
    md_filename: str = "interview_document_draft.md",
    docx_filename: str = "interview_document_draft.docx",
    output_filename: Optional[str] = None,  # Backward compatibility
) -> Dict[str, Any]:
    project = str(selected_playground_project or "default")
    if output_filename and not md_filename:
        md_filename = output_filename

    result: Dict[str, Any] = {
        "saved_md": False,
        "saved_docx": False,
        "saved_files": [],
        "saved": False,
        "error": None,
        "errors": [],
    }

    try:
        # 1. Save Markdown (.md) to playground questions area (section="playground")
        if save_format_choice in ("both", "markdown") and mako_markdown:
            try:
                playground_area, _ = _get_playground_storage(project, section="playground")
                playground_area.write_content(
                    str(mako_markdown or ""),
                    filename=md_filename,
                    project=project,
                    save=False,
                )
                playground_area.finalize()
                result["saved_md"] = True
                result["saved_files"].append(md_filename)
            except Exception as err:
                result["errors"].append(f"Failed to save Markdown file '{md_filename}': {err}")

        # 2. Save DOCX (.docx) to playground templates area (section="playgroundtemplate")
        if save_format_choice in ("both", "docx") and docx_path and os.path.isfile(docx_path):
            try:
                template_area, the_directory = _get_playground_storage(project, section="playgroundtemplate")
                filepath = os.path.join(the_directory, docx_filename)
                with open(docx_path, "rb") as src, open(filepath, "wb") as dst:
                    dst.write(src.read())
                template_area.finalize()
                result["saved_docx"] = True
                result["saved_files"].append(docx_filename)
            except Exception as err:
                result["errors"].append(f"Failed to save DOCX file '{docx_filename}': {err}")

        result["saved"] = result["saved_md"] or result["saved_docx"]
        if result["errors"]:
            result["error"] = "; ".join(result["errors"])

    except Exception as err:
        result["error"] = str(err)

    return result


def extract_interview_metadata_info(
    yaml_texts: List[str],
    primary_filename: Optional[str] = None,
) -> Dict[str, str]:
    short_title = None
    for yaml_text in yaml_texts:
        if not yaml_text:
            continue
        # 1. Match code block assignment: interview_short_title = "..."
        m = re.search(r"interview_short_title\s*=\s*[\"']([^\"']+)[\"']", yaml_text)
        if m:
            short_title = m.group(1).strip()
            break
        # 2. Match YAML key: interview_short_title: ...
        m2 = re.search(r"^interview_short_title:\s*([^\n#]+)", yaml_text, re.MULTILINE)
        if m2:
            short_title = m2.group(1).strip().strip("\"'")
            break

    def _slugify(text: str) -> str:
        s = re.sub(r"[^\w\s-]", "", str(text or "")).strip().lower()
        return re.sub(r"[-\s]+", "_", s)

    if short_title:
        base_name = _slugify(short_title)
        display_title = f"{short_title} Draft"
    elif primary_filename:
        clean_file = primary_filename.split("/")[-1].rsplit(".", 1)[0]
        base_name = _slugify(clean_file)
        display_title = f"{clean_file.replace('_', ' ').replace('-', ' ').title()} Draft"
    else:
        base_name = "interview_document"
        display_title = "Interview Document Draft"

    return {
        "short_title": short_title or "",
        "display_title": display_title,
        "md_filename": f"{base_name}_draft.md",
        "docx_filename": f"{base_name}_draft.docx",
    }


def get_playground_yaml_texts(
    selected_playground_project: str,
    selected_filenames: Iterable[str],
) -> List[str]:
    project = str(selected_playground_project or "default")
    filenames = [str(f) for f in selected_filenames]
    if not filenames:
        return []

    _, project_root = _get_playground_storage(project, section="playground")
    yaml_texts: List[str] = []
    for filename in filenames:
        source_path = os.path.realpath(os.path.join(project_root, filename))
        if not source_path.startswith(project_root + os.sep):
            raise ValueError("Refusing to read files outside selected playground project.")
        with open(source_path, "r", encoding="utf-8") as f:
            yaml_texts.append(f.read())
    return yaml_texts


def generate_and_save_playground_variable_report(
    selected_filenames: Iterable[str],
    *,
    selected_playground_project: str,
    save_to_playground: bool = False,
    save_playground_project: str = "",
    save_format_choice: str = "both",
    md_filename: str = "interview_document_draft.md",
    docx_filename: str = "interview_document_draft.docx",
    output_filename: Optional[str] = None,
    infer_assemblyline: bool = True,
    show_variable_names: bool = False,
    show_variable_types: bool = False,
    max_list_cols: int = 4,
    report_title: Optional[str] = None,
) -> Dict[str, Any]:
    project = str(selected_playground_project or "default")
    filenames = [str(f) for f in selected_filenames]

    if not filenames:
        raise ValueError("Select at least one YAML file.")

    yaml_texts = get_playground_yaml_texts(project, filenames)

    res = generate_variable_report(
        yaml_texts,
        report_title=report_title,
        infer_assemblyline=infer_assemblyline,
        show_variable_names=show_variable_names,
        show_variable_types=show_variable_types,
        max_list_cols=max_list_cols,
    )

    if save_to_playground:
        save_res = save_variable_report_to_playground(
            mako_markdown=res["mako_markdown"],
            docx_path=res["docx_path"],
            selected_playground_project=save_playground_project or project,
            save_format_choice=save_format_choice,
            md_filename=output_filename or md_filename,
            docx_filename=docx_filename,
        )
        res["saved"] = save_res["saved"]
        res["saved_files"] = save_res["saved_files"]
        res["save_error"] = save_res["error"]

    return res
